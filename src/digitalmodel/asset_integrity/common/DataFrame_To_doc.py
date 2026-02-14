import lxml
import pandas as pd
from docx import Document

"""
References:
https://python-docx.readthedocs.io/en/latest/
https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document
Error creating new document. So always need a root document.
https://stackoverflow.com/questions/43578811/python-docx-opc-exceptions-packagenotfounderror-package-not-found-when-opening
"""
def DataFrame_To_doc(df):
    # open an existing document
    doc = open('test.docx')
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    paraObj1 = doc.add_paragraph('This is a second paragraph.')
    paraObj2 = doc.add_paragraph('This is a yet another paragraph.')

    t = doc.add_table(df.shape[0]+1, df.shape[1])

    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    # save the doc
    doc.save('./test1.docx')

if __name__ == '__main__':
    import numpy as np

    df = pd.DataFrame(np.random.random([3, 3]), 
        columns=['A', 'B', 'C'], index=['first', 'second', 'third'])
    
    DataFrame_To_doc(df)

