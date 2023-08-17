import datetime
import yaml
import operator
from functools import reduce

from digitalmodel.common.utilities import is_file_valid_func


class ReadFromExcel():

    def from_xlsx(self, cfg, file_index=0):
        cfg_template_single_xlsx = {'io': 'data_manager/data/sample_data.xlsx'}
        import pandas as pd

        if cfg.__contains__('files'):
            cfg_temp = cfg['files']['from_xlsx'][file_index]
        else:
            cfg_temp = cfg

        if cfg_temp.__contains__('sheet_name'):
            result = pd.read_excel(cfg_temp['io'],
                                   sheet_name=cfg_temp['sheet_name'],
                                   skiprows=cfg_temp['skiprows'],
                                   skipfooter=cfg_temp['skipfooter'],
                                   index_col=cfg_temp['index_col'])
        else:
            xls = pd.ExcelFile(cfg_temp['io'])
            result = {}
            for sheet_name in xls.sheet_names:
                result[sheet_name] = xls.parse(sheet_name)

        return result


class ReadData():

    def df_filter_by_column_values(self, cfg, df, file_index=0):
        cfg_template = {}
        if cfg.__contains__('files'):
            filter_dict = cfg['files']['from_xlsx'][file_index]['filter']
        else:
            cfg_template = {'filter': None}
            filter_dict = cfg['filter']
        if filter_dict is not None:
            for filter_index in range(0, len(filter_dict)):
                column = filter_dict[filter_index]['column']
                value = filter_dict[filter_index]['value']
                df = df[df[column] == value].copy()
                df.reset_index(inplace=True, drop=True)

        return df

    def from_df_delete_unwanted_columns(self, df, dropped_column_array):
        sample_input_dropped_columns = [3, 5]
        df.drop(df.columns[dropped_column_array], axis=1, inplace=True)

        return df

    def from_xlsx_get_line_number_containing_keyword(self, cfg):
        """
        https://stackoverflow.com/questions/38056233/python-search-excel-sheet-for-specific-strings-in-a-column-and-extract-those-ro/38056526
        https://www.tutorialspoint.com/How-to-check-if-multiple-strings-exist-in-another-string-in-Python
        """
        # import xlrd
        from openpyxl import load_workbook
        wb = load_workbook(cfg['io'])

        sh = wb[cfg['sheet_name']]
        keyword_row_number = self.from_xlsx_get_WorkSheetRowNumberWithText(
            sh, cfg["key_words"])
        return keyword_row_number

    def xlsx_to_df_by_keyword_search(self, cfg):
        import pandas as pd
        cfg.update(cfg['start_row'])
        row_number_with_keyword = self.from_xlsx_get_line_number_containing_keyword(
            cfg)
        start_row = cfg['start_row']['transform']['scale'] * row_number_with_keyword + \
                    cfg['start_row']['transform']['shift']
        cfg.update(cfg['end_row'])
        row_number_with_keyword = self.from_xlsx_get_line_number_containing_keyword(
            cfg)
        end_row = cfg['end_row']['transform']['scale'] * row_number_with_keyword + \
                  cfg['end_row']['transform']['shift']
        number_of_rows = end_row - start_row + 1
        df = pd.read_excel(io=cfg['io'],
                           sheet_name=cfg['sheet_name'],
                           skiprows=start_row - 1,
                           nrows=number_of_rows)
        if cfg.__contains__('column'):
            if cfg['column'].__contains__('drop_unwanted_columns'):
                dropped_column_array = [
                    column for column in range(len(cfg['column']['names']),
                                               len(df.columns))
                ]
                df = self.from_df_delete_unwanted_columns(
                    df, dropped_column_array)
            if cfg['column'].__contains__('names'):
                df.columns = cfg['column']['names']

        return df

    def superseded_xlsx_to_df_by_keyword_search(self, data):
        """
        https://stackoverflow.com/questions/38056233/python-search-excel-sheet-for-specific-strings-in-a-column-and-extract-those-ro/38056526
        https://www.tutorialspoint.com/How-to-check-if-multiple-strings-exist-in-another-string-in-Python
        """
        import pandas as pd
        import xlrd
        ReadData = []
        wb = xlrd.open_workbook(data['FileName'])

        sh = wb.sheet_by_name(data['SheetName'])
        KeyWordRowNumber = self.from_xlsx_get_WorkSheetRowNumberWithText(
            sh, data["KeyWords"])

        if KeyWordRowNumber == None:
            raise Exception("Error in keyword provided for search criteria")

        StartRowNumber = KeyWordRowNumber + data['RowsToSkip']
        EndRowNumber = KeyWordRowNumber + data['RowsToSkip'] + data['RowsToRead']
        if EndRowNumber > sh.nrows:
            EndRowNumber = sh.nrows
        for rownum in range(StartRowNumber, EndRowNumber):
            ReadData.append((sh.row_values(rownum)))

        df = pd.DataFrame(ReadData)

        # Assign columns
        if data['Columns'] != None:
            if data['Columns'] == '1st Row':
                df.columns = df.iloc[0]
                df.drop([0], inplace=True)
                df.reset_index(inplace=True)
            elif len(data['Columns']) <= len(df.columns):
                AdditionalColumns = list(
                    range(len(data['Columns']), len(df.columns)))
                df.columns = data['Columns'] + AdditionalColumns
            else:
                df.columns = data['Columns'][0:len(df.columns)]

        return df

    def from_xlsx_get_WorkSheetRowNumberWithText(self, sh, KeyWordArray):
        """
        Objective: To obtain the row number of the worksheet with specified keyword(s)
        """
        rownum = None
        for rownum in range(0, sh.nrows):
            if any(keyword in sh.row_values(rownum)
                   for keyword in KeyWordArray):
                return rownum
                break

    def read_yml_file(self, data):
        import yaml
        with open(data['io'], 'r') as ymlfile:
            data_as_dictionary = yaml.load(ymlfile, Loader=yaml.Loader)

        return data_as_dictionary

    def extract_from_dictionary(self, data_dictionary, map_list):
        return reduce(operator.getitem, map_list, data_dictionary)

    def key_chain(self, data, *args, default=None):
        for key in args:
            if isinstance(data, dict):
                data = data.get(key, default)
            elif isinstance(data, (list, tuple)) and isinstance(key, int):
                try:
                    data = data[key]
                except IndexError:
                    return default
            else:
                return default
        return data

    def from_pdf(self, cfg, file_index=0):
        if cfg['files']['from_pdf'][file_index].__contains__('package'):
            if cfg['files']['from_pdf'][file_index]['package'] in [
                    None, 'tabula'
            ]:
                df = self.from_pdf_tabula(cfg, file_index)
            elif cfg['files']['from_pdf'][file_index]['package'] == 'camelot':
                df = self.from_pdf_camelot(cfg, file_index)
        else:
            df = self.from_pdf_tabula(cfg, file_index)

        return df

    def from_pdf_tabula(self, cfg, file_index=0):
        import tabula
        df = tabula.read_pdf(cfg['files']['from_pdf'][file_index]['io'],
                             pages=cfg['files']['from_pdf'][file_index]['page'],
                             multiple_tables=True)
        return df

    def from_pdf_camelot(self, cfg, file_index=0):
        import camelot
        df = camelot.read_pdf(
            cfg['files']['from_pdf'][file_index]['io'],
            pages=cfg['files']['from_pdf'][file_index]['page'],
            suppress_stdout=False)
        print(df)
        return df

    def update_extension_of_filename(self):
        # TODO add facility to update file extension from one to another
        # old_extension = .plt
        # new_extension = .out
        pass

    def from_ascii_file_get_line_number_containing_keywords(self, cfg):
        sample_data_1 = {
            'io':
                'data_manager\\data\\shear7\\lid_02_cp01_2500ft_WT0750_064pcf.out',
            'line': {
                'key_words': ['5. Fundamental natural frequency'],
                'transform': {
                    'scale': 1,
                    'shift': 0
                }
            }
        }

        all_lines_as_strings = self.from_ascii_file_get_lines_as_string_arrays(
            cfg)
        keyword_line_number = self.get_array_rows_containing_keywords(
            all_lines_as_strings, cfg['line']['key_words'])

        if cfg['line'].__contains__('transform'):
            keyword_line_number = keyword_line_number * cfg['line'][
                'transform']['scale'] + cfg['line']['transform']['shift']

        return keyword_line_number

    def get_array_rows_containing_keywords(self, array, key_words):
        for rownum in range(0, len(array)):
            if any(keyword in array[rownum] for keyword in key_words):
                return rownum + 1
                break

    def from_ascii_file_get_value(self, cfg):
        sample_cfg_format_1 = {'io': 'data_manager\data\orcaflex\common.mds'}
        sample_cfg_format_2 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'end_line': None
        }
        sample_cfg_format_3 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'end_line': 10
        }

        from_string = FromString()
        line_text = self.from_ascii_file_get_lines_as_string_arrays(cfg)
        line_text = from_string.remove_next_line_values(line_text)
        cfg_temp = {'text': line_text}
        cfg_temp.update(cfg['filter'])
        result = from_string.get_value_by_delimiter(cfg_temp)

        return result

    def from_ascii_file_get_structured_data_delimited_white_space(self, cfg):
        import pandas as pd
        sample_cfg_format_1 = {'io': 'data_manager\data\orcaflex\common.mds'}
        sample_cfg_format_2 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'DataFrame': True
        }
        sample_cfg_format_2 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'end_line': None,
            'DataFrame': True
        }    # Get only 1 line
        sample_cfg_format_3 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'end_line': 10,
            'DataFrame': True
        }

        all_lines_as_strings = self.from_ascii_file_get_lines_as_string_arrays(
            cfg)
        all_lines_float_objects = []
        for line_index in range(0, len(all_lines_as_strings)):
            if cfg['delimiter'] == 'space':
                line_string_objects = all_lines_as_strings[line_index].split()
            else:
                print("Unknown delimiter")
            all_lines_float_objects.append(
                [float(item) for item in line_string_objects])

        result = []
        for line_float_object_index in range(0,
                                             len(all_lines_float_objects[0])):
            result.append([])

        for line_index in range(0, len(all_lines_float_objects)):
            for line_float_object_index in range(
                    0, len(all_lines_float_objects[0])):
                result[line_float_object_index].append(
                    all_lines_float_objects[line_index]
                    [line_float_object_index])

        if cfg['DataFrame']:
            df = pd.DataFrame(columns=cfg['columns'])
            for column_index in range(0, len(df.columns)):
                df[df.columns[column_index]] = result[column_index]
            result = df

        return result

    def from_ascii_file_get_lines_as_string_arrays(self, cfg):
        sample_cfg_format_1 = {'io': 'data_manager\data\orcaflex\common.mds'}
        sample_cfg_format_2 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'end_line': None
        }
        sample_cfg_format_2 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5
        }    # Get till end of lines
        sample_cfg_format_3 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'start_line': 5,
            'end_line': 10
        }
        sample_cfg_format_3 = {
            'io': 'data_manager\data\orcaflex\common.mds',
            'lines_from_end': 2
        }

        all_lines = []
        if type(cfg['io']) is list:
            from common.ETL_components import ETL_components
            etl_components = ETL_components(cfg=None)
            cfg['io'] = etl_components.convert_text_list_to_combined_text(
                cfg['io'])
        with open(cfg['io']) as f:
            for line in f:
                all_lines.append(line)

        if cfg.__contains__('start_line'):
            start_line = cfg['start_line']
        else:
            start_line = 1

        if cfg.__contains__('end_line'):
            end_line = cfg['end_line']
        else:
            end_line = len(all_lines)

        if cfg.__contains__('lines_from_end'):
            end_line = len(all_lines)
            start_line = end_line - cfg['lines_from_end']

        if end_line == None:
            result = all_lines[start_line - 1]
        else:
            result = all_lines[start_line - 1:end_line]

        return result

    def get_file_list_from_folder(self,
                                  folder_with_file_type,
                                  with_path=True,
                                  with_extension=True):
        folder_with_file_type_example = 'Q:\projects\Mole\log_files\*.log'

        import glob
        import os
        file_list = []
        for file in glob.glob(folder_with_file_type):
            file_list.append(file)
        if with_path:
            if with_extension:
                return file_list
            else:
                file_list = [os.path.splitext(file)[0] for file in file_list]
                return file_list
        else:
            if with_extension:
                file_list = [os.path.basename(file) for file in file_list]
                return file_list
            else:
                file_list = [
                    os.path.splitext(os.path.basename(file))[0]
                    for file in file_list
                ]
                return file_list


class GetData():

    def download_file_from_url(self, cfg):
        import os
        import time
        from pathlib import Path

        import wget
        cfg_temp = {
            'url': 'https://www.data.bsee.gov/Well/Files/APIRawData.zip',
            'download_to': os.path.abspath(Path("../data_manager/data/bsee"))
        }

        url = cfg['url']
        filename = os.path.join(cfg['download_to'] + '/' +
                                os.path.basename(url))

        if os.path.exists(filename):
            os.remove(filename)

        start_time = time.perf_counter()
        print("Dowloading file: {}".format(filename))
        wget.download(url, out=filename)
        end_time = time.perf_counter()
        print("Downloading file: {} .... COMPLETE".format(filename))
        print("Time Taken to download: {} .... COMPLETE".format(
            (end_time - start_time).__round__(3)))
        return {"filename": filename, "result": True}


class FromString():

    def using_regex(self, ref_text, string):
        import re
        result = re.findall(ref_text, string, re.IGNORECASE)
        if len(result) > 0:
            return result[0]

    def convert_fraction_to_float(self, frac_str):
        try:
            return float(frac_str)
        except ValueError:
            num, denom = frac_str.split('/')
            try:
                leading, num = num.split(' ')
                whole = float(leading)
            except ValueError:
                whole = 0
            frac = float(num) / float(denom)
            return whole - frac if whole < 0 else whole + frac

    def remove_strings(self, text, string_array):
        for string_index in range(0, len(string_array)):
            string_to_be_removed = string_array[string_index]
            if text is not None:
                text = self.remove_string(text, string_to_be_removed)
        return text

    def remove_string(self, text, string_to_be_removed):
        new_string = text.replace(string_to_be_removed, "")
        return new_string

    def get_value_by_delimiter(self, cfg):
        cfg_sample_1 = {
            'text': '          97        2001',
            'delimiter': ' ',
            'column': 2,
            'data_type': 'float'
        }
        cfg_sample_2 = {
            'text': '0.017300(Hz)',
            'delimiter': '(',
            'column': 1,
            'data_type': 'float'
        }

        if cfg['delimiter'] == ' ':
            result = cfg['text'].split()[cfg['column'] - 1]
            if cfg['data_type'] == 'float':
                result = float(result)
        else:
            delimiter = cfg['delimiter']
            result = cfg['text'].split(delimiter)[cfg['column'] - 1]
            if cfg['data_type'] == 'float':
                result = float(result)

        return result

    def remove_next_line_values(self, text):
        return (text.replace('\n', ''))


class SaveData():

    def saveDataJson(self, data, fileName):
        import json
        with open(fileName + '.json', 'w') as f:
            json.dump(data, f)

    def saveDataYaml(self, data, fileName, default_flow_style=False):
        import yaml
        if default_flow_style == None:
            with open(fileName + '.yml', 'w') as f:
                yaml.dump(data, f)
        elif default_flow_style == 'NonAlias':
            with open(fileName + '.yml', 'w') as f:
                yaml.dump(data, f, Dumper=noalias_dumper)
        elif default_flow_style == 'ruamel':
            with open(fileName + '.yml', 'w') as f:
                ruamel.yaml.dump(data, f)
        elif default_flow_style == 'round_trip_dump':
            with open(fileName + '.yml', 'w') as f:
                ruamel.yaml.round_trip_dump(data, f)
        else:
            with open(fileName + '.yml', 'w') as f:
                yaml.dump(data, f, default_flow_style=default_flow_style)

    def saveDataFrame(self, df, fileName):
        df.to_csv(fileName + '.csv')

    def save_ascii_file_from_array(self, array, file_name, extension=None):
        with open(file_name + extension, "w") as f:
            for line_index in range(0, len(array)):
                if type(array[line_index]) is list:
                    from common.ETL_components import ETL_components
                    etl_components = ETL_components(cfg=None)
                    array[
                        line_index] = etl_components.convert_text_list_to_combined_text(
                            array[line_index])
                f.write(array[line_index])
                if "\n" not in array[line_index]:
                    f.write('\n')

    def DataFrame_To_xlsx_xlsxwriter(df, data):
        # from openpyxl.styles.borders import Border, Side, BORDER_THIN
        """
        https://xlsxwriter.readthedocs.io/working_with_pandas.html
        """
        writer = pd.ExcelWriter(data['FileName'], engine='xlsxwriter')

        try:
            # WorkSheet = wb.get_sheet_by_name(data['SheetName'])
            WorkSheet = wb[data['SheetName']]
        except:
            wb.create_sheet(data['SheetName'])
            WorkSheet = wb[data['SheetName']]

        df.to_excel(writer, data['SheetName'])
        writer.save()

    """
    References:
    https://stackoverflow.com/questions/42370977/how-to-save-a-new-sheet-in-an-existing-excel-file-using-pandas/42371251
    https://stackoverflow.com/questions/36814050/openpyxl-get-sheet-by-name
    """

    def DataFrame_To_xlsx_openpyxl(self, df, data):
        import pandas as pd
        from openpyxl import load_workbook

        try:
            wb = load_workbook(data['FileName'])
            writer = pd.ExcelWriter(data['FileName'], engine='openpyxl')
            writer.wb = wb
        except:
            writer = pd.ExcelWriter(data['FileName'])

        try:
            # WorkSheet = wb.get_sheet_by_name(data['SheetName'])
            WorkSheet = wb[data['SheetName']]
        except:
            wb.create_sheet(data['SheetName'])
            WorkSheet = wb[data['SheetName']]

            #  For xlsxwriter
            # WorkSheet = wb.add_worksheet(data['SheetName'])

        df.to_excel(writer, data['SheetName'])
        writer.save()

    def DataFrameArray_To_xlsx_openpyxl(self, dfArray, data):
        import pandas as pd
        print("Opening new workbook")
        writer = pd.ExcelWriter(data['FileName'], engine='openpyxl')

        for dfIndex in range(0, len(dfArray)):
            dfArray[dfIndex].to_excel(writer, data['SheetNames'][dfIndex])
            """
            References:
            https://stackoverflow.com/questions/24917201/applying-borders-to-a-cell-in-openpyxl
            """
            # property cell.border should be used instead of cell.style.border
            # if data['thin_border']:
            #     ws = wb.active
            #     thin_border = Border(
            #         left=Side(border_style=BORDER_THIN, color='00000000'),
            #         right=Side(border_style=BORDER_THIN, color='00000000'),
            #         top=Side(border_style=BORDER_THIN, color='00000000'),
            #         bottom=Side(border_style=BORDER_THIN, color='00000000')
            #     )
            #     rows = len(dfArray[dfIndex])+1
            #     columns = len(dfArray[dfIndex].columns)+1
            #     ws.cell(row=rows, column=columns).border = thin_border

        writer.save()

    def df_to_sheet_in_existing_workbook(self, data):
        # Adds a dataframe to an existing Excel
        # Uses Excel program instance Works only on windows machine as it needs Excel program.
        # As of 2020-Feb, the pacakge is very slow and is NOT recommended for use on even simple excel jobs
        import xlwings as xw
        wb = xw.Book(data['template_file_name'])
        ws = wb.sheets[data['sheetname']]
        ws.range('A1').options(index=True).value = data['df']
        wb.save(data['saved_file_name'])
        # app = xw.apps.active
        # app.quit()
        xw.apps.active.quit()

        # Not working loop using Pandas below for reference code
        # import pandas as pd
        # from openpyxl import load_workbook
        # book = load_workbook(summary_file_name)
        # writer = pd.ExcelWriter(summary_file_name, engine='openpyxl')
        # writer.book = book
        # writer.sheets = dict((ws.title, ws) for ws in book.worksheets)
        # self.well_summary_by_api_df.to_excel(writer, sheet_name=sheetname, cols=['Diff1', 'Diff2'])
        # writer.save()
        # writer = pd.ExcelWriter(summary_file_name, engine='xlsxwriter')
        # self.well_summary_by_api_df.to_excel(writer, sheet_name=sheetname, startrow=0, startcol=0, index=True, engine='openpyx1')
        # writer.save()

    def df_to_table_as_image(self, df, cfg):
        (ax, fig) = self.render_df_table(df,
                                         header_columns=0,
                                         col_width=2.0,
                                         font_size=8)
        fig.savefig(cfg['file_name'])
        self.plt.close()

    def df_to_table_as_docx(self, df, cfg):
        from docx import Document
        document = Document()
        document.add_heading("Heading")
        table = document.add_table(
            rows=(df.shape[0] + 1),
            cols=df.shape[1])    # First row are table headers!
        # Method 1
        # for i, column in enumerate(df):
        #     for row in range(df.shape[0]):
        #         table.cell(row, i).text = str(df[column][row])

        # Method 2
        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        document.save(cfg['file_name'] + '.docx')

    def render_df_table(self,
                        data,
                        col_width=3.0,
                        row_height=0.625,
                        font_size=14,
                        header_color='#40466e',
                        row_colors=['#f1f1f2', 'w'],
                        edge_color='w',
                        bbox=[0, 0, 1, 1],
                        header_columns=0,
                        ax=None,
                        **kwargs):

        import matplotlib.pyplot as plt
        import numpy as np
        import six
        self.plt = plt

        if ax is None:
            size = (np.array(data.shape[::-1]) + np.array([0, 1])) * np.array(
                [col_width, row_height])
            fig, ax = plt.subplots(figsize=size)
            ax.axis('off')

        mpl_table = ax.table(cellText=data.values,
                             bbox=bbox,
                             colLabels=data.columns,
                             **kwargs)

        mpl_table.auto_set_font_size(False)
        mpl_table.set_fontsize(font_size)

        for k, cell in six.iteritems(mpl_table._cells):
            cell.set_edgecolor(edge_color)
            if k[0] == 0 or k[1] < header_columns:
                cell.set_text_props(weight='bold', color='w')
                cell.set_facecolor(header_color)
            else:
                cell.set_facecolor(row_colors[k[0] % len(row_colors)])
        return ax, fig

    def write_ascii_file_from_text(self, text, file_name):
        f = open(file_name, 'wb')
        f.write(text.encode())
        f.close()

    def WriteOrcaFlexYMLFile(self, Files, output_filename):
        print('Write file: "{0}" .... '.format(output_filename))
        with open(output_filename, 'w') as f:
            for file in Files:
                if isinstance(file, str):
                    with open(file, "r") as infile:
                        f.write(infile.read())
                else:
                    yaml.dump(file, f)
        print('Write file: "{0}" .... COMPLETE'.format(output_filename))


class DefineData():

    def empty_data_frame(self, columns=None):
        import pandas as pd
        data_frame = pd.DataFrame(columns=columns)
        return data_frame

    def set_value_in_dictionary(self, data_dictionary, value, map_list=None):
        read_data = ReadData()
        # TODO Custom define maplist here
        read_data.extract_from_dictionary(data_dictionary,
                                          map_list[:-1])[map_list[-1]] = value

        return data_dictionary


class AttributeDict(dict):

    def __init__(self, *args, **kwargs):
        super(AttributeDict, self).__init__(*args, **kwargs)
        self.__dict__ = self


class objdict(dict):
    # TODO Test this attribute dictionary method
    def __getattr__(self, name):
        if name in self:
            return self[name]
        else:
            raise AttributeError("No such attribute: " + name)

    def __setattr__(self, name, value):
        self[name] = value

    def __delattr__(self, name):
        if name in self:
            del self[name]
        else:
            raise AttributeError("No such attribute: " + name)


class DateTimeUtility():

    def last_day_of_month(self, any_day):
        import datetime
        next_month = any_day.replace(day=28) + datetime.timedelta(
            days=4)    # this will never fail
        return next_month - datetime.timedelta(days=next_month.day)


class Transform():

    def numpy_interp(self,):
        # TODO Add for when the x and y are in ascending or descending order to handle correct interpolation
        pass

    def gis_deg_to_distance(self, df, cfg):
        cfg_sample = {
            'Longitude': 'BOT_LONG',
            'Latitude': 'BOT_LAT',
            'label': 'BOT'
        }
        import utm
        longitude_column = cfg['Longitude']
        latitude_column = cfg['Latitude']
        x_array = []
        y_array = []
        zone_array = []
        ut_array = []
        for df_row in range(0, len(df)):
            try:
                latitude = float(df[latitude_column].iloc[df_row])
                longitude = float(df[longitude_column].iloc[df_row])
                x, y, zone, ut = utm.from_latlon(latitude, longitude)
            except:
                x, y, zone, ut = (None, None, None, None)
            x_array.append(x)
            y_array.append(y)
            zone_array.append(zone)
            ut_array.append(ut)

        df[cfg['label'] + '_x'] = x_array
        df[cfg['label'] + '_y'] = y_array
        df[cfg['label'] + '_zone'] = zone_array
        df[cfg['label'] + '_ut'] = ut_array
        return df

    def gis_distance_to_deg(self):
        # TODO
        pass

    def get_gis_converted_df_superseded(self, data_set_cfg, df):
        if data_set_cfg.__contains__('gis'):
            import pyproj
            p = pyproj.Proj(proj='utm',
                            zone=data_set_cfg['gis']['zone'],
                            ellps='WGS84')
            if data_set_cfg['gis']['long_lat_to_northing_easting']['flag']:
                longitude_column = data_set_cfg['gis'][
                    'long_lat_to_northing_easting']['Longitude']
                latitude_column = data_set_cfg['gis'][
                    'long_lat_to_northing_easting']['Latitude']
                df['Easting'], df['Northing'] = p(
                    df[longitude_column].astype(float).tolist(),
                    df[latitude_column].astype(float).tolist())
            elif data_set_cfg['gis']['long_lat_to_x_y']['flag']:
                import utm
                longitude_column = data_set_cfg['gis']['long_lat_to_x_y'][
                    'Longitude']
                latitude_column = data_set_cfg['gis']['long_lat_to_x_y'][
                    'Latitude']
                df['x'], df['y'], zone, ut = utm.from_latlon(
                    df[longitude_column].astype(float).tolist(),
                    df[latitude_column].astype(float).tolist())
            elif data_set_cfg['gis']['northing_easting_to_long_lat']['flag']:
                Easting_column = data_set_cfg['gis'][
                    'long_lat_to_northing_easting']['Easting']
                Northing_column = data_set_cfg['gis'][
                    'long_lat_to_northing_easting']['Northing']
                df['Longitude'], df['Latitude'] = p(
                    df[Easting_column].tolist(),
                    df[Northing_column].tolist(),
                    inverse=True)

        return df

    def dataframe_to_dataframe(self, df, cfg=None):
        df_transposed = self.transpose_df(df, cfg)
        df_transposed = self.add_column_to_df(df_transposed, cfg)

        return df_transposed

    def dataframe_to_dict(self, df, cfg={}):
        json_dict = {}
        if df is not None:
            import json
            orient = cfg.get('orient', 'records')
            if len(df.columns.unique()) == len(df.columns):
                json_string = df.to_json(orient=orient)
                json_dict = json.loads(json_string)

        return json_dict

    def df_JSON_strings_values_to_dict(self, df, cfg_settings):
        if df is not None:
            if cfg_settings.__contains__('json_transformation'):
                import json
                for column in cfg_settings['json_transformation']['columns']:
                    json_array = []
                    for df_row in range(0, len(df)):
                        json_array.append(json.loads(df.iloc[df_row][column]))
                    df[column] = json_array

        return df.copy()

    def dataframe_to_json(self, df, cfg={}):
        import pandas as pd
        json_string = ""
        if df is not None:
            import json
            orient = cfg.get('orient', 'records')
            index = cfg.get('index', True)
            if index:
                df.insert(0, column='   ', value=list(df.index))

            if len(df.columns.unique()) == len(df.columns):
                json_string = df.apply(lambda x: pd.Series(x.dropna()),
                                       axis=1).to_json(orient=orient)
            else:
                df = self.df_transform_repeat_columns_to_unique_columns(df)
                json_string = df.to_json(orient=orient)
        return json_string

    def df_transform_repeat_columns_to_unique_columns(
            self, df, transform_character='trailing_alphabet'):
        old_columns = list(df.columns)
        cfg = {'list': old_columns, 'transform_character': transform_character}
        new_columns = self.transform_list_to_unique_list(cfg)
        df.columns = new_columns
        return df

    def transform_list_to_unique_list(self, cfg):
        cfg_temp = {
            'list': ['001', '001', '002'],
            'transform_character': 'trailing_alphabet',
            'result': ['001', '001a', '002']
        }
        old_list = cfg['list']
        old_list.reverse()
        new_list = old_list.copy()
        transform_character = cfg['transform_character']
        for list_index in range(0, len(old_list)):
            list_element = old_list[list_index]
            repeat_element_count = new_list.count(list_element)
            if repeat_element_count == 1:
                pass
            elif repeat_element_count > 1:
                if transform_character == 'trailing_space':
                    new_list[list_index] = new_list[list_index] + ' ' * (
                        repeat_element_count - 1)
                elif transform_character == 'leading_space':
                    new_list[list_index] = ' ' * (repeat_element_count -
                                                  1) + new_list[list_index]
                elif transform_character == 'leading_alphabet':
                    new_list[list_index] = 'a' * (repeat_element_count -
                                                  1) + new_list[list_index]
                elif transform_character == 'trailing_alphabet':
                    new_list[list_index] = new_list[list_index] + 'a' * (
                        repeat_element_count - 1)

        new_list.reverse()
        return new_list

    def add_column_to_df(self, df, cfg):
        if df is not None:
            add_column_to_transposed_df = cfg.get('add_column_to_transposed_df',
                                                  None)
            if add_column_to_transposed_df is not None:
                location = add_column_to_transposed_df['location']
                header = add_column_to_transposed_df['header']
                values = add_column_to_transposed_df['values']
                df.insert(loc=location, column=header, value=values)

        return df

    def transpose_df(self, df, cfg):
        df_transposed = df
        if df is not None:
            transposed_df_column_name = cfg.get('transposed_df_column_name',
                                                None)
            if transposed_df_column_name is not None:
                transpose_df_columns = df[
                    transposed_df_column_name['column']].tolist()
                if transposed_df_column_name['drop']:
                    df.drop(transposed_df_column_name['column'],
                            axis=1,
                            inplace=True)
                df_transposed = df.T.copy()
                df_transposed.columns = transpose_df_columns
        return df_transposed

    def dataframe_to_html(self, df=None, cfg_settings={}):
        if df is None:
            import pandas as pd
            df = pd.DataFrame({
                'name': ['Somu', 'Kiku', 'Amol', 'Lini'],
                'physics': [68, 74, 77, 78],
                'chemistry': [84, 56, 73, 69],
                'algebra': [78, 88, 82, 87]
            })
        index = cfg_settings.get('index', True)
        justify = cfg_settings.get('justify', 'center')
        classes = cfg_settings.get(
            'classes',
            'table table-bordered table-responsive-sm table-sm table-striped table-condensed'
        )
        max_cols = cfg_settings.get('max_cols', None)

        if (not cfg_settings.__contains__('json_transformation')) or (
                cfg_settings['json_transformation'] is None):
            html = df.to_html(index=index,
                              justify=justify,
                              classes=classes,
                              max_cols=max_cols)
            thead = cfg_settings.get('thead', '<thead class="thead-light">')
            html = html.replace('<thead>', thead)
            return html
        else:
            import json
            for column in cfg_settings['json_transformation']['columns']:
                json_array = []
                for df_row in range(0, len(df)):
                    import pandas as pd
                    try:
                        temp_df = pd.DataFrame.from_dict(
                            df.iloc[df_row][column], orient='columns')
                    except:
                        temp_df = pd.DataFrame.from_dict(
                            df.iloc[df_row][column], orient='index')
                    html = self.dataframe_to_html(temp_df)
                    json_array.append(html)
                df[column] = json_array

            df_dict = self.dataframe_to_dict(df)
            return df_dict

    def convert_numpy_types_to_native_python_types(self, cfg):
        if cfg['datatype'] == dict:
            for key in cfg['data'].keys():
                if type(cfg['data'][key]) not in [str, int, float]:
                    cfg['data'].update({key: cfg['data'][key].item()})
            return cfg['data']
        else:
            print("data types not supported")


class TransformData():

    def __init__(self):
        pass

    def get_transformed_data(self, cfg):
        calculate_function = getattr(self, cfg['type'])
        calculate_function(cfg)

    def linear(self, cfg):
        cfg_example1 = {
            'scale': 1,
            'shift': 3,
            'data': [1, 2, 3, 4],
            'type': 'linear'
        }
        cfg_example2 = {'scale': 1, 'shift': 3, 'data': 1, 'type': 'linear'}

        scale = cfg['scale']
        shift = cfg['shift']
        if type(cfg['data']) is list:
            cfg['data'] = [item * scale + shift for item in cfg['data']]
        else:
            cfg['data'] = cfg['data'] * scale + shift

        return cfg


class TransferDataFromExcelToWord():

    def __init__(self):
        pass

    def transfer_table_from_excel_to_word(self, cfg=None):
        if cfg is None:
            cfg = {
                'excel_file': 'K:\\0173 KM Extreme\\SLWR\\Fatigue\\Test.xlsx',
                'sheet_name': 'Sheet1',
                'table_name': 'Table1',
                'word_file': 'K:\\0173 KM Extreme\\SLWR\\Fatigue\\Test.docx',
                'word_table_name': 'Table1'
            }
        #TODO implement this function


if __name__ == '__main__':
    # write better tests
    FileName = 'K:\\0173 KM Extreme\\SLWR\\Fatigue\\Test.xlsx'
    Columns = [
        'Arc Length', 'S-N Curve', 'Theta', 'Overall Damage', 'Life (years)'
    ]
    CustomData = {
        "FileName": FileName,
        "SheetName": "Sheet1",
        "KeyWords": ['Arc Length'],
        "RowsToSkip": 2,
        "RowsToRead": 1000,
        "Columns": Columns
    }
    # df = xlsx_To_DataFrame(CustomData)
    # print(df)

    import pandas as pd
    save_data = SaveData()
    example_data = {
        'first_name': ['Jason', 'Molly', 'Tina', 'Jake', 'Amy'],
        'last_name': ['Miller', 'Jacobson', 'Ali', 'Milner', 'Cooze'],
        'age': [42, 52, 36, 24, 73],
        'preTestScore': [4, 24, 31, 2, 3],
        'postTestScore': [25, 94, 57, 62, 70]
    }
    df = pd.DataFrame(example_data,
                      columns=[
                          'first_name', 'last_name', 'age', 'preTestScore',
                          'postTestScore'
                      ])
    cfg = {'file_name': 'test_file'}
    save_data.df_to_table_as_docx(df, cfg)


class PandasChainedAssignent:

    def __init__(self, chained=None):
        acceptable = [None, 'warn', 'raise']
        assert chained in acceptable, "chained must be in " + str(acceptable)
        self.swcw = chained

    def __enter__(self):
        import pandas as pd
        self.saved_swcw = pd.options.mode.chained_assignment
        pd.options.mode.chained_assignment = self.swcw
        return self

    def __exit__(self, *args):
        import pandas as pd
        pd.options.mode.chained_assignment = self.saved_swcw


def transform_df_datetime_to_str(df, date_format='%Y-%m-%d %H:%M:%S'):
    df = df.copy()
    if len(df) > 0:
        df_columns = list(df.columns)
        for column in df_columns:
            if isinstance(df[column].iloc[0], datetime.datetime) or isinstance(
                    df[column].iloc[0], datetime.date):
                df[column] = [
                    item.strftime(date_format) for item in df[column].to_list()
                ]

    return df


class CopyAndPasteFiles:
    '''
    Class to copy and paste files from 1 directory to another.
    '''

    def __init__(self):
        pass

    def iterate_all_cfgs(self, cfgs):
        copy_cfgs = cfgs['copy_cfgs'].copy()
        for cfg in copy_cfgs:
            self.copy_files(cfg)

    def copy_files(self, cfg):
        '''
        cfg = {
            'source_dir': 'C:\\Users\\kylem\\Desktop\\Test\\',
            'destination_dirs': 'C:\\Users\\kylem\\Desktop\\Test2\\',
            'file_names': ['test1.txt', 'test2.txt']
        }
        '''
        import shutil
        file_names = cfg['files']
        for file_name in file_names:
            for destination_dir in cfg['destination_dirs']:
                shutil.copy(cfg['source_dir'] + file_name,
                            destination_dir + file_name)
