# ABOUTME: B1528 SIROCCO current-heading/rudder force-component sweep report.
# ABOUTME: Generates ship-fixed force/moment tables plus interactive dropdown charts.
"""B1528/SIROCCO current-heading/rudder force component report utilities.

This module extends the existing B1528 moored-current rudder-only calculation to
an approved visualization sweep over current speed, heading offset, and rudder
angle. It reports rudder-induced loads and generic/reference OCIMF hull-current
review components in a ship-fixed COG frame. It is not a validated oblique-current
hull/rudder/MMG model and deliberately excludes mooring-line stiffness, tug loads,
bank effects, and class compliance conclusions.
"""

from __future__ import annotations

import csv
import json
import math
import os
from dataclasses import asdict, dataclass
from functools import lru_cache
from importlib import resources
from pathlib import Path
from typing import Any

import yaml

from digitalmodel.citations.schema import Citation, validate_citation
from digitalmodel.naval_architecture.b1528_sirocco_yaw_report import (
    GITHUB_REPO_BLOB,
    KNOT_TO_M_PER_S,
    NON_ROTATING_PROPELLER_CR,
    PROPELLER_ROTATION_FACTOR_NOTE,
)

REPORT_REVIEW_TARGET_DATE = "2026-05-09"
DEFAULT_CHART_RUDDER_ANGLE_DEG = 28.0
REPORT_ARTIFACT_STEM = "b1528_sirocco_current_rudder_force"
REPORT_DISPLAY_TITLE = "B1528 SIROCCO Current/Rudder Force Review — Issue #2760"
REPORT_DISPLAY_HEADING = "B1528 SIROCCO current/rudder force review — Issue #2760"
REPORT_SCOPE = (
    "SIROCCO issue #2760 current/rudder force-component review at COG; "
    "rudder-induced ship-fixed X/Y/N components are derived by rotating local "
    "current-frame loads by heading offset, and generic/reference OCIMF tanker-current "
    "review loads are estimated separately from the licensed off-repo workbook route; "
    "bank effect, tug loads, mooring-line stiffness, current-profile variation, "
    "propeller race, and class compliance conclusion are excluded"
)
ZERO_EFFECTIVE_ANGLE_NOTE = (
    "When rudder_angle_deg equals heading_offset_deg, alpha is zero and this "
    "rudder-induced component is zero; that is not total hull current load."
)
OCIMF_WORKBOOK_PATH_ENV = "OCIMF_WORKBOOK_PATH"
OCIMF_WORKBOOK_SOURCE_ID = "licensed-off-repo-ocimf-workbook"
OCIMF_PROVENANCE_README = "docs/data/OCIMF_CORPUS_README.md"


def _repo_root() -> Path:
    return Path(__file__).resolve().parents[3]


def _resolve_ocimf_workbook_path() -> Path:
    raw_path = os.environ.get(OCIMF_WORKBOOK_PATH_ENV)
    if not raw_path:
        raise FileNotFoundError(
            f"Set {OCIMF_WORKBOOK_PATH_ENV} to the licensed OCIMF coefficient workbook before "
            "running the #2760 current/rudder calculation. The workbook is intentionally "
            "kept off-repo and is not emitted in report artifacts."
        )
    return Path(raw_path).expanduser()


def _resolve_ocimf_provenance_readme() -> Path:
    return _repo_root() / OCIMF_PROVENANCE_README


@dataclass(frozen=True)
class B1528CurrentHeadingRudderConfig:
    """Validated B1528 current-heading/rudder report configuration."""

    case_id: str
    aliases: tuple[str, ...]
    source_pack_issue: str
    report_issue: str
    plan_path: str
    lbp_m: float
    beam_m: float
    draft_m: float
    water_depth_m: float
    yaw_lever_m: float
    rudder_area_m2: float
    rudder_span_m: float
    ship_sog_kn: float
    current_speeds_kn: tuple[float, ...]
    chart_default_current_speed_kn: float
    default_speed_policy: str
    heading_offsets_deg: tuple[float, ...]
    rudder_angles_deg: tuple[float, ...]
    rho_kg_m3: float
    beta: float
    prop_rotation_factor: float
    prop_rotation_factor_logic: str
    force_convention: dict[str, str]
    limitations: tuple[str, ...]
    raw: dict[str, Any]


def load_packaged_b1528_current_heading_rudder_config() -> B1528CurrentHeadingRudderConfig:
    """Load packaged B1528 SIROCCO current-heading/rudder YAML."""

    resource = resources.files("digitalmodel.naval_architecture.data").joinpath(
        "b1528_sirocco_current_heading_rudder.yml"
    )
    return validate_b1528_current_heading_rudder_config(
        yaml.safe_load(resource.read_text(encoding="utf-8"))
    )


def issue_2760_source_preflight() -> dict[str, Any]:
    """Fail-closed source/citation preflight for approved issue #2760 implementation.

    Env-var and file-existence checks run on every call so monkeypatched tests
    still observe fail-closed behavior. The expensive Citation construction +
    `validate_citation` work is memoized in `_issue_2760_source_preflight_cached`
    keyed on (path, mtime) tuples; cache invalidates when the workbook or
    provenance README is replaced.
    """

    workbook_path = _resolve_ocimf_workbook_path()
    provenance_readme = _resolve_ocimf_provenance_readme()
    if not workbook_path.exists():
        raise FileNotFoundError(
            f"Required licensed OCIMF workbook from {OCIMF_WORKBOOK_PATH_ENV} is missing"
        )
    if not provenance_readme.exists():
        raise FileNotFoundError(f"Required OCIMF provenance README is missing: {OCIMF_PROVENANCE_README}")
    return _issue_2760_source_preflight_cached(
        str(workbook_path),
        workbook_path.stat().st_mtime_ns,
        str(provenance_readme),
        provenance_readme.stat().st_mtime_ns,
    )


@lru_cache(maxsize=4)
def _issue_2760_source_preflight_cached(
    workbook_path_text: str,
    workbook_mtime_ns: int,
    readme_path_text: str,
    readme_mtime_ns: int,
) -> dict[str, Any]:
    """Cached citation construction/validation; cache key invalidates on mtime change."""

    del workbook_path_text, workbook_mtime_ns, readme_path_text, readme_mtime_ns
    ocimf_citation = Citation(
        code_id="OCIMF-MEG4",
        publisher="Oil Companies International Marine Forum",
        revision="4th edition, 2018",
        section="Annex A current coefficient figures A9/A10/A11",
        wiki_path="wikis/marine-engineering/wiki/standards/ocimf-meg4.md",
        note="Generic/reference tanker-current coefficient basis for #2760; not SIROCCO-specific.",
    )
    rudder_citation = Citation(
        code_id="B1528-SIROCCO-RUDDER-YAW-INPUTS",
        publisher="ACMA project source pack",
        revision="2026-05-01",
        section="Rudder force and yaw-moment workbook inputs",
        wiki_path="wikis/acma-projects/wiki/concepts/b1528-sirocco-rudder-yaw-moment-inputs.md",
        note="Project-specific B1528 rudder geometry and normal-force constant used for screening calculation.",
    )
    validate_citation(ocimf_citation)
    validate_citation(rudder_citation)

    return {
        "ocimf": {
            "workbook_source_id": OCIMF_WORKBOOK_SOURCE_ID,
            "provenance_readme": OCIMF_PROVENANCE_README,
            "citation": ocimf_citation,
            "license_boundary": "pointer-only-no-coefficient-corpus",
            "basis": "generic/reference tanker-current coefficient route for screening review",
        },
        "rudder": {
            "citation": rudder_citation,
            "basis": "normal-force rudder model, rpm=0, Cr=1.0",
        },
    }


def validate_b1528_current_heading_rudder_config(
    payload: dict[str, Any],
) -> B1528CurrentHeadingRudderConfig:
    """Validate the B1528 current-heading/rudder YAML contract."""

    required = {
        "case",
        "source_pack",
        "current_heading_rudder_report",
        "vessel",
        "rudder",
        "scenario",
        "environment",
        "workbook_regression",
        "force_convention",
        "limitations",
    }
    missing = sorted(required - set(payload))
    if missing:
        raise ValueError(f"missing required top-level sections: {missing}")

    scenario = payload["scenario"]
    workbook = payload["workbook_regression"]
    cfg = B1528CurrentHeadingRudderConfig(
        case_id=str(payload["case"]["id"]),
        aliases=tuple(str(value) for value in payload["case"].get("aliases", [])),
        source_pack_issue=str(payload["source_pack"]["issue"]),
        report_issue=str(payload["current_heading_rudder_report"]["issue"]),
        plan_path=str(payload["current_heading_rudder_report"]["plan"]),
        lbp_m=_positive("lbp_m", payload["vessel"]["lbp_m"]),
        beam_m=_positive("beam_m", payload["vessel"].get("beam_m", 32.26)),
        draft_m=_positive("draft_m", payload["vessel"].get("draft_m", 12.2)),
        water_depth_m=_positive("water_depth_m", payload["environment"].get("water_depth_m", 100.0)),
        yaw_lever_m=_positive("yaw_lever_m", payload["vessel"]["yaw_lever_m"]),
        rudder_area_m2=_positive("rudder_area_m2", payload["rudder"]["area_m2"]),
        rudder_span_m=_positive("rudder_span_m", payload["rudder"]["span_m"]),
        ship_sog_kn=_nonnegative("ship_sog_kn", scenario["ship_sog_kn"]),
        current_speeds_kn=tuple(float(value) for value in scenario["current_speeds_kn"]),
        chart_default_current_speed_kn=float(scenario["chart_default_current_speed_kn"]),
        default_speed_policy=str(scenario["default_speed_policy"]),
        heading_offsets_deg=tuple(float(value) for value in scenario["heading_offsets_deg"]),
        rudder_angles_deg=tuple(float(value) for value in scenario["rudder_angles_deg"]),
        rho_kg_m3=_positive("rho_kg_m3", payload["environment"]["rho_kg_m3"]),
        beta=_positive("beta", workbook["beta"]),
        prop_rotation_factor=_positive("prop_rotation_factor", workbook["prop_rotation_factor"]),
        prop_rotation_factor_logic=str(workbook["prop_rotation_factor_logic"]),
        force_convention={key: str(value) for key, value in payload["force_convention"].items()},
        limitations=tuple(str(value) for value in payload["limitations"]),
        raw=payload,
    )
    if cfg.current_speeds_kn != (0.0, 1.0, 2.0, 3.0, 3.08, 4.0, 5.0):
        raise ValueError("current_speeds_kn must match the approved issue #2760 0..5 kn sweep including 3.08 kn (Pass H sensitivity extension)")
    expected_heading_angles = tuple(float(value) for value in range(-5, 6))
    if cfg.heading_offsets_deg != expected_heading_angles:
        raise ValueError("heading_offsets_deg must be -5..+5 deg in 1 deg steps")
    expected_rudder_angles = tuple(float(value) for value in range(-28, 29, 2))
    if cfg.rudder_angles_deg != expected_rudder_angles:
        raise ValueError("rudder_angles_deg must be -28..+28 deg in 2 deg steps (Pass H+ symmetric sweep)")
    if cfg.chart_default_current_speed_kn != 3.08:
        raise ValueError("chart default current speed must be exact 3.08 kn")
    if cfg.prop_rotation_factor != NON_ROTATING_PROPELLER_CR:
        raise ValueError("current-heading/rudder report requires neutral Cr=1.0")
    return cfg


def speed_planes_for_calculation(cfg: B1528CurrentHeadingRudderConfig) -> tuple[float, ...]:
    """Return engineering speeds plus exact chart-default speed if absent."""

    speeds = list(cfg.current_speeds_kn)
    if cfg.chart_default_current_speed_kn not in speeds:
        speeds.append(cfg.chart_default_current_speed_kn)
    return tuple(sorted(speeds))


def select_ocimf_loaded_tanker_current_basis(cfg: B1528CurrentHeadingRudderConfig) -> dict[str, Any]:
    """Document the geometry-driven OCIMF tanker-current basis for #2760.

    MEG Annex A publishes loaded-tanker current curves by water-depth/draft
    bucket. For the approved #2760 screening use, B1528 is off-class and the
    deepest available loaded-tanker bucket is used when the configured WD/T
    exceeds that domain.
    """

    wd_over_t = cfg.water_depth_m / cfg.draft_m
    if wd_over_t <= 0.0:
        raise ValueError("water_depth_m / draft_m must be positive for OCIMF basis selection")
    if wd_over_t > 6.0:
        selected_bucket = ">6"
        selected_figures = ["A9", "A10", "A11"]
        rejected = ["A5-A8 shallower loaded-tanker Cxc buckets", "A12-A14 ballast 40%T tanker curves"]
    else:
        raise ValueError(
            "#2760 requires a documented loaded-tanker current curve basis; "
            f"configured WD/T={wd_over_t:.3f} is outside the approved >6 screening bucket"
        )
    return {
        "selection_rule": "max_available_wd_over_t_for_loaded_tanker_current",
        "wd_over_t": wd_over_t,
        "water_depth_m": cfg.water_depth_m,
        "draft_m": cfg.draft_m,
        "selected_wd_over_t_bucket": selected_bucket,
        "selected_figures": selected_figures,
        "selected_curve_family": "loaded_tanker_current",
        "rejected_alternatives": rejected,
        "limitation": (
            "B1528 SIROCCO is not an OCIMF tanker-class vessel; loaded-tanker "
            "curves are used off-class as a generic/reference screening basis only."
        ),
    }


def run_b1528_current_heading_rudder_report(
    config: B1528CurrentHeadingRudderConfig | None = None,
) -> dict[str, Any]:
    """Run the current speed × heading offset × rudder angle sweep."""

    cfg = config or load_packaged_b1528_current_heading_rudder_config()
    rows = [
        _row(cfg, speed, heading, rudder)
        for speed in speed_planes_for_calculation(cfg)
        for heading in cfg.heading_offsets_deg
        for rudder in cfg.rudder_angles_deg
    ]
    result = {"metadata": _metadata(cfg), "rows": rows}
    result["summary"] = _summary(rows)
    result["sample_working_example"] = _sample_working_example(result)
    return result


def write_b1528_current_heading_rudder_report(
    result: dict[str, Any], output_dir: str | Path
) -> dict[str, str]:
    """Write CSV/JSON/provenance/manifest plus Markdown, interactive HTML, and PDF."""

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    csv_path = out / f"{REPORT_ARTIFACT_STEM}_results.csv"
    json_path = out / f"{REPORT_ARTIFACT_STEM}_results.json"
    provenance_path = out / f"{REPORT_ARTIFACT_STEM}_provenance.json"
    citation_path = out / f"{REPORT_ARTIFACT_STEM}_citations.json"
    md_path = out / f"{REPORT_ARTIFACT_STEM}_report.md"
    docx_path = out / f"{REPORT_ARTIFACT_STEM}_report.docx"
    html_path = out / f"{REPORT_ARTIFACT_STEM}_report.html"
    pdf_path = out / f"{REPORT_ARTIFACT_STEM}_report.pdf"
    manifest_path = out / f"{REPORT_ARTIFACT_STEM}_manifest.json"

    headers = list(result["rows"][0])
    with csv_path.open("w", newline="", encoding="utf-8") as stream:
        writer = csv.DictWriter(stream, fieldnames=headers, lineterminator="\n")
        writer.writeheader()
        writer.writerows(result["rows"])
    json_path.write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")
    provenance_path.write_text(json.dumps(_provenance(result), indent=2) + "\n", encoding="utf-8")
    citation_path.write_text(json.dumps(_citation_sidecar(result), indent=2) + "\n", encoding="utf-8")
    md_path.write_text(_markdown_report(result) + "\n", encoding="utf-8")
    # Pass H follow-up "add pictures": write HTML first so we can screenshot SVG
    # schematics via Playwright, then embed PNGs in DOCX. PDF is rendered last and
    # also benefits from the screenshot pass (Chrome rasterises the SVGs at print time).
    html_path.write_text(_html_report(result), encoding="utf-8")
    schematic_pngs = _capture_schematic_pngs(html_path, out / "schematic-pngs")
    _write_docx_report(result, docx_path, schematic_pngs=schematic_pngs)
    _write_pdf_from_html(html_path, pdf_path)
    manifest = {
        "csv": str(csv_path),
        "json": str(json_path),
        "provenance": str(provenance_path),
        "citation_sidecar": str(citation_path),
        "markdown_report": str(md_path),
        "docx_report": str(docx_path),
        "html_report": str(html_path),
        "pdf_report": str(pdf_path),
        "manifest": str(manifest_path),
    }
    manifest_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    return manifest


def _write_docx_report(
    result: dict[str, Any],
    docx_path: Path,
    schematic_pngs: dict[str, Path] | None = None,
) -> None:
    """Write a Word .docx report matching the MD/HTML canonical 6-section layout.

    Sections: (1) Introduction, (2) Design Data & Assumptions, (3) Axes & Sign
    Conventions, (4) Load Due to Current, (5) Load Due to Rudder, (6) Limitations.
    Reading level: freshman engineering grad — short sentences, defined symbols,
    professional citation references (OCIMF MEG3/MEG4 Annex A specifics, not
    vague "workbook" prose).
    """

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:  # pragma: no cover - dependency availability is environment-specific
        raise RuntimeError("python-docx is required to render the B1528 SIROCCO Word report") from exc

    schematic_pngs = schematic_pngs or {}

    def _add_schematic_picture(schematic_id: str, caption: str) -> None:
        """Embed a schematic PNG with caption if Playwright capture succeeded."""
        png_path = schematic_pngs.get(schematic_id)
        if png_path and Path(png_path).exists():
            document.add_picture(str(png_path), width=Inches(5.5))
            cap_p = document.add_paragraph()
            cap_run = cap_p.add_run(caption)
            cap_run.italic = True

    metadata = result["metadata"]
    sample = result["sample_working_example"]
    design = metadata["design_data"]
    links = metadata["traceability_links"]

    default_speed = metadata["chart_default_current_speed_kn"]
    default_heading = 5.0
    default_rudder = DEFAULT_CHART_RUDDER_ANGLE_DEG
    default_alpha = default_rudder - default_heading
    table_angle = 180.0 - abs(default_heading)
    v_m_s = default_speed * KNOT_TO_M_PER_S
    q_pa = 0.5 * design["rho_kg_m3"] * v_m_s ** 2
    a_f = design["beam_m"] * design["draft_m"]
    a_l = design["lbp_m"] * design["draft_m"]
    wd_over_t = design["water_depth_to_draft_ratio"]

    document = Document()
    document.add_heading(REPORT_DISPLAY_TITLE, level=0)
    document.add_paragraph(f"Prepared for engineer review on {metadata['review_target_date']}.")

    # §1 Introduction
    document.add_heading("1. Introduction", level=1)
    document.add_paragraph(
        "The B1528 SIROCCO is a moored vessel. When current flows past it, two separate "
        "loads act on the ship:"
    )
    document.add_paragraph(
        "A hull current load — the current pushes on the hull, producing a longitudinal "
        "force (X), a transverse force (Y), and a yaw moment (N) about the centre of gravity.",
        style="List Number",
    )
    document.add_paragraph(
        "A rudder-induced load — current also flows past the rudder. When the rudder is "
        "held at an angle, it generates an additional force at the stern that contributes "
        "to X, Y, and N.",
        style="List Number",
    )
    document.add_paragraph("This report estimates both load components for the approved review case:")
    for label, value in (
        ("Current speed", f"{default_speed:.2f} knots"),
        ("Current heading ψ", f"+{default_heading:.0f}° off the bow (port positive)"),
        ("Rudder angle δ", f"+{default_rudder:.0f}° (port)"),
        ("Propeller", f"stopped (rpm = 0; rotation factor Cr = {design['prop_rotation_factor']:.1f})"),
    ):
        document.add_paragraph(f"{label}: {value}", style="List Bullet")
    document.add_paragraph(
        "All forces are reported in ship-fixed axes about the centre of gravity. This is a "
        "screening calculation intended to support mooring reaction-load reviews — it is "
        "not a certified hydrodynamic model. Limitations are listed in §6."
    )

    # §2 Design Data & Assumptions
    document.add_heading("2. Design Data & Assumptions", level=1)
    document.add_paragraph(
        "Values used by the calculation are sourced from the packaged B1528 SIROCCO YAML "
        "input pack (see Bibliography §B.4)."
    )
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Value"
    for label, value in (
        ("LBP (length between perpendiculars)", f"{design['lbp_m']:.1f} m"),
        ("Beam", f"{design['beam_m']:.2f} m"),
        ("Draft", f"{design['draft_m']:.2f} m"),
        ("Water depth", f"{design['water_depth_m']:.1f} m"),
        ("Water depth / Draft (WD/T)", f"{wd_over_t:.2f}"),
        ("Centre of gravity (longitudinal datum)", "midship"),
        ("Centre of gravity (vertical, above keel)", "6.1 m"),
        ("Yaw lever arm (rudder force to CoG)", f"{design['yaw_lever_m']:.2f} m"),
        ("Rudder area A_R", f"{design['rudder_area_m2']:.2f} m²"),
        ("Rudder span", f"{design['rudder_span_m']:.2f} m"),
        ("Water density ρ", f"{design['rho_kg_m3']:.1f} kg/m³"),
        ("Whicker-Fehlner constant β", f"{design['beta']:.1f}"),
        ("Propeller rotation factor Cr", f"{design['prop_rotation_factor']:.1f} (rpm = 0)"),
        ("Current speed (default)", f"{default_speed:.2f} kn"),
        ("Current heading ψ (default)", f"+{default_heading:.0f}° (port positive)"),
        ("Rudder angle δ (default)", f"+{default_rudder:.0f}° (port positive)"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    # §3 Axes & Sign Conventions
    document.add_heading("3. Axes & Sign Conventions", level=1)
    document.add_paragraph(
        "All forces and moments are reported in ship-fixed axes at the centre of gravity "
        "using the right-hand rule."
    )
    convention_table = document.add_table(rows=1, cols=3)
    hdr = convention_table.rows[0].cells
    hdr[0].text = "Symbol"
    hdr[1].text = "Meaning"
    hdr[2].text = "Sign convention"
    for symbol, meaning, sign in (
        ("+X", "longitudinal force (along ship centreline)", "forward"),
        ("+Y", "transverse force (lateral, across ship)", "port"),
        ("+N", "yaw moment about CoG", "bow-to-port"),
        ("ψ", "current heading offset from bow", "port positive"),
        ("δ", "rudder angle relative to ship centreline", "port positive"),
        ("α", "effective rudder inflow angle = δ - ψ", "(derived)"),
    ):
        cells = convention_table.add_row().cells
        cells[0].text = symbol
        cells[1].text = meaning
        cells[2].text = sign
    document.add_paragraph(
        "See the per-section schematics below (§3 axes-conventions, §4 current-loading "
        "+ current-moment, §5 rudder-loading) for the ship outline with annotated arrows "
        "showing each axis and angle."
    )
    _add_schematic_picture(
        "schematic-axes-conventions",
        "Figure 0 — Axes & Sign Conventions (per OCIMF MEG4 [1] Annex A §A.1).",
    )

    # §4 Load Due to Current
    document.add_heading("4. Load Due to Current", level=1)
    document.add_paragraph(
        "The current load on the hull is estimated using generic-reference loaded-tanker "
        "current coefficients interpolated from OCIMF Mooring Equipment Guidelines, 4th "
        "edition (OCIMF MEG4, 2018), Annex A figures A9 (Cxc), A10 (Cyc), and A11 (Cxyc). "
        "These curves represent a tanker-class basis — they are not vessel-specific to "
        "SIROCCO and are used here as a screening estimate only. See Bibliography §B.1."
    )
    document.add_paragraph(
        f"Workbook bucket selection: loaded-tanker family, water-depth-to-draft ratio "
        f"WD/T > 6 (vessel WD/T = {wd_over_t:.2f}). Coefficients are interpolated at the "
        f"table angle 180° − |ψ| = {table_angle:.0f}° for the default heading "
        f"ψ = +{default_heading:.0f}°."
    )

    document.add_heading("4.1. Force calculation", level=2)
    document.add_paragraph(
        "Dynamic pressure: q = 0.5 × ρ × V². Frontal projected area: A_f = beam × draft. "
        "Lateral projected area: A_l = LBP × draft. Force components: longitudinal "
        "Xc = q × A_f × Cxc; transverse Yc = q × A_l × Cyc."
    )
    document.add_paragraph("Sample calculation at default values:")
    xc_kN = (q_pa * a_f * -0.0324) / 1000.0
    yc_kN = (q_pa * a_l * 0.0341) / 1000.0
    for line in (
        f"V = {default_speed:.2f} kn × {KNOT_TO_M_PER_S:.5f} = {v_m_s:.4f} m/s",
        f"q = 0.5 × {design['rho_kg_m3']:.0f} × ({v_m_s:.4f})² = {q_pa:.1f} Pa",
        f"A_f = {design['beam_m']:.2f} × {design['draft_m']:.2f} = {a_f:.2f} m²",
        f"A_l = {design['lbp_m']:.1f} × {design['draft_m']:.2f} = {a_l:.2f} m²",
        f"Cxc({table_angle:.0f}°) ≈ −0.0324, Cyc({table_angle:.0f}°) ≈ +0.0341 (OCIMF MEG4 Annex A Fig. A9/A10, loaded tanker WD/T>6)",
        f"Longitudinal Xc ≈ {xc_kN:.0f} kN (sign per Cxc<0 convention)",
        f"Transverse Yc ≈ {yc_kN:.0f} kN (+port per Cyc>0)",
    ):
        document.add_paragraph(line, style="List Bullet")
    _add_schematic_picture(
        "schematic-current-loading",
        "Figure 1 — Current loading at default case (OCIMF MEG4 [1] Annex A sign convention).",
    )

    document.add_heading("4.2. Yaw moment about CoG", level=2)
    document.add_paragraph(
        f"Transverse current force Yc input — the HTML report's §4.2 leads with a Plotly "
        f"chart showing Yc across the heading sweep ψ ∈ [-{abs(default_heading):.0f}°, "
        f"+{default_heading:.0f}°] at 1° steps for the selected current speed (default "
        f"{default_speed:.2f} kn). Yc is the input that drives Method B's "
        "force × lever-arm yaw moment; data is also exposed in the generated CSV "
        "(ocimf_current_force_y_ship_port_N column)."
    )
    document.add_paragraph(
        "Two methods estimate the current yaw moment, shown side by side as a sanity check:"
    )
    document.add_paragraph(
        f"Method A — OCIMF direct (default): Nc_A = q × A_l × LBP × Cxyc({table_angle:.0f}°), "
        f"per OCIMF MEG4 Annex A Fig. A11.",
        style="List Bullet",
    )
    document.add_paragraph(
        f"Method B — force × lever arm: Nc_B = Yc × CoP_lever_arm, where CoP_lever_arm = "
        f"{OCIMF_CURRENT_COP_LEVER_ARM_FRACTION:g}·LBP is an approximate center-of-pressure "
        "lever arm for loaded-tanker near-head-current per OCIMF MEG3/MEG4 Annex A figures.",
        style="List Bullet",
    )
    document.add_paragraph(
        "These two methods rest on different assumptions and may differ. The HTML report's "
        "§4.2.1 chart shows both as an overlay with an explicit non-equality caption — this "
        "is a sanity-check overlay, NOT an equality test."
    )
    _add_schematic_picture(
        "schematic-current-moment",
        "Figure 2 — Current yaw moment about CoG: CoP, lever arm, and +M_XY direction.",
    )

    # §5 Load Due to Rudder
    document.add_heading("5. Load Due to Rudder", level=1)
    document.add_paragraph(
        f"The rudder-induced force comes from current flowing past the rudder blade at the "
        f"effective inflow angle α = δ - ψ. At the default case "
        f"(δ = +{default_rudder:.0f}° port, ψ = +{default_heading:.0f}° port): "
        f"α = +{default_alpha:.0f}°."
    )
    _add_schematic_picture(
        "schematic-rudder-loading",
        "Figure 3 — Rudder loading at default deflection δ=+28° port; normal force F decomposed into +F_X and +F_Y at the rudder pivot.",
    )
    document.add_paragraph(
        "Two screening rudder models are presented side by side:"
    )
    document.add_paragraph(
        f"Model A — Whicker-Fehlner normal-force basis (default): F = β × A_R × V² × Cr × sin(α). "
        f"Source: B1528 SIROCCO source pack (β = {design['beta']:.0f}, Cr = "
        f"{design['prop_rotation_factor']:.1f}). See Bibliography §B.4.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Model B — thin-plate drag/lift (comparison): Cn(α) ≈ 2π·sin(α) for small angles, "
        "with stall handling at ~25–30°. Reference: Faltinsen, Sea Loads on Ships and "
        "Offshore Structures, Cambridge University Press, 1990, §6.5. See Bibliography §B.2.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Both models are screening-level; neither is a validated rudder hydrodynamic model. "
        "Differences at large angles (>20°) reflect stall-region behaviour and the "
        "simplifying assumptions of each."
    )
    document.add_heading("5.1. Sample calculation at default values (Whicker-Fehlner - Model A)", level=2)
    for line in (
        f"α = δ - ψ = {default_rudder:.0f}° - {default_heading:.0f}° = {default_alpha:.0f}°",
        f"F = β × A_R × V² × Cr = {design['beta']:.0f} × {design['rudder_area_m2']:.2f} × "
        f"({v_m_s:.4f})² × {design['prop_rotation_factor']:.1f} = {sample['base_force_N']:.0f} N",
        f"Fn (normal force) = F × sin(α) = {sample['normal_force_N']:.0f} N",
        f"Longitudinal X_rudder ≈ {sample['force_x_ship_N']/1000.0:.0f} kN",
        f"Transverse Y_rudder ≈ {sample['force_y_ship_port_N']/1000.0:.0f} kN (+port)",
        f"Yaw moment N_rudder ≈ {sample['moment_n_yaw_bow_port_kN_m']:.0f} kN·m (+bow-to-port)",
    ):
        document.add_paragraph(line, style="List Bullet")

    # Pass H DOCX parity: locate default-case row + envelope rows + simple-plate row
    # for §5.2-§5.6 + Appendix A tables (mirrors the HTML pre-render so DOCX/PDF
    # readers see actual values, not "—" placeholders).
    default_row = next(
        (r for r in result["rows"]
         if abs(r["current_speed_kn"] - default_speed) < 1e-6
         and abs(r["heading_offset_deg"] - default_heading) < 1e-6
         and abs(r["rudder_angle_deg"] - default_rudder) < 1e-6),
        result["rows"][0],
    )
    default_speed_rows = [r for r in result["rows"] if abs(r["current_speed_kn"] - default_speed) < 1e-6]
    sensitivity_rows = [r for r in result["rows"]
                        if abs(r["heading_offset_deg"] - default_heading) < 1e-6
                        and abs(r["rudder_angle_deg"] - default_rudder) < 1e-6]

    def _kn(value_N: float) -> str:
        kn = value_N / 1000.0
        return f"{kn:.0f}" if abs(kn) >= 1 else f"{kn:.2f}"

    def _kNm(value_kNm: float) -> str:
        return f"{value_kNm:.0f}" if abs(value_kNm) >= 1 else f"{value_kNm:.2f}"

    # §5.2 Rudder force with rudder angle (NEW per user request — δ ∈ [-28°, +28°])
    document.add_heading("5.2. Rudder force with rudder angle", level=2)
    document.add_paragraph(
        f"Key variables: current speed V = {default_speed:.2f} kn · heading ψ = +{default_heading:.0f}° · "
        f"rudder sweep δ ∈ [-{default_rudder:.0f}°, +{default_rudder:.0f}°] at 2° steps. "
        "The HTML report's §5.2 Plotly chart shows rudder-induced longitudinal X_rudder and "
        f"transverse Y_rudder force vs rudder angle, with the default case (δ = +{default_rudder:.0f}°) "
        "marked. Both X_rudder and Y_rudder are signed per the Whicker-Fehlner sin(α) projection: "
        "at δ = +ψ (here +5°), α = 0 and rudder force is zero; sign reverses for δ < ψ. CSV exposes "
        "force_x_ship_N and force_y_ship_port_N columns for every rudder angle."
    )

    # §5.3 Rudder force over heading
    document.add_heading("5.3. Rudder force over heading", level=2)
    document.add_paragraph(
        f"Rudder-induced longitudinal X and transverse Y force across the heading sweep "
        f"ψ ∈ [-{abs(default_heading):.0f}°, +{default_heading:.0f}°] at 1° steps, at the "
        f"default rudder δ = +{default_rudder:.0f}° port and selected current speed "
        f"V = {default_speed:.2f} kn. Interactive chart in HTML §5.2; CSV exposes "
        "force_x_ship_N and force_y_ship_port_N columns for every row in the sweep."
    )

    # §5.3 Rudder yaw moment over heading
    document.add_heading("5.4. Rudder yaw moment over heading", level=2)
    document.add_paragraph(
        f"Signed rudder-induced yaw moment N (kN·m, +bow-to-port) across the heading sweep "
        f"at the selected current speed. Interactive chart in HTML §5.3; CSV exposes "
        "moment_n_yaw_bow_port_kN_m column."
    )

    # §5.4 Selected-speed envelope summary
    document.add_heading("5.5. Selected-speed envelope summary", level=2)
    document.add_paragraph(
        f"Key variables for this section: current speed V = {default_speed:.2f} kn (default) · "
        f"heading sweep ψ ∈ [-{abs(default_heading):.0f}°, +{default_heading:.0f}°] at 1° steps · "
        f"rudder sweep δ ∈ [0°, +{default_rudder:.0f}°] at 2° steps. Envelope is the "
        "absolute-max across the full heading × rudder grid at the selected speed."
    )
    if default_speed_rows:
        max_y = max(default_speed_rows, key=lambda r: abs(r["force_y_ship_port_N"]))
        max_n = max(default_speed_rows, key=lambda r: abs(r["moment_n_yaw_bow_port_kN_m"]))
        max_x = max(default_speed_rows, key=lambda r: abs(r["force_x_ship_N"]))
        env_table = document.add_table(rows=1, cols=3)
        hdr = env_table.rows[0].cells
        hdr[0].text = "Metric"; hdr[1].text = "Envelope case"; hdr[2].text = "Value"
        for label, row, val_str in (
            (f"Max |Y| at V = {default_speed:.2f} kn",
             f"ψ={max_y['heading_offset_deg']:.0f}° · δ={max_y['rudder_angle_deg']:.0f}°",
             f"{_kn(max_y['force_y_ship_port_N'])} kN"),
            (f"Max |N| at V = {default_speed:.2f} kn",
             f"ψ={max_n['heading_offset_deg']:.0f}° · δ={max_n['rudder_angle_deg']:.0f}°",
             f"{_kNm(max_n['moment_n_yaw_bow_port_kN_m'])} kN·m"),
            (f"Max |X| at V = {default_speed:.2f} kn",
             f"ψ={max_x['heading_offset_deg']:.0f}° · δ={max_x['rudder_angle_deg']:.0f}°",
             f"{_kn(max_x['force_x_ship_N'])} kN"),
        ):
            cells = env_table.add_row().cells
            cells[0].text = label; cells[1].text = row; cells[2].text = val_str

    # §5.5 Selected-case force breakdown
    document.add_heading("5.6. Selected-case force breakdown", level=2)
    document.add_paragraph(
        f"Key variables: V = {default_speed:.2f} kn · ψ = +{default_heading:.0f}° · "
        f"δ = +{default_rudder:.0f}° · α = +{default_alpha:.0f}°. Components below are "
        "at the default case (HTML/JS updates these when the user changes the speed or "
        "rudder selector)."
    )
    bd_table = document.add_table(rows=1, cols=4)
    hdr = bd_table.rows[0].cells
    hdr[0].text = "Component"; hdr[1].text = "X (kN, long.)"
    hdr[2].text = "Y (kN, port)"; hdr[3].text = "N (kN·m, +bow-port)"
    for label, x_val, y_val, n_val in (
        ("Current (OCIMF MEG4 [1])",
         _kn(default_row["ocimf_current_force_x_ship_N"]),
         _kn(default_row["ocimf_current_force_y_ship_port_N"]),
         _kNm(default_row["ocimf_current_moment_n_yaw_bow_port_kN_m"])),
        ("Rudder (Whicker-Fehlner [3])",
         _kn(default_row["force_x_ship_N"]),
         _kn(default_row["force_y_ship_port_N"]),
         _kNm(default_row["moment_n_yaw_bow_port_kN_m"])),
        ("Total (current + rudder)",
         _kn(default_row["total_force_x_ship_N"]),
         _kn(default_row["total_force_y_ship_port_N"]),
         _kNm(default_row["total_moment_n_yaw_bow_port_kN_m"])),
    ):
        cells = bd_table.add_row().cells
        cells[0].text = label; cells[1].text = x_val
        cells[2].text = y_val; cells[3].text = n_val

    # §5.6 Current-speed sensitivity (0..5 knots)
    document.add_heading("5.7. Current-speed sensitivity (0..5 knots)", level=2)
    document.add_paragraph(
        f"Key variables: heading ψ = +{default_heading:.0f}° (fixed) · rudder "
        f"δ = +{default_rudder:.0f}° (fixed) · current speed V swept 0..5 kn. The HTML "
        "report renders two interactive Plotly charts; the static table below summarises "
        f"the same values at each speed. Selected case marker at V = {default_speed:.2f} kn."
    )
    sens_table = document.add_table(rows=1, cols=5)
    hdr = sens_table.rows[0].cells
    hdr[0].text = "V (kn)"; hdr[1].text = "Xc (kN, long.)"
    hdr[2].text = "Yc (kN, port)"; hdr[3].text = "X_rudder (kN)"
    hdr[4].text = "Y_rudder (kN, port)"
    for r in sorted(sensitivity_rows, key=lambda x: x["current_speed_kn"]):
        cells = sens_table.add_row().cells
        cells[0].text = f"{r['current_speed_kn']:.2f}"
        cells[1].text = _kn(r["ocimf_current_force_x_ship_N"])
        cells[2].text = _kn(r["ocimf_current_force_y_ship_port_N"])
        cells[3].text = _kn(r["force_x_ship_N"])
        cells[4].text = _kn(r["force_y_ship_port_N"])
    document.add_paragraph(
        "Both current and rudder loads scale as V² (quadratic in current speed) per the "
        "underlying physics: current load uses q = 0.5·ρ·V²; rudder load uses "
        "F = β·A_R·V²·Cr·sin(α)."
    )

    # §6 Limitations
    document.add_heading("6. Limitations", level=1)
    for line in (
        "Generic-reference OCIMF MEG4 tanker-current coefficients are not vessel-specific to "
        "SIROCCO. The report basis is an off-class screening tier.",
        "Both rudder models are screening-level; neither is a validated rudder hydrodynamic model.",
        "Component sums (X_total = Xc + Xr, etc.) are reported for engineering review; they are "
        "not a validated whole-vessel force balance.",
        metadata["default_speed_policy"],
        metadata["zero_effective_angle_note"],
        "This report excludes: hull current force at oblique headings beyond the generic basis "
        "range, mooring-line stiffness, tug loads, bank effects, current-profile variation, "
        "propeller race, IMO/class compliance conclusions.",
    ):
        document.add_paragraph(line, style="List Bullet")

    # Appendix A — Rudder model comparison (Pass H: relocated from old §5.6)
    document.add_heading("Appendix A — Rudder model comparison (Whicker-Fehlner vs thin-plate)", level=1)
    document.add_paragraph(
        f"Side-by-side comparison of the two screening rudder models at the default case "
        f"(V = {default_speed:.2f} kn, ψ = +{default_heading:.0f}°, δ = +{default_rudder:.0f}°, "
        f"α = +{default_alpha:.0f}°). The HTML report's Appendix A also includes an "
        f"interactive Plotly chart sweeping rudder angle 0..+{default_rudder:.0f}° at fixed "
        f"ψ = +{default_heading:.0f}°, comparing Model A's force_y_ship_port_N and "
        "Model B's simple_plate_force_y_ship_port_N. Both models are screening-level; "
        "neither is a validated rudder hydrodynamic model. Differences at large angles "
        "(>20°) reflect each model's simplifying assumptions and the stall cap in Model B."
    )
    appendix_table = document.add_table(rows=1, cols=3)
    hdr = appendix_table.rows[0].cells
    hdr[0].text = "Quantity"
    hdr[1].text = "Model A — Whicker-Fehlner [3]"
    hdr[2].text = "Model B — thin-plate (Faltinsen [2] §6.5)"
    for label, a_val, b_val in (
        ("Normal force F (N)",
         f"{default_row['normal_force_N']:.0f}",
         f"{default_row['simple_plate_normal_force_N']:.0f}"),
        ("Longitudinal X_rudder (kN)",
         _kn(default_row["force_x_ship_N"]),
         _kn(default_row["simple_plate_force_x_ship_N"])),
        ("Transverse Y_rudder (kN, port)",
         _kn(default_row["force_y_ship_port_N"]),
         _kn(default_row["simple_plate_force_y_ship_port_N"])),
        ("Yaw moment N_rudder (kN·m, +bow-to-port)",
         _kNm(default_row["moment_n_yaw_bow_port_kN_m"]),
         _kNm(default_row["simple_plate_moment_n_yaw_bow_port_kN_m"])),
    ):
        cells = appendix_table.add_row().cells
        cells[0].text = label; cells[1].text = a_val; cells[2].text = b_val

    # References (numbered by first citation in text; standards + literature in [n], project
    # documents and code in [Pn], per consulting-report convention — see Bibliography research
    # 2026-05-22, agent output for #2760 Pass G).
    document.add_heading("References", level=1)
    for citation in (
        "[1]  OCIMF (2018). Mooring Equipment Guidelines, 4th edition (MEG4). "
        "Witherby Publishing Group, Livingston, UK. ISBN 978-1-85609-771-0. "
        "Annex A — Wind and Current Coefficients for VLCC; figures A9 (Cxc), "
        "A10 (Cyc), A11 (Cxyc); loaded-tanker family, water-depth-to-draft "
        "regime applied per §A.1 conventions.",
        "[2]  Faltinsen, O. M. (1990). Sea Loads on Ships and Offshore Structures. "
        "Cambridge University Press, Cambridge Ocean Technology Series. §6.5 — "
        "Rudder forces (thin-plate small-angle approximation Cn(α) ≈ 2π·sin(α)).",
        "[3]  Whicker, L. F. and Fehlner, L. F. (1958). Free-stream characteristics "
        "of a family of low-aspect-ratio, all-movable control surfaces for "
        "application to ship design. David Taylor Model Basin Report 933. "
        "(Basis for the β·A_R·V²·Cr·sin(α) rudder normal-force formulation "
        "used here as Model A.)",
    ):
        document.add_paragraph(citation, style="List Number")
    document.add_heading("Project Documents", level=2)
    for citation in (
        "[P1] B1528 SIROCCO — Vessel Geometry and Rudder Particulars (workspace-hub "
        "issue #2569 source-pack). Project-internal vessel particulars, rudder "
        "dimensions, and Whicker-Fehlner β/Cr inputs. (Proprietary; available on "
        "request.)",
        f"[P2] digitalmodel (open-source). Report generator: "
        f"{links['report_generator']}. Function "
        "run_b1528_current_heading_rudder_report interpolates Annex A figures "
        "from the OCIMF MEG4 workbook held off-repo at the licensed publisher "
        "path; calculation-time fail-closed if the workbook or citation route "
        "cannot be resolved.",
        f"[P3] Approved revision plan: {links['plan']} "
        "(workspace-hub issue #2760 plan, dated 2026-05-20).",
    ):
        document.add_paragraph(citation, style="List Number")

    document.save(str(docx_path))


def _write_pdf_from_html(html_path: Path, pdf_path: Path) -> None:
    """Render the interactive HTML report to a static PDF using Playwright."""

    try:
        from playwright.sync_api import sync_playwright
    except ImportError as exc:  # pragma: no cover - dependency availability is environment-specific
        raise RuntimeError("Playwright is required to render the B1528 SIROCCO PDF report") from exc

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1600, "height": 1000}, device_scale_factor=1)
        page.goto(html_path.resolve().as_uri(), wait_until="networkidle")
        page.emulate_media(media="print")
        page.pdf(path=str(pdf_path), format="A4", landscape=True, print_background=True)
        browser.close()


SCHEMATIC_IDS = (
    "schematic-axes-conventions",
    "schematic-current-loading",
    "schematic-current-moment",
    "schematic-rudder-loading",
)


def _capture_schematic_pngs(html_path: Path, png_dir: Path) -> dict[str, Path]:
    """Capture each per-section schematic SVG as a PNG via Playwright element screenshot.

    Returns a dict of {schematic_id: png_path}. Pass H DOCX embedding consumer.
    Returns empty dict if Playwright is unavailable (DOCX writer degrades gracefully).
    """

    try:
        from playwright.sync_api import sync_playwright
    except ImportError:  # pragma: no cover - degrade gracefully
        return {}

    png_dir.mkdir(parents=True, exist_ok=True)
    captured: dict[str, Path] = {}
    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1200, "height": 1000}, device_scale_factor=2)
        page.goto(html_path.resolve().as_uri(), wait_until="networkidle")
        for schematic_id in SCHEMATIC_IDS:
            try:
                element = page.locator(f"#{schematic_id}")
                element.wait_for(state="visible", timeout=5000)
                png_path = png_dir / f"{schematic_id}.png"
                element.screenshot(path=str(png_path), omit_background=False)
                captured[schematic_id] = png_path
            except Exception:  # pragma: no cover - screenshot failures shouldn't break report gen
                continue
        browser.close()
    return captured


def _load_ocimf_workbook_basis() -> dict[str, list[tuple[float, float]]]:
    """Load the minimal #2760 OCIMF coefficient curves from the licensed workbook."""

    issue_2760_source_preflight()
    workbook_path = _resolve_ocimf_workbook_path()
    return _load_ocimf_workbook_basis_for_path(str(workbook_path), workbook_path.stat().st_mtime_ns)


@lru_cache(maxsize=4)
def _load_ocimf_workbook_basis_for_path(
    workbook_path_text: str, workbook_mtime_ns: int
) -> dict[str, list[tuple[float, float]]]:
    """Read report-specific OCIMF curves; cache is keyed by path and mtime."""

    del workbook_mtime_ns  # value participates in the cache key.
    workbook_path = Path(workbook_path_text)
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover - dependency availability is environment-specific
        raise RuntimeError("openpyxl is required to resolve OCIMF workbook coefficients") from exc

    wb = openpyxl.load_workbook(workbook_path, data_only=True, read_only=True)
    required_sheets = {"Data 5a-9a", "Data 10a-14a"}
    missing_sheets = sorted(required_sheets.difference(wb.sheetnames))
    if missing_sheets:
        raise ValueError(f"OCIMF workbook missing required #2760 sheets: {missing_sheets}")
    sheet_a5_a9 = wb["Data 5a-9a"]
    sheet_a10_a14 = wb["Data 10a-14a"]

    def points(ws: Any, angle_col: int, value_col: int, start_row: int = 3) -> list[tuple[float, float]]:
        curve: list[tuple[float, float]] = []
        for row_idx in range(start_row, ws.max_row + 1):
            angle = ws.cell(row_idx, angle_col).value
            value = ws.cell(row_idx, value_col).value
            if isinstance(angle, (int, float)) and isinstance(value, (int, float)):
                curve.append((float(angle), float(value)))
        if len(curve) < 2:
            raise ValueError(f"OCIMF workbook curve at columns {angle_col}/{value_col} has insufficient numeric data")
        return sorted(curve)

    return {
        # Figure A9: loaded tanker longitudinal current Cxc, WD/T >4.4, conventional bow.
        "cxc": points(sheet_a5_a9, 23, 24),
        # Figure A10: loaded tanker lateral current Cyc, WD/T >6.
        "cyc": points(sheet_a10_a14, 1, 7),
        # Figure A11: loaded tanker yaw moment Cxyc, WD/T >6.
        "cxyc": points(sheet_a10_a14, 9, 15),
    }


OCIMF_CURRENT_COP_LEVER_ARM_FRACTION = 0.3
"""Approximate current center-of-pressure lever arm as a fraction of LBP.

OCIMF MEG3/MEG4 Annex A figures imply a CoP forward of midship at
near-head-current headings; typical loaded-tanker range is 0.2..0.4 LBP.
0.3 is a screening-level mid-range value used only for the §4.2
side-by-side "force × lever arm" yaw moment estimate that compares
against OCIMF's direct Cxyc-based yaw moment. This is NOT an equality
test — the two methods rest on different assumptions per approved
plan §126.
"""


SIMPLE_PLATE_STALL_ANGLE_DEG = 28.0
"""Stall cap for the thin-plate rudder model.

Selected to match the issue #2760 sweep upper bound. Above |α| > stall, the
formula returns the value at ±stall (no growth) to prevent the linear
sin(α)·2π formula from over-predicting in regions where it has no validity.
"""


def _rudder_simple_plate_force_N(
    area_m2: float, velocity_m_s: float, alpha_deg: float, rho_kg_m3: float
) -> float:
    """Thin-plate rudder normal-force screening estimate per Faltinsen 1990 §6.5.

    Reference: Faltinsen, "Sea Loads on Ships and Offshore Structures",
    Cambridge University Press, 1990, §6.5 (thin-plate small-angle
    approximation `Cn(α) ≈ 2π·sin(α)`, with empirical stall cap above the
    angle where the formula loses validity).

    F_normal = 0.5 · ρ · A_R · V² · Cn(α)

    This is a screening-level model independent of the Whicker-Fehlner
    `β·A_R·V²·Cr·sin(α)` basis used as the report default. It is intended
    for side-by-side sanity-checking only; neither model is a validated
    rudder hydrodynamic model. Differences at large angles reflect each
    model's simplifying assumptions.
    """

    effective_alpha_deg = max(
        min(alpha_deg, SIMPLE_PLATE_STALL_ANGLE_DEG),
        -SIMPLE_PLATE_STALL_ANGLE_DEG,
    )
    cn = 2.0 * math.pi * math.sin(math.radians(effective_alpha_deg))
    q = 0.5 * rho_kg_m3 * velocity_m_s ** 2
    return q * area_m2 * cn


def _interp_curve(curve: list[tuple[float, float]], x: float) -> float:
    """Linearly interpolate a numeric OCIMF curve."""

    if x <= curve[0][0]:
        return curve[0][1]
    if x >= curve[-1][0]:
        return curve[-1][1]
    for (x0, y0), (x1, y1) in zip(curve, curve[1:]):
        if x0 <= x <= x1:
            if x1 == x0:
                return y0
            frac = (x - x0) / (x1 - x0)
            return y0 + frac * (y1 - y0)
    raise ValueError(f"interpolation value {x} outside OCIMF curve domain")


def resolve_ocimf_loaded_tanker_current_coefficients(
    cfg: B1528CurrentHeadingRudderConfig,
    heading_offset_deg: float,
) -> dict[str, Any]:
    """Resolve generic/reference OCIMF loaded-tanker current coefficients.

    Issue #2760 defines heading as current off bow, port positive. The OCIMF
    workbook curves used here are tabulated on a 0..180 degree heading basis,
    with head-current cases near 180 degrees, so ±5 degrees off bow maps to 175.
    The workbook publishes Cyc as positive magnitude and Cxyc as signed values
    that encode yaw-moment direction across heading; the per-call sign factor
    only mirrors port vs. starboard heading and MUST preserve the workbook's
    own Cxyc sign (otherwise yaw direction inverts in the Cxyc<0 regime — see
    Annex A figure A11 below ~98° and above ~180° where Cxyc is negative).
    """

    preflight = issue_2760_source_preflight()
    basis_selection = select_ocimf_loaded_tanker_current_basis(cfg)
    curves = _load_ocimf_workbook_basis()
    heading = float(heading_offset_deg)
    sign = 1.0 if heading >= 0.0 else -1.0
    table_angle = 180.0 - abs(heading)
    cxc = _interp_curve(curves["cxc"], table_angle)
    if abs(heading) < 1e-12:
        cyc = 0.0
        cxyc = 0.0
    else:
        cyc = sign * _interp_curve(curves["cyc"], table_angle)
        cxyc = sign * _interp_curve(curves["cxyc"], table_angle)
    return {
        "angle_table_deg": table_angle,
        "cxc": cxc,
        "cyc": cyc,
        "cxyc": cxyc,
        "source_workbook": preflight["ocimf"]["workbook_source_id"],
        "source_code_id": preflight["ocimf"]["citation"].code_id,
        "basis": "OCIMF loaded tanker current curves: A9 Cxc WD/T>4.4 conventional, A10/A11 WD/T>6",
        "basis_selection": basis_selection,
    }


def _row(
    cfg: B1528CurrentHeadingRudderConfig,
    current_speed_kn: float,
    heading_offset_deg: float,
    rudder_angle_deg: float,
) -> dict[str, Any]:
    speed_m_s = current_speed_kn * KNOT_TO_M_PER_S
    base_force_N = cfg.beta * cfg.rudder_area_m2 * speed_m_s**2 * cfg.prop_rotation_factor
    psi_rad = math.radians(heading_offset_deg)
    alpha_deg = rudder_angle_deg - heading_offset_deg
    alpha_rad = math.radians(alpha_deg)
    normal_force_N = base_force_N * math.sin(alpha_rad)
    x_local_N = base_force_N * math.sin(alpha_rad) ** 2
    y_local_N = base_force_N * math.sin(alpha_rad) * math.cos(alpha_rad)
    x_ship_N = x_local_N * math.cos(psi_rad) - y_local_N * math.sin(psi_rad)
    y_ship_N = x_local_N * math.sin(psi_rad) + y_local_N * math.cos(psi_rad)
    n_ship_Nm = y_ship_N * cfg.yaw_lever_m
    dynamic_pressure_Pa = 0.5 * cfg.rho_kg_m3 * speed_m_s**2
    frontal_area_current_m2 = cfg.beam_m * cfg.draft_m
    lateral_area_current_m2 = cfg.lbp_m * cfg.draft_m
    ocimf_coefficients = resolve_ocimf_loaded_tanker_current_coefficients(
        cfg, heading_offset_deg=heading_offset_deg
    )
    ocimf_cx = ocimf_coefficients["cxc"]
    ocimf_cy = ocimf_coefficients["cyc"]
    ocimf_cm = ocimf_coefficients["cxyc"]
    current_x_N = dynamic_pressure_Pa * frontal_area_current_m2 * ocimf_cx
    current_y_N = dynamic_pressure_Pa * lateral_area_current_m2 * ocimf_cy
    current_n_Nm = dynamic_pressure_Pa * lateral_area_current_m2 * cfg.lbp_m * ocimf_cm
    # Method B (Y_current × lever_arm) yaw moment for §4.2 side-by-side
    # comparison against OCIMF direct Method A (Cxyc-based). NOT an equality
    # test — methods rest on different assumptions per plan §126.
    current_cop_lever_arm_m = OCIMF_CURRENT_COP_LEVER_ARM_FRACTION * cfg.lbp_m
    current_y_times_lever_arm_n_Nm = current_y_N * current_cop_lever_arm_m
    # Model B (simple thin-plate, Faltinsen 1990 §6.5) — independent reference
    # for side-by-side comparison only; does NOT contribute to total_* sums
    # (Whicker-Fehlner Model A remains the default rudder force basis per plan §128).
    simple_plate_normal_force_N = _rudder_simple_plate_force_N(
        cfg.rudder_area_m2, speed_m_s, alpha_deg, cfg.rho_kg_m3
    )
    simple_plate_x_local_N = simple_plate_normal_force_N * math.sin(alpha_rad)
    simple_plate_y_local_N = simple_plate_normal_force_N * math.cos(alpha_rad)
    simple_plate_x_ship_N = (
        simple_plate_x_local_N * math.cos(psi_rad)
        - simple_plate_y_local_N * math.sin(psi_rad)
    )
    simple_plate_y_ship_N = (
        simple_plate_x_local_N * math.sin(psi_rad)
        + simple_plate_y_local_N * math.cos(psi_rad)
    )
    simple_plate_n_ship_Nm = simple_plate_y_ship_N * cfg.yaw_lever_m
    total_x_N = current_x_N + x_ship_N
    total_y_N = current_y_N + y_ship_N
    total_n_Nm = current_n_Nm + n_ship_Nm
    is_extra_default = (
        current_speed_kn == cfg.chart_default_current_speed_kn
        and current_speed_kn not in cfg.current_speeds_kn
    )
    return {
        "case_id": cfg.case_id,
        "current_speed_kn": current_speed_kn,
        "is_chart_default_extra_speed": is_extra_default,
        "heading_offset_deg": heading_offset_deg,
        "rudder_angle_deg": rudder_angle_deg,
        "effective_rudder_inflow_angle_deg": alpha_deg,
        "current_speed_m_s": speed_m_s,
        "prop_rotation_factor": cfg.prop_rotation_factor,
        "prop_rotation_factor_basis": cfg.prop_rotation_factor_logic,
        "beta": cfg.beta,
        "rudder_area_m2": cfg.rudder_area_m2,
        "yaw_lever_m": cfg.yaw_lever_m,
        "ocimf_first_cut_dynamic_pressure_Pa": dynamic_pressure_Pa,
        "ocimf_first_cut_frontal_area_current_m2": frontal_area_current_m2,
        "ocimf_first_cut_lateral_area_current_m2": lateral_area_current_m2,
        "ocimf_first_cut_cx_current": ocimf_cx,
        "ocimf_first_cut_cy_current": ocimf_cy,
        "ocimf_first_cut_cm_current": ocimf_cm,
        "ocimf_current_force_x_ship_N": current_x_N,
        "ocimf_current_force_y_ship_port_N": current_y_N,
        "ocimf_current_moment_n_yaw_bow_port_Nm": current_n_Nm,
        "ocimf_current_moment_n_yaw_bow_port_kN_m": current_n_Nm / 1000.0,
        # Pass E: Method B (Y×lever_arm) yaw moment for §4.2 side-by-side
        "current_cop_lever_arm_m": current_cop_lever_arm_m,
        "current_y_times_lever_arm_moment_n_yaw_bow_port_Nm": current_y_times_lever_arm_n_Nm,
        "current_y_times_lever_arm_moment_n_yaw_bow_port_kN_m": current_y_times_lever_arm_n_Nm / 1000.0,
        "base_force_N": base_force_N,
        "normal_force_N": normal_force_N,
        "force_x_local_downstream_N": x_local_N,
        "force_y_local_port_of_current_N": y_local_N,
        "force_x_ship_N": x_ship_N,
        "force_y_ship_port_N": y_ship_N,
        "force_z_heave_N": 0.0,
        "moment_k_roll_Nm": 0.0,
        "moment_m_pitch_Nm": 0.0,
        "moment_n_yaw_bow_port_Nm": n_ship_Nm,
        "moment_n_yaw_bow_port_kN_m": n_ship_Nm / 1000.0,
        # Model B (simple thin-plate, Faltinsen 1990 §6.5) — comparison only
        "simple_plate_normal_force_N": simple_plate_normal_force_N,
        "simple_plate_force_x_ship_N": simple_plate_x_ship_N,
        "simple_plate_force_y_ship_port_N": simple_plate_y_ship_N,
        "simple_plate_moment_n_yaw_bow_port_Nm": simple_plate_n_ship_Nm,
        "simple_plate_moment_n_yaw_bow_port_kN_m": simple_plate_n_ship_Nm / 1000.0,
        "total_force_x_ship_N": total_x_N,
        "total_force_y_ship_port_N": total_y_N,
        "total_moment_n_yaw_bow_port_Nm": total_n_Nm,
        "total_moment_n_yaw_bow_port_kN_m": total_n_Nm / 1000.0,
        "mooring_reaction_x_N": -total_x_N,
        "mooring_reaction_y_N": -total_y_N,
        "mooring_reaction_n_Nm": -total_n_Nm,
        "force_component_basis": "issue #2760 rudder induced plus generic/reference OCIMF tanker-current review comparison; components are summed in ship-fixed COG X/Y/N",
    }


def _metadata(cfg: B1528CurrentHeadingRudderConfig) -> dict[str, Any]:
    return {
        "case_id": cfg.case_id,
        "aliases": list(cfg.aliases),
        "review_target_date": REPORT_REVIEW_TARGET_DATE,
        "scope": REPORT_SCOPE,
        "zero_effective_angle_note": ZERO_EFFECTIVE_ANGLE_NOTE,
        "method": (
            "V=kn*0.51444; F=beta*A_R*V^2*Cr; alpha=delta-psi; "
            "X_local=F*sin(alpha)^2; Y_local=F*sin(alpha)*cos(alpha); "
            "X_ship=X_local*cos(psi)-Y_local*sin(psi); "
            "Y_ship=X_local*sin(psi)+Y_local*cos(psi); N=Y_ship*yaw_lever"
        ),
        "current_speeds_kn": list(cfg.current_speeds_kn),
        "chart_default_current_speed_kn": cfg.chart_default_current_speed_kn,
        "chart_default_extra_speed_included": cfg.chart_default_current_speed_kn not in cfg.current_speeds_kn,
        "default_speed_policy": cfg.default_speed_policy,
        "heading_offsets_deg": list(cfg.heading_offsets_deg),
        "rudder_angles_deg": list(cfg.rudder_angles_deg),
        "design_data": {
            "lbp_m": cfg.lbp_m,
            "beam_m": cfg.beam_m,
            "draft_m": cfg.draft_m,
            "water_depth_m": cfg.water_depth_m,
            "water_depth_to_draft_ratio": cfg.water_depth_m / cfg.draft_m,
            "yaw_lever_m": cfg.yaw_lever_m,
            "rudder_area_m2": cfg.rudder_area_m2,
            "rudder_span_m": cfg.rudder_span_m,
            "ship_sog_kn": cfg.ship_sog_kn,
            "rho_kg_m3": cfg.rho_kg_m3,
            "beta": cfg.beta,
            "prop_rotation_factor": cfg.prop_rotation_factor,
            "force_convention": cfg.force_convention,
            "ocimf_current_cx_basis": "interpolated OCIMF loaded-tanker A9 longitudinal-current curve from approved workbook route",
            "ocimf_current_cm_basis": "interpolated OCIMF loaded-tanker A11 yaw-moment curve from approved workbook route",
            "ocimf_basis_selection": select_ocimf_loaded_tanker_current_basis(cfg),
        },
        "analysis_assumptions": [
            "Vessel is moored with ship speed over ground equal to 0 kn.",
            "Heading/rudder effective-angle convention: α = rudder_angle_deg - heading_offset_deg.",
            "Positive heading rotates the local downstream current-force axis toward port from +X_ship.",
            "Local-to-ship transform rotates local current-frame X/Y loads into ship-fixed COG axes.",
            cfg.default_speed_policy,
            "OCIMF reference hull current loads are first-cut review loads using projected beam*draft and LBP*draft areas with loaded-tanker current coefficients interpolated from the approved OCIMF workbook route; they are not vessel-specific OCIMF/current-coefficient curves or certified coefficients.",
            "Mooring reactions are equal and opposite loads for static-equilibrium context only.",
        ],
        "limitations": list(cfg.limitations),
        "traceability_links": {
            "source_pack_issue": cfg.source_pack_issue,
            "current_heading_rudder_issue": cfg.report_issue,
            "plan": f"{GITHUB_REPO_BLOB}/{cfg.plan_path}",
            "durable_report": f"{GITHUB_REPO_BLOB}/docs/domains/marine-engineering/b1528-sirocco-current-rudder-force-report.md",
            "generated_markdown_report": f"{GITHUB_REPO_BLOB}/outputs/b1528_sirocco/current_rudder_force/{REPORT_ARTIFACT_STEM}_report.md",
            "generated_html_report": f"{GITHUB_REPO_BLOB}/outputs/b1528_sirocco/current_rudder_force/{REPORT_ARTIFACT_STEM}_report.html",
            "packaged_input_yaml": f"{GITHUB_REPO_BLOB}/src/digitalmodel/naval_architecture/data/b1528_sirocco_current_heading_rudder.yml",
            "report_generator": f"{GITHUB_REPO_BLOB}/src/digitalmodel/naval_architecture/b1528_sirocco_current_heading_rudder_report.py",
            "moored_current_reference": f"{GITHUB_REPO_BLOB}/src/digitalmodel/naval_architecture/b1528_sirocco_moored_current_report.py",
        },
        "prop_rotation_factor_note": PROPELLER_ROTATION_FACTOR_NOTE,
    }


def _summary(rows: list[dict[str, Any]]) -> dict[str, Any]:
    max_y = max(rows, key=lambda row: abs(row["force_y_ship_port_N"]))
    max_n = max(rows, key=lambda row: abs(row["moment_n_yaw_bow_port_Nm"]))
    max_total_n = max(rows, key=lambda row: abs(row["total_moment_n_yaw_bow_port_Nm"]))
    return {
        "row_count": len(rows),
        "requested_engineering_row_count": sum(not row["is_chart_default_extra_speed"] for row in rows),
        "extra_default_row_count": sum(row["is_chart_default_extra_speed"] for row in rows),
        "max_abs_ship_sway_force_N": abs(max_y["force_y_ship_port_N"]),
        "max_abs_ship_sway_signed_force_N": max_y["force_y_ship_port_N"],
        "max_abs_ship_sway_case": _case_label(max_y),
        "max_abs_yaw_moment_kN_m": abs(max_n["moment_n_yaw_bow_port_kN_m"]),
        "max_abs_yaw_signed_moment_kN_m": max_n["moment_n_yaw_bow_port_kN_m"],
        "max_abs_yaw_case": _case_label(max_n),
        "max_abs_total_yaw_moment_kN_m": abs(max_total_n["total_moment_n_yaw_bow_port_kN_m"]),
        "max_abs_total_yaw_signed_moment_kN_m": max_total_n["total_moment_n_yaw_bow_port_kN_m"],
        "max_abs_total_yaw_case": _case_label(max_total_n),
    }


def _sample_working_example(result: dict[str, Any]) -> dict[str, Any]:
    row = next(
        item
        for item in result["rows"]
        if item["current_speed_kn"] == 3.08
        and item["heading_offset_deg"] == 5.0
        and item["rudder_angle_deg"] == 28.0
    )
    return {
        "data_point": "issue #2760 default 3.08 kn, heading +5 deg, rudder +28 deg, Cr=1.0",
        "current_speed_kn": row["current_speed_kn"],
        "current_speed_m_s": row["current_speed_m_s"],
        "heading_offset_deg": row["heading_offset_deg"],
        "rudder_angle_deg": row["rudder_angle_deg"],
        "effective_rudder_inflow_angle_deg": row["effective_rudder_inflow_angle_deg"],
        "beta": row["beta"],
        "rudder_area_m2": row["rudder_area_m2"],
        "cr": row["prop_rotation_factor"],
        "base_force_N": row["base_force_N"],
        "normal_force_N": row["normal_force_N"],
        "ocimf_current_force_x_ship_N": row["ocimf_current_force_x_ship_N"],
        "ocimf_current_force_y_ship_port_N": row["ocimf_current_force_y_ship_port_N"],
        "ocimf_current_moment_n_yaw_bow_port_kN_m": row["ocimf_current_moment_n_yaw_bow_port_kN_m"],
        "force_x_ship_N": row["force_x_ship_N"],
        "force_y_ship_port_N": row["force_y_ship_port_N"],
        "moment_n_yaw_bow_port_kN_m": row["moment_n_yaw_bow_port_kN_m"],
        "total_force_x_ship_N": row["total_force_x_ship_N"],
        "total_force_y_ship_port_N": row["total_force_y_ship_port_N"],
        "total_moment_n_yaw_bow_port_kN_m": row["total_moment_n_yaw_bow_port_kN_m"],
    }


def _provenance(result: dict[str, Any]) -> dict[str, Any]:
    metadata = result["metadata"]
    return {
        "scope": metadata["scope"],
        "method": metadata["method"],
        "heading_rudder_effective_angle_convention": "α = rudder_angle_deg - heading_offset_deg",
        "local_to_ship_transform": {
            "x_ship": "X_local*cos(psi)-Y_local*sin(psi)",
            "y_ship": "X_local*sin(psi)+Y_local*cos(psi)",
            "n_ship": "Y_ship*yaw_lever_m",
        },
        "ocimf_first_cut_current_loads": {
            "dynamic_pressure": "q=0.5*rho*V^2",
            "surge_force": "X_current=q*(beam*draft)*Cxc_interp",
            "sway_force": "Y_current=q*(LBP*draft)*Cyc_interp_signed",
            "yaw_moment_about_cog": "N_current=q*(LBP*draft)*LBP*Cxyc_interp_signed",
            "coefficient_caveat": "Generic/reference OCIMF tanker-current coefficients resolved through the approved licensed off-repo workbook route; not vessel-specific SIROCCO current-coefficient curves",
            "component_sum": "total X/Y/N components = current review load + rudder-induced component, summed in ship-fixed COG axes",
        },
        "default_speed_policy": metadata["default_speed_policy"],
        "scope_exclusions": metadata["limitations"],
        "design_data": metadata["design_data"],
        "analysis_assumptions": metadata["analysis_assumptions"],
        "traceability_links": metadata["traceability_links"],
        "sample_working_example": result["sample_working_example"],
        "summary": result["summary"],
    }



def _citation_sidecar(result: dict[str, Any]) -> dict[str, Any]:
    preflight = issue_2760_source_preflight()
    metadata = result["metadata"]
    citations = [
        asdict(preflight["ocimf"]["citation"]),
        asdict(preflight["rudder"]["citation"]),
    ]
    return {
        "citations": citations,
        "ocimf_basis_selection": metadata["design_data"]["ocimf_basis_selection"],
        "license_boundary": preflight["ocimf"]["license_boundary"],
        "source_artifacts": {
            "ocimf_workbook": preflight["ocimf"]["workbook_source_id"],
            "provenance_readme": preflight["ocimf"]["provenance_readme"],
        },
    }


def _markdown_report(result: dict[str, Any]) -> str:
    """Generate Markdown report with the canonical 6-section layout per
    issue #2760 thread comment 4492317819 and approved plan §3.1.

    Sections: (1) Introduction, (2) Design Data & Assumptions, (3) Axes & Sign
    Conventions, (4) Load Due to Current, (5) Load Due to Rudder, (6) Limitations.
    Reading level targets a freshman engineering grad: short sentences, symbols
    defined on first use, "what this is / what it isn't" framing per section.
    """

    metadata = result["metadata"]
    sample = result["sample_working_example"]
    design = metadata["design_data"]
    links = metadata["traceability_links"]

    default_speed = metadata["chart_default_current_speed_kn"]
    default_heading = 5.0
    default_rudder = DEFAULT_CHART_RUDDER_ANGLE_DEG
    default_alpha = default_rudder - default_heading
    table_angle = 180.0 - abs(default_heading)
    v_m_s = default_speed * KNOT_TO_M_PER_S
    q_pa = 0.5 * design["rho_kg_m3"] * v_m_s ** 2
    wd_over_t = design["water_depth_to_draft_ratio"]

    lines = [
        f"# {REPORT_DISPLAY_TITLE}",
        "",
        f"Prepared for engineer review on {metadata['review_target_date']}.",
        "",
        "## 1. Introduction",
        "",
        "The B1528 SIROCCO is a moored vessel. When current flows past it, two separate loads act on the ship:",
        "",
        "1. A **hull current load** — the current pushes on the hull, producing a longitudinal force (X), a transverse force (Y), and a yaw moment (N) about the centre of gravity.",
        "2. A **rudder-induced load** — current also flows past the rudder. When the rudder is held at an angle, it generates an additional force at the stern that contributes to X, Y, and N.",
        "",
        "This report estimates both load components for the approved review case:",
        "",
        f"- Current speed: **{default_speed:.2f} knots**",
        f"- Current heading ψ: **+{default_heading:.0f}° off the bow** (port positive)",
        f"- Rudder angle δ: **+{default_rudder:.0f}°** (port)",
        "- Propeller: **stopped** (rpm = 0; rotation factor Cr = 1.0)",
        "",
        "All forces are reported in ship-fixed axes about the centre of gravity. This is a **screening calculation** intended to support mooring reaction-load reviews — it is not a certified hydrodynamic model. Limitations are in §6.",
        "",
        "## 2. Design Data & Assumptions",
        "",
        "Values used by the calculation are sourced from the packaged B1528 SIROCCO YAML input pack.",
        "",
        "| Parameter | Value |",
        "|---|---:|",
        f"| LBP (length between perpendiculars) | {design['lbp_m']:.1f} m |",
        f"| Beam | {design['beam_m']:.2f} m |",
        f"| Draft | {design['draft_m']:.2f} m |",
        f"| Water depth | {design['water_depth_m']:.1f} m |",
        f"| Water depth / Draft (WD/T) | {wd_over_t:.2f} |",
        "| Centre of gravity (longitudinal datum) | midship |",
        "| Centre of gravity (vertical, above keel) | 6.1 m |",
        f"| Yaw lever arm (rudder force to CoG) | {design['yaw_lever_m']:.2f} m |",
        f"| Rudder area | {design['rudder_area_m2']:.3f} m² |",
        f"| Rudder span | {design['rudder_span_m']:.2f} m |",
        f"| Water density ρ | {design['rho_kg_m3']:.1f} kg/m³ |",
        f"| Whicker-Fehlner constant β | {design['beta']:.1f} |",
        f"| Propeller rotation factor Cr | {design['prop_rotation_factor']:.1f} (rpm = 0) |",
        f"| Current speed (default) | {default_speed:.2f} kn |",
        f"| Current heading ψ (default) | +{default_heading:.0f}° (port positive) |",
        f"| Rudder angle δ (default) | +{default_rudder:.0f}° (port positive) |",
        "",
        "## 3. Axes & Sign Conventions",
        "",
        "All forces and moments are reported in **ship-fixed axes at the centre of gravity** using the right-hand rule.",
        "",
        "| Symbol | Meaning | Sign convention |",
        "|---|---|---|",
        "| **+X** | longitudinal force (along ship centreline) | forward |",
        "| **+Y** | transverse force (lateral, across ship) | port |",
        "| **+N** | yaw moment about CoG | bow-to-port |",
        "| **ψ** | current heading offset from bow | port positive |",
        "| **δ** | rudder angle relative to ship centreline | port positive |",
        "| **α** | effective rudder inflow angle = δ - ψ | (derived) |",
        "",
        "**See the schematic in the HTML report (§3 schematic-axes-conventions)** for the ship outline with arrows showing each axis and angle.",
        "",
        "## 4. Load Due to Current",
        "",
        "The current load on the hull is estimated using loaded-tanker current coefficients from **OCIMF MEG4 (2018) [1], Annex A** (figures A9/A10/A11). These curves are a tanker-class basis — they are not vessel-specific to SIROCCO and are used here as a screening estimate only.",
        "",
        "Three coefficients drive the load:",
        "",
        "- **Cxc** (longitudinal current coefficient) from Annex A figure A9",
        "- **Cyc** (transverse current coefficient) from Annex A figure A10",
        "- **Cxyc** (yaw moment coefficient) from Annex A figure A11",
        "",
        f"Workbook bucket selection: loaded-tanker family, water-depth-to-draft ratio WD/T > 6 (vessel WD/T = {wd_over_t:.2f}). Coefficients are interpolated at the table angle `180° − |ψ|`, which equals **{table_angle:.0f}°** for the default heading ψ = +{default_heading:.0f}°.",
        "",
        "### 4.1. Force calculation",
        "",
        "Dynamic pressure: `q = 0.5 × ρ × V²`  ",
        "Frontal projected area: `A_f = beam × draft`  ",
        "Lateral projected area: `A_l = LBP × draft`",
        "",
        "Force components:",
        "",
        "- **Longitudinal (X):** `Xc = q × A_f × Cxc`",
        "- **Transverse (Y):** `Yc = q × A_l × Cyc`",
        "",
        "**Sample calculation at default values:**",
        "",
        f"- V = {default_speed:.2f} kn × {KNOT_TO_M_PER_S:.5f} = **{v_m_s:.4f} m/s**",
        f"- q = 0.5 × {design['rho_kg_m3']:.0f} kg/m³ × ({v_m_s:.4f} m/s)² = **{q_pa:.1f} Pa**",
        f"- A_f = {design['beam_m']:.2f} × {design['draft_m']:.2f} = **{design['beam_m'] * design['draft_m']:.2f} m²**",
        f"- A_l = {design['lbp_m']:.1f} × {design['draft_m']:.2f} = **{design['lbp_m'] * design['draft_m']:.2f} m²**",
        f"- Cxc({table_angle:.0f}°) ≈ **−0.0324**, Cyc({table_angle:.0f}°) ≈ **+0.0341**",
        f"- **Xc** ≈ {q_pa * design['beam_m'] * design['draft_m'] * -0.0324 / 1000.0:.1f} kN (longitudinal current force on ship)",
        f"- **Yc** ≈ {q_pa * design['lbp_m'] * design['draft_m'] * 0.0341 / 1000.0:.1f} kN (transverse current force on ship)",
        "",
        "### 4.2. Yaw moment about CoG",
        "",
        "Two methods estimate the current yaw moment, shown side by side as a sanity check:",
        "",
        f"- **Method A — OCIMF direct (default):** `Nc_A = q × A_l × LBP × Cxyc({table_angle:.0f}°)`",
        f"- **Method B — force × lever arm:** `Nc_B = Yc × CoP_lever_arm`",
        "",
        "These two methods rest on different assumptions and **may differ**. The HTML report's Chart 4 shows both with a caption that explicitly says this comparison is **not equality-based** — it is a sanity-check overlay.",
        "",
        "**See the per-section schematic in the HTML report (§4 schematic-current-loading and schematic-current-moment).**",
        "",
        "## 5. Load Due to Rudder",
        "",
        "The rudder-induced force comes from current flowing past the rudder blade at the effective inflow angle:",
        "",
        "`α = δ - ψ`",
        "",
        f"At the default case (δ = +{default_rudder:.0f}° port, ψ = +{default_heading:.0f}° port): **α = +{default_alpha:.0f}°**.",
        "",
        "Two screening rudder models are presented side by side (Pass C wires Model B numerically; this Pass A lays out the structure):",
        "",
        f"- **Model A — Whicker-Fehlner normal-force basis** (current default): `F = β × A_R × V² × Cr × sin(α)`",
        f"  - Source: B1528 SIROCCO source pack (`β = {design['beta']:.0f}`, `Cr = {design['prop_rotation_factor']:.1f}`)",
        "  - Captures the bulk of normal force at small-to-moderate α; simple sin(α) projection",
        "- **Model B — thin-plate drag/lift** (Pass C): `Cn(α) ≈ 2π·sin(α)` for small angles, with stall handling at ~25–30°",
        "  - Reference: Faltinsen, *Sea Loads on Ships and Offshore Structures*, 1990, §6.5",
        "  - Independent reference model for sanity-checking Model A",
        "",
        "**Both models are screening-level**; neither is a validated rudder hydrodynamic model. Differences at large angles (>20°) reflect stall-region behaviour and the simplifying assumptions of each.",
        "",
        "**Sample calculation at default values (Whicker-Fehlner - Model A):**",
        "",
        f"- α = δ - ψ = {default_rudder:.0f}° − {default_heading:.0f}° = **{default_alpha:.0f}°**",
        f"- F = β × A_R × V² × Cr = {design['beta']:.0f} × {design['rudder_area_m2']:.3f} × ({v_m_s:.4f})² × {design['prop_rotation_factor']:.1f} = **{sample['base_force_N']:.1f} N**",
        f"- Fn (normal) = F × sin(α) = **{sample['normal_force_N']:.1f} N**",
        f"- Longitudinal X_rudder ≈ **{sample['force_x_ship_N'] / 1000.0:.2f} kN**",
        f"- Transverse Y_rudder ≈ **{sample['force_y_ship_port_N'] / 1000.0:.2f} kN**",
        f"- Yaw moment N_rudder ≈ **{sample['moment_n_yaw_bow_port_kN_m']:.2f} kN·m** (+bow-to-port)",
        "",
        "**See the per-section schematic in the HTML report (§5 schematic-rudder-loading).**",
        "",
        "## 6. Limitations",
        "",
        "- Generic-reference OCIMF tanker-current coefficients are **not vessel-specific to SIROCCO**. The report basis is an off-class screening tier.",
        "- Both rudder models are **screening-level**; neither is a validated rudder hydrodynamic model.",
        "- Component sums (X_total = Xc + Xr, etc.) are reported for engineering review; they are **not a validated whole-vessel force balance**.",
        f"- {metadata['default_speed_policy']}",
        f"- {metadata['zero_effective_angle_note']}",
        "- This report excludes: hull current force at oblique headings beyond the generic basis range, mooring-line stiffness, tug loads, bank effects, current-profile variation, propeller race, IMO/class compliance conclusions.",
        "",
        "Original limitations from the project input pack:",
        "",
        *[f"- {item}" for item in metadata["limitations"]],
        "",
        "---",
        "",
        "## References",
        "",
        "[1] OCIMF (2018). *Mooring Equipment Guidelines*, 4th edition (MEG4). Witherby Publishing Group, Livingston, UK. ISBN 978-1-85609-771-0. Annex A — Wind and Current Coefficients for VLCC; figures A9 (Cxc), A10 (Cyc), A11 (Cxyc); loaded-tanker family, water-depth-to-draft regime per §A.1 conventions.",
        "",
        "[2] Faltinsen, O. M. (1990). *Sea Loads on Ships and Offshore Structures*. Cambridge University Press, Cambridge Ocean Technology Series. §6.5 — Rudder forces (thin-plate small-angle approximation `Cn(α) ≈ 2π·sin(α)`).",
        "",
        "[3] Whicker, L. F. and Fehlner, L. F. (1958). *Free-stream characteristics of a family of low-aspect-ratio, all-movable control surfaces for application to ship design*. David Taylor Model Basin Report 933. (Basis for the `β·A_R·V²·Cr·sin(α)` rudder normal-force formulation used here as Model A.)",
        "",
        "### Project Documents",
        "",
        f"[P1] B1528 SIROCCO — Vessel Geometry and Rudder Particulars ([workspace-hub #2569]({links['source_pack_issue']}) source-pack). Project-internal vessel particulars, rudder dimensions, and Whicker-Fehlner β/Cr inputs. (Proprietary; available on request.)",
        "",
        f"[P2] digitalmodel (open-source). Report generator: [b1528_sirocco_current_heading_rudder_report.py]({links['report_generator']}). Function `run_b1528_current_heading_rudder_report` interpolates Annex A figures from the OCIMF MEG4 workbook held off-repo at the licensed publisher path; calculation-time fail-closed if the workbook or citation route cannot be resolved.",
        "",
        f"[P3] Approved revision plan: [workspace-hub #2760 plan]({links['plan']}) (dated 2026-05-20).",
        "",
        f"[P4] Packaged input YAML: [b1528_sirocco_current_heading_rudder.yml]({links['packaged_input_yaml']}).",
        "",
        f"- Sweep coverage: {result['summary']['row_count']} rows ({result['summary']['requested_engineering_row_count']} engineering + {result['summary']['extra_default_row_count']} chart-default).",
        f"- GitHub issue: [workspace-hub #2760]({links['current_heading_rudder_issue']}).",
    ]
    return "\n".join(lines)


def _html_report(result: dict[str, Any]) -> str:
    metadata = result["metadata"]
    rows_json = json.dumps(result["rows"], separators=(",", ":"))
    speed_options = "\n".join(
        _speed_option(speed, metadata["chart_default_current_speed_kn"])
        for speed in sorted({row["current_speed_kn"] for row in result["rows"]})
    )
    rudder_options = "\n".join(
        f'<option value="{float(angle):.1f}"{ " selected" if float(angle) == DEFAULT_CHART_RUDDER_ANGLE_DEG else ""}>{float(angle):.1f} deg</option>'
        for angle in metadata["rudder_angles_deg"]
    )
    summary_json = json.dumps(result["summary"], indent=2)
    design_data = metadata["design_data"]
    design = design_data  # alias used in restructured section content
    links = metadata["traceability_links"]  # Pass G G6: References section
    sample = result["sample_working_example"]
    # Default-case scalars for the restructured 6-section content
    default_speed = metadata["chart_default_current_speed_kn"]
    default_heading = 5.0
    default_rudder = DEFAULT_CHART_RUDDER_ANGLE_DEG
    default_alpha = default_rudder - default_heading
    table_angle = 180.0 - abs(default_heading)
    v_m_s = default_speed * KNOT_TO_M_PER_S
    q_pa = 0.5 * design["rho_kg_m3"] * v_m_s ** 2
    wd_over_t = design["water_depth_to_draft_ratio"]
    limitations_list_html = "\n".join(
        f"<li>{item}</li>" for item in metadata["limitations"]
    )
    # Pass G G5: pre-render default-case values into static table cells (otherwise PDF
    # readers see "—" placeholders that only update via JS in interactive HTML).
    default_row = next(
        (r for r in result["rows"]
         if abs(r["current_speed_kn"] - default_speed) < 1e-6
         and abs(r["heading_offset_deg"] - default_heading) < 1e-6
         and abs(r["rudder_angle_deg"] - default_rudder) < 1e-6),
        result["rows"][0],
    )
    speed_rows = [r for r in result["rows"] if abs(r["current_speed_kn"] - default_speed) < 1e-6]
    max_y_row = max(speed_rows, key=lambda r: abs(r["force_y_ship_port_N"])) if speed_rows else default_row
    max_n_row = max(speed_rows, key=lambda r: abs(r["moment_n_yaw_bow_port_kN_m"])) if speed_rows else default_row
    max_x_row = max(speed_rows, key=lambda r: abs(r["force_x_ship_N"])) if speed_rows else default_row

    def _fmt_kn_static(value_N: float) -> str:
        kn = value_N / 1000.0
        return f"{kn:.0f}" if abs(kn) >= 1 else f"{kn:.2f}"

    def _fmt_moment_static(value_kNm: float) -> str:
        return f"{value_kNm:.0f}" if abs(value_kNm) >= 1 else f"{value_kNm:.2f}"

    def _case_label_static(row: dict) -> str:
        return f"ψ={row['heading_offset_deg']:.0f}° · δ={row['rudder_angle_deg']:.0f}°"

    envelope_default_speed = f"{default_speed:.2f} kn · {len(speed_rows)} rows"
    envelope_max_y_case = _case_label_static(max_y_row)
    envelope_max_y_value = f"{_fmt_kn_static(max_y_row['force_y_ship_port_N'])} kN"
    envelope_max_n_case = _case_label_static(max_n_row)
    envelope_max_n_value = f"{_fmt_moment_static(max_n_row['moment_n_yaw_bow_port_kN_m'])} kN·m"
    envelope_max_h_case = _case_label_static(max_x_row)
    envelope_max_h_value = f"{_fmt_kn_static(max_x_row['force_x_ship_N'])} kN"

    breakdown_default_current_x = _fmt_kn_static(default_row["ocimf_current_force_x_ship_N"])
    breakdown_default_current_y = _fmt_kn_static(default_row["ocimf_current_force_y_ship_port_N"])
    breakdown_default_current_n = _fmt_moment_static(default_row["ocimf_current_moment_n_yaw_bow_port_kN_m"])
    breakdown_default_rudder_x = _fmt_kn_static(default_row["force_x_ship_N"])
    breakdown_default_rudder_y = _fmt_kn_static(default_row["force_y_ship_port_N"])
    breakdown_default_rudder_n = _fmt_moment_static(default_row["moment_n_yaw_bow_port_kN_m"])
    breakdown_default_total_x = _fmt_kn_static(default_row["total_force_x_ship_N"])
    breakdown_default_total_y = _fmt_kn_static(default_row["total_force_y_ship_port_N"])
    breakdown_default_total_n = _fmt_moment_static(default_row["total_moment_n_yaw_bow_port_kN_m"])
    breakdown_default_label = (
        f"Default case: V={default_speed:.2f} kn · {_case_label_static(default_row)} · "
        f"total yaw reaction for mooring = "
        f"{_fmt_moment_static(-default_row['total_moment_n_yaw_bow_port_kN_m'])} kN·m."
    )

    comp_default_a_normal = f"{default_row['normal_force_N']:.0f}"
    comp_default_a_x = _fmt_kn_static(default_row["force_x_ship_N"])
    comp_default_a_y = _fmt_kn_static(default_row["force_y_ship_port_N"])
    comp_default_a_n = _fmt_moment_static(default_row["moment_n_yaw_bow_port_kN_m"])
    comp_default_b_normal = f"{default_row['simple_plate_normal_force_N']:.0f}"
    comp_default_b_x = _fmt_kn_static(default_row["simple_plate_force_x_ship_N"])
    comp_default_b_y = _fmt_kn_static(default_row["simple_plate_force_y_ship_port_N"])
    comp_default_b_n = _fmt_moment_static(default_row["simple_plate_moment_n_yaw_bow_port_kN_m"])

    # Pass G G2: design data table aligned with MD §2 (English labels, professional
    # precision per Pass D rounding rule, explicit CoG rows, source-pack citation).
    input_rows_html = "\n".join(
        [
            f"<tr><th>LBP (length between perpendiculars)</th><td>{design_data['lbp_m']:.1f} m</td></tr>",
            f"<tr><th>Beam</th><td>{design_data['beam_m']:.2f} m</td></tr>",
            f"<tr><th>Draft</th><td>{design_data['draft_m']:.2f} m</td></tr>",
            f"<tr><th>Water depth</th><td>{design_data['water_depth_m']:.1f} m</td></tr>",
            f"<tr><th>Water depth / Draft (WD/T)</th><td>{wd_over_t:.2f}</td></tr>",
            f"<tr><th>Centre of gravity (longitudinal datum)</th><td>midship</td></tr>",
            f"<tr><th>Centre of gravity (vertical, above keel)</th><td>6.1 m</td></tr>",
            f"<tr><th>Yaw lever arm (rudder force to CoG)</th><td>{design_data['yaw_lever_m']:.2f} m</td></tr>",
            f"<tr><th>Rudder area A_R</th><td>{design_data['rudder_area_m2']:.2f} m²</td></tr>",
            f"<tr><th>Rudder span</th><td>{design_data['rudder_span_m']:.2f} m</td></tr>",
            f"<tr><th>Water density ρ</th><td>{design_data['rho_kg_m3']:.1f} kg/m³</td></tr>",
            f"<tr><th>Whicker-Fehlner constant β <em>(see [3] in References)</em></th><td>{design_data['beta']:.1f}</td></tr>",
            f"<tr><th>Propeller rotation factor Cr</th><td>{design_data['prop_rotation_factor']:.1f} (rpm = 0)</td></tr>",
            f"<tr><th>Current speed (default)</th><td>{default_speed:.2f} kn</td></tr>",
            f"<tr><th>Current heading ψ (default)</th><td>+{default_heading:.0f}° (port positive)</td></tr>",
            f"<tr><th>Rudder angle δ (default)</th><td>+{default_rudder:.0f}° (port positive)</td></tr>",
            f"<tr><th>Current speed sweep range</th><td>{min(metadata['current_speeds_kn']):.0f}..{max(metadata['current_speeds_kn']):.0f} kn (plotted bound, not the design value)</td></tr>",
            f"<tr><th>Heading sweep</th><td>{min(metadata['heading_offsets_deg']):.0f}° .. +{max(metadata['heading_offsets_deg']):.0f}° in 1° steps</td></tr>",
            f"<tr><th>Rudder sweep</th><td>{min(metadata['rudder_angles_deg']):.0f}° .. +{max(metadata['rudder_angles_deg']):.0f}° port in 2° steps</td></tr>",
            f"<tr><th>Input source</th><td>Project Document — B1528 SIROCCO vessel particulars source-pack ([P1] in References)</td></tr>",
        ]
    )
    # Pass G G3: sample calc rewrite — English labels, rounded display values, professional
    # citations inline. Values match Pass D rounding rule (kN/kN·m to 0 decimals when ≥1).
    sample_rows_html = "\n".join(
        [
            f"<li>Speed conversion: <code>V = {sample['current_speed_kn']:.2f} kn × {KNOT_TO_M_PER_S:.5f} = {sample['current_speed_m_s']:.4f} m/s</code></li>",
            f"<li>Effective rudder inflow angle: <code>α = δ - ψ = {sample['rudder_angle_deg']:.0f}° - {sample['heading_offset_deg']:.0f}° = {sample['effective_rudder_inflow_angle_deg']:.0f}°</code></li>",
            f"<li>Whicker-Fehlner base force (Model A): <code>F = β × A_R × V² × Cr = {sample['beta']:.0f} × {sample['rudder_area_m2']:.2f} × ({sample['current_speed_m_s']:.4f})² × {sample['cr']:.1f} ≈ {sample['base_force_N']/1000.0:.0f} kN</code></li>",
            f"<li>Rudder normal force: <code>Fn = F × sin(α) ≈ {sample['normal_force_N']/1000.0:.0f} kN</code></li>",
            f"<li>Rudder load resolved to ship-fixed axes at CoG: <strong>longitudinal X_rudder</strong> ≈ <code>{sample['force_x_ship_N']/1000.0:.0f} kN</code>, <strong>transverse Y_rudder</strong> ≈ <code>{sample['force_y_ship_port_N']/1000.0:.0f} kN</code> (+port), <strong>yaw moment N_rudder</strong> ≈ <code>{sample['moment_n_yaw_bow_port_kN_m']:.0f} kN·m</code> (+bow-to-port)</li>",
            f"<li>Current load (OCIMF MEG4 [1] Annex A Fig. A9/A10/A11 interpolation at table angle {table_angle:.0f}°): <strong>X_current</strong> ≈ <code>{sample['ocimf_current_force_x_ship_N']/1000.0:.0f} kN</code> (longitudinal), <strong>Y_current</strong> ≈ <code>{sample['ocimf_current_force_y_ship_port_N']/1000.0:.0f} kN</code> (+port), <strong>N_current</strong> ≈ <code>{sample['ocimf_current_moment_n_yaw_bow_port_kN_m']:.0f} kN·m</code> (+bow-to-port)</li>",
            f"<li>Combined at CoG (sum of current + rudder components): <strong>X_total</strong> ≈ <code>{sample['total_force_x_ship_N']/1000.0:.0f} kN</code>, <strong>Y_total</strong> ≈ <code>{sample['total_force_y_ship_port_N']/1000.0:.0f} kN</code>, <strong>N_total</strong> ≈ <code>{sample['total_moment_n_yaw_bow_port_kN_m']:.0f} kN·m</code> — these sums are reported for engineering review only; not a validated whole-vessel force balance</li>",
        ]
    )
    return f"""<!doctype html>
<html lang=\"en\">
<head>
<meta charset=\"utf-8\">
<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
<title>{REPORT_DISPLAY_TITLE}</title>
<script src=\"https://cdn.plot.ly/plotly-2.35.2.min.js\"></script>
<style>
:root {{ --page-bg:#eef2f7; --paper:#fff; --ink:#172033; --muted:#607085; --line:#d7dee8; --accent:#155e95; --soft:#e8f2fb; }}
* {{ box-sizing:border-box; }}
body {{ margin:0; background:var(--page-bg); color:var(--ink); font-family:Aptos, 'Segoe UI', system-ui, sans-serif; line-height:1.55; }}
.report-shell {{ width:min(1180px, calc(100% - 56px)); margin:0 auto; padding:34px 0 50px; }}
.report-page {{ background:var(--paper); border:1px solid var(--line); border-radius:10px; padding:36px 44px; }}
h1 {{ margin:0 0 10px; font-size:2rem; }}
h2 {{ margin:28px 0 10px; font-size:1.25rem; }}
.eyebrow {{ color:var(--accent); font-size:.78rem; font-weight:700; letter-spacing:.08em; text-transform:uppercase; }}
.note-panel {{ background:var(--soft); border-left:4px solid var(--accent); padding:12px 16px; margin:16px 0; }}
.controls {{ display:flex; flex-wrap:wrap; gap:14px; align-items:end; padding:14px; border:1px solid var(--line); border-radius:8px; background:#f8fafc; }}
.controls label {{ display:flex; flex-direction:column; gap:5px; font-weight:700; color:var(--muted); }}
select {{ min-width:190px; padding:8px 10px; border:1px solid var(--line); border-radius:6px; background:white; color:var(--ink); }}
.chart {{ width:100%; height:460px; margin:16px 0 30px; border:1px solid var(--line); border-radius:8px; background:white; }}
.data-table {{ width:100%; border-collapse:collapse; margin:12px 0; font-size:.92rem; }}
.data-table th,.data-table td {{ border-bottom:1px solid var(--line); padding:8px 10px; text-align:left; vertical-align:top; }}
.method-block {{ background:#f8fafc; border:1px solid var(--line); border-radius:6px; padding:14px 16px; white-space:pre-wrap; overflow-x:auto; }}
.print-section {{ break-inside:avoid; page-break-inside:avoid; }}
.print-page {{ break-after:page; page-break-after:always; }}
.chart-section {{ break-inside:avoid; page-break-inside:avoid; margin-top:18px; }}
.calc-list {{ margin:10px 0 0 20px; padding:0; }}
.calc-list li {{ margin:8px 0; }}
.schematic-card {{ border:1px solid var(--line); border-radius:10px; background:#fbfdff; padding:16px; margin:16px 0; }}
.schematic-grid {{ display:grid; grid-template-columns:minmax(320px, 1fr) minmax(240px, .62fr); gap:18px; align-items:center; }}
.schematic-svg {{ width:100%; min-height:300px; border:1px solid var(--line); border-radius:8px; background:linear-gradient(180deg,#f8fbff 0%,#eef6ff 100%); }}
.schematic-legend {{ margin:0; padding-left:20px; color:var(--muted); }}
.schematic-legend li {{ margin:8px 0; }}
.schematic-readout {{ display:grid; grid-template-columns:1fr 1fr; gap:8px; margin-top:12px; }}
.schematic-readout div {{ background:white; border:1px solid var(--line); border-radius:6px; padding:8px 10px; }}
.axis-line {{ stroke:#172033; stroke-width:2; stroke-dasharray:5 5; }}
.current-line {{ stroke:#0b6bcb; stroke-width:4; marker-end:url(#arrow-blue); }}
.rudder-line {{ stroke:#c2410c; stroke-width:4; marker-end:url(#arrow-orange); }}
.force-line {{ stroke:#15803d; stroke-width:3; marker-end:url(#arrow-green); }}
.ship-hull {{ fill-opacity:0.18; fill:#cfe2f3; stroke:#172033; stroke-width:2.5; }}
.ship-hull-transparent {{ fill:#b87060; fill-opacity:0.85; stroke:#1a2a4a; stroke-width:0.75; }}
.rudder-blade {{ fill:#fed7aa; stroke:#c2410c; stroke-width:2; }}
.rudder-blade-transparent {{ fill:#1a2a4a; fill-opacity:0.85; stroke:#1a2a4a; stroke-width:0.75; }}
/* OCIMF Annex A schematic style — hairline navy arrows, symbol-only labels.
   Per OCIMF MEG3/MEG4 Annex A "Figure 1 / Sign Convention" canonical style:
   solid terracotta hull, hairline 0.75pt stroke, 1.5pt force arrows with
   small slim arrowheads, symbols-only on figure, English mapping in caption. */
.ocimf-arrow {{ stroke:#1a2a4a; stroke-width:1.5; fill:none; }}
.ocimf-arrow-current {{ stroke:#1a2a4a; stroke-width:1.5; fill:none; }}
.ocimf-arrow-force {{ stroke:#1a2a4a; stroke-width:1.5; fill:none; }}
.ocimf-arrow-rudder {{ stroke:#1a2a4a; stroke-width:1.5; fill:none; }}
.ocimf-arrow-yaw {{ stroke:#1a2a4a; stroke-width:1.5; fill:none; }}
/* Legacy bold-arrow classes retained — used by §3 interactive schematic */
.schematic-bold-arrow-blue {{ stroke:#0b6bcb; stroke-width:5; fill:none; }}
.schematic-bold-arrow-orange {{ stroke:#c2410c; stroke-width:5; fill:none; }}
.schematic-bold-arrow-green {{ stroke:#15803d; stroke-width:5; fill:none; }}
.schematic-bold-arrow-purple {{ stroke:#7c3aed; stroke-width:5; fill:none; }}
.lever-arm {{ stroke:#1a2a4a; stroke-width:0.75; stroke-dasharray:4 3; fill:none; }}
.svg-label {{ font:600 12px 'Times New Roman', serif; fill:#1a2a4a; }}
.svg-label-ocimf {{ font:600 11px 'Times New Roman', serif; fill:#1a2a4a; }}
.svg-label-symbol {{ font:italic 600 12px 'Times New Roman', serif; fill:#1a2a4a; }}
.svg-label-blue {{ font:700 13px Aptos, Segoe UI, sans-serif; fill:#0b6bcb; }}
.svg-label-orange {{ font:700 13px Aptos, Segoe UI, sans-serif; fill:#c2410c; }}
.svg-label-green {{ font:700 13px Aptos, Segoe UI, sans-serif; fill:#15803d; }}
.svg-label-purple {{ font:700 13px Aptos, Segoe UI, sans-serif; fill:#7c3aed; }}
.svg-muted {{ font:11px 'Times New Roman', serif; fill:#1a2a4a; }}
.svg-default-annot {{ font:700 12px 'Times New Roman', serif; fill:#1a2a4a; }}
.per-section-schematic {{ width:100%; max-width:380px; min-height:420px; border:1px solid var(--line); border-radius:4px; background:#ffffff; display:block; margin:12px auto; }}
.schematic-caption {{ font:11px Aptos, Segoe UI, sans-serif; color:#1a2a4a; background:#f8fafc; border:1px solid var(--line); border-radius:4px; padding:10px 14px; margin:6px auto 18px; max-width:680px; }}
.schematic-caption strong {{ color:#1a2a4a; }}
.schematic-caption code {{ font-family:'Times New Roman', serif; font-style:italic; }}
@media print {{ @page {{ size:A4 landscape; margin:12mm; }} body {{ background:white; }} .report-shell {{ width:100%; padding:0; }} .report-page {{ border:0; padding:0; }} .print-section,.chart-section {{ break-inside:avoid; page-break-inside:avoid; }} .chart-section {{ break-after:page; page-break-after:always; }} .chart {{ break-inside:avoid; height:300px; margin-bottom:18px; }} .schematic-grid {{ grid-template-columns:1fr 1fr; }} }}
</style>
</head>
<body>
<div class=\"report-shell\"><main class=\"report-page\">
<p class=\"eyebrow\">B1528 SIROCCO · Issue #2760</p>
<h1>{REPORT_DISPLAY_HEADING}</h1>

<section id=\"introduction-section\" class=\"print-section\">
<h2>1. Introduction</h2>
<p>The B1528 SIROCCO is a moored vessel. When current flows past it, two separate loads act on the ship:</p>
<ol>
<li>A <strong>hull current load</strong> — the current pushes on the hull, producing a longitudinal force (X), a transverse force (Y), and a yaw moment (N) about the centre of gravity.</li>
<li>A <strong>rudder-induced load</strong> — current also flows past the rudder. When the rudder is held at an angle, it generates an additional force at the stern that contributes to X, Y, and N.</li>
</ol>
<p>This report estimates both load components for the approved review case:</p>
<ul>
<li>Current speed: <strong>{default_speed:.2f} knots</strong></li>
<li>Current heading <code>ψ</code>: <strong>+{default_heading:.0f}° off the bow</strong> (port positive)</li>
<li>Rudder angle <code>δ</code>: <strong>+{default_rudder:.0f}°</strong> (port)</li>
<li>Propeller: <strong>stopped</strong> (rpm = 0; rotation factor Cr = 1.0)</li>
</ul>
<p>All forces are reported in ship-fixed axes about the centre of gravity. This is a <strong>screening calculation</strong> intended to support mooring reaction-load reviews — it is not a certified hydrodynamic model. Limitations are in §6.</p>
<div class=\"note-panel\"><strong>Default speed policy:</strong> {metadata['default_speed_policy']}</div>
</section>

<section id=\"design-data-section\" class=\"print-section\">
<h2>2. Design Data &amp; Assumptions</h2>
<p>Values used by the calculation are sourced from the packaged B1528 SIROCCO YAML input pack.</p>
<table class=\"data-table\" aria-label=\"Design data for B1528 SIROCCO current-heading rudder calculation\">
<tbody>
{input_rows_html}
</tbody>
</table>
</section>

<section id=\"axes-conventions-section\" class=\"print-section schematic-card\">
<h2>3. Axes &amp; Sign Conventions</h2>
<p>All forces and moments are reported in <strong>ship-fixed axes at the centre of gravity</strong> using the right-hand rule.</p>
<table class=\"data-table\" aria-label=\"Axes and sign conventions\">
<thead><tr><th>Symbol</th><th>Meaning</th><th>Sign convention</th></tr></thead>
<tbody>
<tr><td><strong>+X</strong></td><td>longitudinal force (along ship centreline)</td><td>forward</td></tr>
<tr><td><strong>+Y</strong></td><td>transverse force (lateral, across ship)</td><td>port</td></tr>
<tr><td><strong>+N</strong></td><td>yaw moment about CoG</td><td>bow-to-port</td></tr>
<tr><td><code>ψ</code></td><td>current heading offset from bow</td><td>port positive</td></tr>
<tr><td><code>δ</code></td><td>rudder angle relative to ship centreline</td><td>port positive</td></tr>
<tr><td><code>α</code></td><td>effective rudder inflow angle = δ - ψ</td><td>(derived)</td></tr>
</tbody>
</table>
<svg id=\"schematic-axes-conventions\" class=\"per-section-schematic\" viewBox=\"0 0 400 500\" role=\"img\" aria-label=\"Figure 0 — Axes and sign conventions per OCIMF MEG4 Annex A\">
<defs>
<marker id=\"arrow-ac\" markerWidth=\"5\" markerHeight=\"5\" refX=\"4.5\" refY=\"2.5\" orient=\"auto\" markerUnits=\"strokeWidth\"><path d=\"M0,0 L0,5 L5,2.5 z\" fill=\"#1a2a4a\" /></marker>
</defs>
<!-- Cardinal heading labels at frame edges per OCIMF Annex A §A.1 -->
<text x=\"200\" y=\"24\" text-anchor=\"middle\" class=\"svg-label-symbol\">180°</text>
<text x=\"380\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">90°</text>
<text x=\"20\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">270°</text>
<text x=\"200\" y=\"485\" text-anchor=\"middle\" class=\"svg-label-symbol\">0°</text>
<!-- Solid muted-terracotta ship silhouette, OCIMF house style -->
<path class=\"ship-hull-transparent\" d=\"M200 70 C232 95 250 140 250 220 C250 305 235 380 200 425 C165 380 150 305 150 220 C150 140 168 95 200 70 Z\" />
<!-- Hairline ship-fixed axes through CoG -->
<line x1=\"200\" y1=\"56\" x2=\"200\" y2=\"460\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<line x1=\"60\" y1=\"250\" x2=\"340\" y2=\"250\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<!-- Current direction line (JS rotates this group) -->
<g id=\"schematic-current-heading-line\" transform=\"rotate(0 200 250)\">
<line x1=\"200\" y1=\"60\" x2=\"200\" y2=\"100\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-ac)\" />
</g>
<!-- Rudder pivot + blade (JS rotates this rect) -->
<circle cx=\"200\" cy=\"420\" r=\"2\" fill=\"#1a2a4a\" />
<rect id=\"schematic-rudder-blade\" x=\"196\" y=\"420\" width=\"8\" height=\"32\" rx=\"1.5\" class=\"rudder-blade-transparent\" transform=\"rotate(0 200 420)\" />
<!-- Hidden JS-target groups (preserved so updateHeadingRudderSchematic doesn't error) -->
<g id=\"schematic-rudder-line\" transform=\"rotate(0 200 420)\" style=\"display:none\"><line x1=\"200\" y1=\"420\" x2=\"200\" y2=\"455\" class=\"ocimf-arrow\" /></g>
<g id=\"schematic-total-force-line\" transform=\"rotate(0 200 250)\" style=\"display:none\"><line x1=\"200\" y1=\"250\" x2=\"160\" y2=\"195\" class=\"ocimf-arrow\" /></g>
<!-- Force/moment symbol cluster at CoG, OCIMF symbol-only style -->
<circle cx=\"200\" cy=\"250\" r=\"2.5\" fill=\"#1a2a4a\" />
<text x=\"205\" y=\"246\" class=\"svg-label-symbol\">CoG</text>
<line x1=\"200\" y1=\"250\" x2=\"200\" y2=\"210\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-ac)\" />
<text x=\"206\" y=\"222\" class=\"svg-label-symbol\">+F_X</text>
<line x1=\"200\" y1=\"250\" x2=\"150\" y2=\"250\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-ac)\" />
<text x=\"138\" y=\"244\" class=\"svg-label-symbol\">+F_Y</text>
<path d=\"M 222 250 A 22 22 0 0 0 200 272\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-ac)\" />
<text x=\"228\" y=\"276\" class=\"svg-label-symbol\">+M_XY</text>
<!-- ψ/δ/α labels (JS updates these text contents) -->
<text id=\"schematic-psi-label\" x=\"300\" y=\"24\" class=\"svg-label-symbol\">ψ = 0°</text>
<text id=\"schematic-delta-label\" x=\"260\" y=\"460\" class=\"svg-label-symbol\">δ = 0°</text>
<text id=\"schematic-alpha-label\" x=\"320\" y=\"460\" class=\"svg-label-symbol\">α = 0°</text>
<text id=\"schematic-force-label\" x=\"30\" y=\"24\" class=\"svg-label-symbol\">total F</text>
</svg>
<div class=\"schematic-caption\">
<strong>Figure 0 — Axes &amp; Sign Conventions.</strong> Top-down plan view, ship-fixed CoG frame per OCIMF MEG4 [1] Annex A §A.1. Cardinal numerals at frame edges (0° stern, 180° bow, 90° starboard, 270° port) define heading convention; <code>ψ</code> = current heading offset from bow (port positive), <code>δ</code> = rudder angle (port positive), <code>α = δ - ψ</code> = effective rudder inflow angle.
<code>+F_X</code> = longitudinal force component; <code>+F_Y</code> = transverse force component (positive port); <code>+M_XY</code> = yaw moment about CoG (positive bow-to-port, right-hand rule).
The current-direction arrow rotates with the heading offset and the rudder blade rotates with the rudder angle in the interactive HTML; PDF/DOCX shows the schematic at the selected default-case angles.
<div class=\"schematic-readout\" aria-label=\"Selected schematic geometry readout\" style=\"display:grid; grid-template-columns:repeat(4, 1fr); gap:8px; margin-top:8px\">
<div><strong>Speed</strong><br><span id=\"schematic-speed-readout\">—</span></div>
<div><strong>Heading ψ</strong><br><span id=\"schematic-heading-readout\">—</span></div>
<div><strong>Rudder δ</strong><br><span id=\"schematic-rudder-readout\">—</span></div>
<div><strong>Effective α</strong><br><span id=\"schematic-alpha-readout\">—</span></div>
</div>
</div>
</section>

<section id=\"load-due-to-current-section\" class=\"print-section\">
<h2>4. Load Due to Current</h2>
<p>The current load on the hull is estimated using <strong>generic-reference OCIMF MEG3/MEG4 loaded-tanker current coefficients</strong> interpolated from the licensed off-repo workbook. These curves are a tanker-class basis — they are <strong>not vessel-specific to SIROCCO</strong>.</p>
<p>Three coefficients drive the load: <strong>Cxc</strong> (longitudinal) from Annex A figure A9, <strong>Cyc</strong> (transverse) from A10, and <strong>Cxyc</strong> (yaw moment) from A11. Workbook bucket selection: loaded-tanker family, water-depth-to-draft ratio WD/T &gt; 6 (vessel WD/T = {wd_over_t:.2f}). Coefficients are interpolated at the table angle <code>180° − |ψ|</code> = <strong>{table_angle:.0f}°</strong> for the default ψ = +{default_heading:.0f}°.</p>

<h3>4.1. Force calculation</h3>
<p>Dynamic pressure <code>q = 0.5 × ρ × V²</code>; frontal area <code>A_f = beam × draft</code>; lateral area <code>A_l = LBP × draft</code>. Force components: <strong>longitudinal Xc = q × A_f × Cxc</strong>; <strong>transverse Yc = q × A_l × Cyc</strong>.</p>
<div class=\"schematic-card\">
<svg id=\"schematic-current-loading\" class=\"per-section-schematic\" viewBox=\"0 0 400 500\" role=\"img\" aria-label=\"Figure 1 — Current loading per OCIMF MEG4 Annex A sign convention, B1528 SIROCCO at default case\">
<defs>
<marker id=\"arrow-cl\" markerWidth=\"5\" markerHeight=\"5\" refX=\"4.5\" refY=\"2.5\" orient=\"auto\" markerUnits=\"strokeWidth\"><path d=\"M0,0 L0,5 L5,2.5 z\" fill=\"#1a2a4a\" /></marker>
</defs>
<!-- Cardinal heading labels at frame edges per OCIMF convention (NOT arc on hull) -->
<text x=\"200\" y=\"24\" text-anchor=\"middle\" class=\"svg-label-symbol\">180°</text>
<text x=\"380\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">90°</text>
<text x=\"20\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">270°</text>
<text x=\"200\" y=\"485\" text-anchor=\"middle\" class=\"svg-label-symbol\">0°</text>
<!-- Solid muted-terracotta ship silhouette (OCIMF house style) -->
<path class=\"ship-hull-transparent\" d=\"M200 70 C232 95 250 140 250 220 C250 305 235 380 200 425 C165 380 150 305 150 220 C150 140 168 95 200 70 Z\" />
<!-- Hairline ship-fixed axes through CoG (X = centerline, Y = perpendicular) -->
<line x1=\"200\" y1=\"56\" x2=\"200\" y2=\"460\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<line x1=\"60\" y1=\"250\" x2=\"340\" y2=\"250\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<!-- CoG marker (small filled dot) -->
<circle cx=\"200\" cy=\"250\" r=\"2.5\" fill=\"#1a2a4a\" />
<!-- Force/moment symbols at CoG cluster, OCIMF symbol-only style -->
<!-- +F_X longitudinal pointing toward bow (aft of CoG since Cxc<0; small arrow) -->
<line x1=\"200\" y1=\"250\" x2=\"200\" y2=\"205\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-cl)\" />
<text x=\"206\" y=\"218\" class=\"svg-label-symbol\">+F_X</text>
<!-- +F_Y transverse pointing port (left) — larger since Cyc>0 dominates -->
<line x1=\"200\" y1=\"250\" x2=\"125\" y2=\"250\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-cl)\" />
<text x=\"118\" y=\"244\" class=\"svg-label-symbol\">+F_Y</text>
<!-- ψ current-attack inset (separate cluster per OCIMF convention) -->
<line x1=\"230\" y1=\"6\" x2=\"210\" y2=\"66\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-cl)\" />
<text x=\"238\" y=\"20\" class=\"svg-label-symbol\">ψ</text>
<!-- Inline scale-bar style L_OA / B dimension hint (small, OCIMF style) -->
<text x=\"200\" y=\"248\" text-anchor=\"middle\" class=\"svg-label-symbol\">CoG</text>
</svg>
<div class=\"schematic-caption\">
<strong>Figure 1 — Sign Convention &amp; Current Loading.</strong> Top-down plan view, B1528 SIROCCO at default case (ψ=+{default_heading:.0f}°, V={default_speed:.2f} kn).
<code>+F_X</code> = longitudinal current force on hull (component of <code>q·A_f·Cxc</code> along ship centreline; sign per OCIMF MEG4 [1] Annex A Fig. A9, Cxc≈−0.0324 at table angle {table_angle:.0f}°).
<code>+F_Y</code> = transverse current force on hull (<code>q·A_l·Cyc</code> across ship; positive port; OCIMF MEG4 [1] Annex A Fig. A10, Cyc≈+0.0341).
Cardinal labels at frame edges (0° stern / 180° bow / 90° starboard / 270° port) define heading convention per OCIMF Annex A §A.1. ψ inset shows the current angle-of-attack relative to bow. CoG at midship; vertical CoG 6.1 m above keel.
</div>
<p><strong>Sample calculation at default values:</strong></p>
<ul class=\"calc-list\">
<li><code>V = {default_speed:.2f} kn × {KNOT_TO_M_PER_S:.5f} = {v_m_s:.4f} m/s</code></li>
<li><code>q = 0.5 × {design['rho_kg_m3']:.0f} × ({v_m_s:.4f})² = {q_pa:.1f} Pa</code></li>
<li><code>A_f = {design['beam_m']:.2f} × {design['draft_m']:.2f} = {design['beam_m'] * design['draft_m']:.2f} m²</code></li>
<li><code>A_l = {design['lbp_m']:.1f} × {design['draft_m']:.2f} = {design['lbp_m'] * design['draft_m']:.2f} m²</code></li>
<li><code>Cxc({table_angle:.0f}°) ≈ −0.0324</code>, <code>Cyc({table_angle:.0f}°) ≈ +0.0341</code></li>
<li><strong>Xc</strong> ≈ <code>{(q_pa * design['beam_m'] * design['draft_m'] * -0.0324) / 1000.0:.1f} kN</code> (longitudinal current force)</li>
<li><strong>Yc</strong> ≈ <code>{(q_pa * design['lbp_m'] * design['draft_m'] * 0.0341) / 1000.0:.1f} kN</code> (transverse current force)</li>
</ul>

<h3>4.2. Yaw moment about CoG</h3>
<p><strong>Transverse current force <code>Yc</code> input</strong> — Method B (force × lever arm) multiplies the transverse current force <code>Yc</code> by a center-of-pressure lever arm. The chart below shows <code>Yc</code> across the heading sweep -5° to +5° in 1° steps at the selected current speed (default {default_speed:.2f} kn). This is the <code>Yc</code> input that drives Method B's yaw moment estimate.</p>
<div id=\"current-transverse-force-chart\" class=\"chart\" aria-label=\"Transverse current force Yc versus heading offset psi at selected current speed\"></div>
<p>Two methods estimate the current yaw moment, shown side by side as a sanity check:</p>
<ul>
<li><strong>Method A — OCIMF direct (default):</strong> <code>Nc_A = q × A_l × LBP × Cxyc({table_angle:.0f}°)</code></li>
<li><strong>Method B — force × lever arm:</strong> <code>Nc_B = Yc × CoP_lever_arm</code></li>
</ul>
<p>These two methods rest on different assumptions and <strong>may differ</strong>. Chart 4 below shows both as an overlay with an explicit non-equality caption.</p>
<div class=\"schematic-card\">
<svg id=\"schematic-current-moment\" class=\"per-section-schematic\" viewBox=\"0 0 400 500\" role=\"img\" aria-label=\"Figure 2 — Current yaw moment about CoG, OCIMF Annex A sign convention with center-of-pressure lever arm\">
<defs>
<marker id=\"arrow-cm\" markerWidth=\"5\" markerHeight=\"5\" refX=\"4.5\" refY=\"2.5\" orient=\"auto\" markerUnits=\"strokeWidth\"><path d=\"M0,0 L0,5 L5,2.5 z\" fill=\"#1a2a4a\" /></marker>
</defs>
<!-- Cardinal heading labels at frame edges -->
<text x=\"200\" y=\"24\" text-anchor=\"middle\" class=\"svg-label-symbol\">180°</text>
<text x=\"380\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">90°</text>
<text x=\"20\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">270°</text>
<text x=\"200\" y=\"485\" text-anchor=\"middle\" class=\"svg-label-symbol\">0°</text>
<!-- Solid muted-terracotta ship silhouette -->
<path class=\"ship-hull-transparent\" d=\"M200 70 C232 95 250 140 250 220 C250 305 235 380 200 425 C165 380 150 305 150 220 C150 140 168 95 200 70 Z\" />
<!-- Hairline ship-fixed axes through CoG -->
<line x1=\"200\" y1=\"56\" x2=\"200\" y2=\"460\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<line x1=\"60\" y1=\"250\" x2=\"340\" y2=\"250\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<!-- CoG marker -->
<circle cx=\"200\" cy=\"250\" r=\"2.5\" fill=\"#1a2a4a\" />
<text x=\"205\" y=\"246\" class=\"svg-label-symbol\">CoG</text>
<!-- Center of pressure marker (forward of CoG per OCIMF tanker convention) -->
<circle cx=\"200\" cy=\"185\" r=\"2\" fill=\"#1a2a4a\" />
<text x=\"205\" y=\"183\" class=\"svg-label-symbol\">CoP</text>
<!-- Lever arm from CoP to CoG (dashed hairline) -->
<line x1=\"200\" y1=\"189\" x2=\"200\" y2=\"246\" class=\"lever-arm\" />
<text x=\"165\" y=\"222\" class=\"svg-label-symbol\">arm</text>
<!-- F_Y transverse force at CoP, symbols-only -->
<line x1=\"200\" y1=\"185\" x2=\"140\" y2=\"185\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-cm)\" />
<text x=\"132\" y=\"178\" class=\"svg-label-symbol\">+F_Y</text>
<!-- +M_XY yaw moment curved arrow at CoG (OCIMF canonical symbol) -->
<path d=\"M 224 250 A 24 24 0 0 0 200 274\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-cm)\" />
<text x=\"232\" y=\"278\" class=\"svg-label-symbol\">+M_XY</text>
</svg>
<div class=\"schematic-caption\">
<strong>Figure 2 — Yaw Moment about CoG.</strong> Two estimation methods for current yaw moment, shown side by side as a sanity check (NOT an equality test).
<strong>Method A (default):</strong> <code>M_XY,A = q·A_l·LBP·Cxyc(table_angle)</code> per OCIMF MEG4 [1] Annex A Fig. A11 (loaded-tanker WD/T&gt;6 family).
<strong>Method B (comparison):</strong> <code>M_XY,B = F_Y × CoP_arm</code> where CoP_arm ≈ {OCIMF_CURRENT_COP_LEVER_ARM_FRACTION:g}·LBP is an approximate center-of-pressure lever arm for loaded-tanker near-head current (OCIMF MEG4 [1] Annex A guidance, typical range 0.2–0.4·LBP). Methods rest on different assumptions and may differ — see §4.2.1 chart.
</div>

<h4>4.2.1. Side-by-side yaw moment chart — Method A vs Method B</h4>
<p>The chart below overlays both methods across the current-heading sweep at the default rudder angle. Method A uses the OCIMF direct yaw moment coefficient Cxyc; Method B multiplies the OCIMF transverse force Yc by an approximate center-of-pressure lever arm (~{OCIMF_CURRENT_COP_LEVER_ARM_FRACTION:g}·LBP). <strong>These methods rest on different assumptions and may differ — this overlay is a sanity check, NOT an equality test.</strong></p>
<div id=\"yaw-side-by-side-chart\" class=\"chart\" aria-label=\"Yaw moment side-by-side — Method A OCIMF direct vs Method B Y times lever arm\"></div>

<h3>4.3. Interactive: current vs rudder vs total side-by-side</h3>
<p>Use the controls below to switch current speed and rudder angle and see how the OCIMF current component, the rudder component, and their sum combine.</p>
<div class=\"controls\">
<label for=\"current-speed-select\">Current speed
<select id=\"current-speed-select\">{speed_options}</select>
</label>
<label for=\"rudder-angle-select\">Rudder angle
<select id=\"rudder-angle-select\">{rudder_options}</select>
</label>
</div>
<div id=\"ocimf-rudder-component-force-chart\" class=\"chart\" aria-label=\"OCIMF current versus rudder versus total Y force chart\"></div>
<div id=\"ocimf-rudder-component-yaw-chart\" class=\"chart\" aria-label=\"OCIMF current versus rudder versus total yaw moment about COG chart\"></div>
</section>

<section id=\"load-due-to-rudder-section\" class=\"print-section\">
<h2>5. Load Due to Rudder</h2>
<p>The rudder-induced force comes from current flowing past the rudder blade at the effective inflow angle <code>α = δ - ψ</code>. At the default case (δ = +{default_rudder:.0f}° port, ψ = +{default_heading:.0f}° port): <strong>α = +{default_alpha:.0f}°</strong>.</p>
<p>Two screening rudder models are presented (Pass C wires Model B numerically; this Pass A lays out the structure):</p>
<ul>
<li><strong>Model A — Whicker-Fehlner normal-force basis</strong> (current default): <code>F = β × A_R × V² × Cr × sin(α)</code>. Source: B1528 SIROCCO source pack (<code>β = {design['beta']:.0f}</code>, <code>Cr = {design['prop_rotation_factor']:.1f}</code>).</li>
<li><strong>Model B — thin-plate drag/lift</strong> (Pass C): <code>Cn(α) ≈ 2π·sin(α)</code> for small angles, with stall handling at ~25–30°. Reference: Faltinsen, <em>Sea Loads on Ships and Offshore Structures</em>, 1990, §6.5.</li>
</ul>
<p><strong>Both models are screening-level</strong>; neither is a validated rudder hydrodynamic model. Differences at large angles (&gt;20°) reflect stall-region behaviour.</p>
<div class=\"schematic-card\">
<svg id=\"schematic-rudder-loading\" class=\"per-section-schematic\" viewBox=\"0 0 400 500\" role=\"img\" aria-label=\"Figure 3 — Rudder loading, OCIMF Annex A sign convention with rudder at default 28 deg port\">
<defs>
<marker id=\"arrow-rl\" markerWidth=\"5\" markerHeight=\"5\" refX=\"4.5\" refY=\"2.5\" orient=\"auto\" markerUnits=\"strokeWidth\"><path d=\"M0,0 L0,5 L5,2.5 z\" fill=\"#1a2a4a\" /></marker>
</defs>
<!-- Cardinal heading labels at frame edges -->
<text x=\"200\" y=\"24\" text-anchor=\"middle\" class=\"svg-label-symbol\">180°</text>
<text x=\"380\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">90°</text>
<text x=\"20\" y=\"253\" text-anchor=\"middle\" class=\"svg-label-symbol\">270°</text>
<text x=\"200\" y=\"485\" text-anchor=\"middle\" class=\"svg-label-symbol\">0°</text>
<!-- Solid muted-terracotta ship silhouette -->
<path class=\"ship-hull-transparent\" d=\"M200 70 C232 95 250 140 250 220 C250 305 235 380 200 425 C165 380 150 305 150 220 C150 140 168 95 200 70 Z\" />
<!-- Hairline ship-fixed axes through CoG -->
<line x1=\"200\" y1=\"56\" x2=\"200\" y2=\"460\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<line x1=\"60\" y1=\"250\" x2=\"340\" y2=\"250\" stroke=\"#1a2a4a\" stroke-width=\"0.5\" stroke-dasharray=\"3 3\" />
<!-- CoG marker -->
<circle cx=\"200\" cy=\"250\" r=\"2.5\" fill=\"#1a2a4a\" />
<text x=\"205\" y=\"246\" class=\"svg-label-symbol\">CoG</text>
<!-- Rudder pivot at stern (small dark mark) -->
<circle cx=\"200\" cy=\"420\" r=\"2\" fill=\"#1a2a4a\" />
<!-- Rudder blade deflected δ=+28° port (rotated counter-clockwise in this +X-up view) -->
<rect class=\"rudder-blade-transparent\" x=\"196\" y=\"420\" width=\"8\" height=\"34\" rx=\"1.5\" transform=\"rotate(-28 200 420)\" />
<text x=\"210\" y=\"446\" class=\"svg-label-symbol\">δ</text>
<!-- ψ current-attack inset (separate cluster per OCIMF convention) -->
<line x1=\"230\" y1=\"6\" x2=\"210\" y2=\"66\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-rl)\" />
<text x=\"238\" y=\"20\" class=\"svg-label-symbol\">ψ</text>
<!-- F normal-force at rudder (perpendicular to blade, hairline) -->
<line x1=\"200\" y1=\"420\" x2=\"148\" y2=\"398\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-rl)\" />
<text x=\"114\" y=\"393\" class=\"svg-label-symbol\">F</text>
<!-- Decomposition arrows from rudder pivot to ship-fixed X/Y components, hairline -->
<line x1=\"200\" y1=\"420\" x2=\"170\" y2=\"420\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-rl)\" />
<text x=\"152\" y=\"414\" class=\"svg-label-symbol\">+F_Y</text>
<line x1=\"200\" y1=\"420\" x2=\"200\" y2=\"445\" class=\"ocimf-arrow\" marker-end=\"url(#arrow-rl)\" />
<text x=\"205\" y=\"442\" class=\"svg-label-symbol\">+F_X</text>
<!-- α effective inflow inset (small, top-right) -->
<text x=\"320\" y=\"24\" class=\"svg-label-symbol\">α</text>
</svg>
<div class=\"schematic-caption\">
<strong>Figure 3 — Rudder Loading.</strong> Top-down view with rudder at default deflection δ=+{default_rudder:.0f}° port, current at ψ=+{default_heading:.0f}° off the bow. Effective rudder inflow angle <code>α = δ - ψ = +{default_alpha:.0f}°</code>.
<code>F</code> = rudder normal force (perpendicular to blade); decomposed at rudder pivot into <code>+F_X</code> (longitudinal, along centreline) and <code>+F_Y</code> (transverse, port).
Two models compute <code>F</code>: <strong>Model A</strong> (Whicker-Fehlner [3], <code>F = β·A_R·V²·Cr·sin(α)</code>) and <strong>Model B</strong> (Faltinsen [2] §6.5, <code>F = 0.5·ρ·A_R·V²·2π·sin(α)</code> with stall cap). Comparison in §5.6.
</div>

<h3>5.1. Sample calculation at default values (Whicker-Fehlner - Model A)</h3>
<p>Representative hand-check point: <strong>{sample['data_point']}</strong>.</p>
<ol class=\"calc-list\">
{sample_rows_html}
</ol>

<h3>5.2. Rudder force with rudder angle</h3>
<p><strong>Key variables for this section:</strong> current speed <code>V</code> = <strong>{default_speed:.2f} kn</strong> (default) · heading <code>ψ</code> = <strong>+{default_heading:.0f}°</strong> (fixed) · rudder sweep <code>δ</code> ∈ [-{default_rudder:.0f}°, +{default_rudder:.0f}°] at 2° steps. The chart below shows the rudder-induced <strong>longitudinal X_rudder</strong> and <strong>transverse Y_rudder</strong> force vs rudder angle (both decomposed at CoG, port-positive). Selected case (δ = +{default_rudder:.0f}°) marked with a star.</p>
<div id=\"rudder-force-vs-angle-chart\" class=\"chart\" aria-label=\"Rudder force longitudinal and transverse vs rudder angle from -28 to +28 deg\"></div>

<h3>5.3. Rudder force over heading</h3>
<p>Shows rudder-induced <strong>longitudinal X</strong> and <strong>transverse Y</strong> in kN for the selected current speed and rudder angle. Default rudder angle: <strong>{default_rudder:.0f}° port</strong>.</p>
<div id=\"force-components-chart\" class=\"chart\" aria-label=\"Rudder-induced ship-fixed force component chart\"></div>

<h3>5.4. Rudder yaw moment over heading</h3>
<p>Shows signed rudder-induced <strong>yaw moment N</strong> (kN·m, +bow-to-port) over heading offset and rudder angle for the selected current speed.</p>
<div id=\"yaw-moment-chart\" class=\"chart\" aria-label=\"Rudder-induced yaw moment chart\"></div>

<h3>5.5. Selected-speed envelope summary</h3>
<p><strong>Key variables for this section:</strong> current speed <code>V</code> = <strong>{default_speed:.2f} kn</strong> (default) · heading sweep <code>ψ</code> ∈ [-{abs(default_heading):.0f}°, +{default_heading:.0f}°] at 1° steps · rudder sweep <code>δ</code> ∈ [0°, +{default_rudder:.0f}°] at 2° steps. Envelope is the absolute-max across the full heading × rudder grid at the selected speed.</p>
<p>Envelope values update with the current-speed selector and use the full heading × rudder grid at that speed. The selected default is <strong>{default_speed:.2f} kn</strong>; 5 kn is the upper bound of the sensitivity range.</p>
<table id=\"selected-speed-envelope-summary\" class=\"data-table\" aria-label=\"Selected speed envelope summary\">
<thead><tr><th>Metric</th><th>Envelope case</th><th>Value</th></tr></thead>
<tbody>
<tr><td>Selected speed</td><td id=\"selected-speed-case\">engineering sweep plane</td><td id=\"selected-speed-value\">{envelope_default_speed}</td></tr>
<tr><td>Max |Y| at selected speed</td><td id=\"selected-speed-max-y-case\">{envelope_max_y_case}</td><td id=\"selected-speed-max-y-value\">{envelope_max_y_value}</td></tr>
<tr><td>Max |N| at selected speed</td><td id=\"selected-speed-max-n-case\">{envelope_max_n_case}</td><td id=\"selected-speed-max-n-value\">{envelope_max_n_value}</td></tr>
<tr><td>Max |X| at selected speed</td><td id=\"selected-speed-max-h-case\">{envelope_max_h_case}</td><td id=\"selected-speed-max-h-value\">{envelope_max_h_value}</td></tr>
</tbody>
</table>
<p class=\"svg-muted\" style=\"margin-top:4px; font-size:11px\">Initial values are at the default current speed ({default_speed:.2f} kn); the interactive HTML updates the row above when the speed selector changes.</p>

<h3>5.6. Selected-case force breakdown</h3>
<p><strong>Key variables for this section:</strong> current speed <code>V</code> = <strong>{default_speed:.2f} kn</strong> (default; user-selectable via dropdown above) · heading <code>ψ</code> = <strong>+{default_heading:.0f}°</strong> · rudder <code>δ</code> = <strong>+{default_rudder:.0f}°</strong> · effective rudder inflow <code>α</code> = δ − ψ = <strong>+{default_alpha:.0f}°</strong>. Static rows below are at the default case; JS updates them when the user changes the speed or rudder selector.</p>
<p>Updates with current speed and rudder angle. Values shown at the heading with maximum absolute total yaw moment for the selected trace, so individual X/Y/N component paths can be reviewed together.</p>
<table class=\"data-table\" aria-label=\"Selected-case OCIMF current rudder component force breakdown\">
<thead><tr><th>Component</th><th>X (kN, longitudinal)</th><th>Y (kN, transverse port)</th><th>N (kN·m, yaw bow-to-port)</th></tr></thead>
<tbody>
<tr><td>Current (OCIMF MEG4 [1])</td><td id=\"breakdown-current-x\">{breakdown_default_current_x}</td><td id=\"breakdown-current-y\">{breakdown_default_current_y}</td><td id=\"breakdown-current-n\">{breakdown_default_current_n}</td></tr>
<tr><td>Rudder (Whicker-Fehlner [3])</td><td id=\"breakdown-rudder-x\">{breakdown_default_rudder_x}</td><td id=\"breakdown-rudder-y\">{breakdown_default_rudder_y}</td><td id=\"breakdown-rudder-n\">{breakdown_default_rudder_n}</td></tr>
<tr><td>Total (current + rudder) / mooring reaction = negated</td><td id=\"breakdown-total-x\">{breakdown_default_total_x}</td><td id=\"breakdown-total-y\">{breakdown_default_total_y}</td><td id=\"breakdown-total-n\">{breakdown_default_total_n}</td></tr>
</tbody>
</table>
<p id=\"breakdown-case-label\" class=\"note-panel\">{breakdown_default_label}</p>

<h3>5.7. Current-speed sensitivity (0..5 knots)</h3>
<p><strong>Key variables for this section:</strong> heading <code>ψ</code> = <strong>+{default_heading:.0f}°</strong> (fixed) · rudder <code>δ</code> = <strong>+{default_rudder:.0f}°</strong> (fixed) · current speed <code>V</code> sweep <strong>0..5 kn</strong>. The plots below show how the current load and rudder load scale with current speed at the default heading/rudder; the selected default case (V = {default_speed:.2f} kn) is marked on each plot.</p>
<div id=\"sensitivity-current-load-chart\" class=\"chart\" aria-label=\"Current load sensitivity to current speed (longitudinal Xc and transverse Yc vs V)\"></div>
<div id=\"sensitivity-rudder-load-chart\" class=\"chart\" aria-label=\"Rudder load sensitivity to current speed (longitudinal X_rudder and transverse Y_rudder vs V)\"></div>
<p class=\"svg-muted\" style=\"font-size:11px; margin-top:4px\">Sensitivity plots use the engineering sweep at fixed heading +{default_heading:.0f}° and rudder +{default_rudder:.0f}°. Both current and rudder loads scale as V² (quadratic in current speed) per the underlying physics.</p>
</section>

<section id=\"limitations-section\" class=\"print-section\">
<h2>6. Limitations</h2>
<ul>
<li>Generic-reference OCIMF tanker-current coefficients are <strong>not vessel-specific to SIROCCO</strong>. The report basis is an off-class screening tier.</li>
<li>Both rudder models are <strong>screening-level</strong>; neither is a validated rudder hydrodynamic model.</li>
<li>Component sums (X_total = Xc + Xr, etc.) are reported for engineering review; they are <strong>not a validated whole-vessel force balance</strong>.</li>
<li>{ZERO_EFFECTIVE_ANGLE_NOTE}</li>
<li>This report excludes: hull current force at oblique headings beyond the generic basis range, mooring-line stiffness, tug loads, bank effects, current-profile variation, propeller race, IMO/class compliance conclusions.</li>
{limitations_list_html}
</ul>
</section>

<section id=\"method-provenance-section\" class=\"print-section\">
<h2>Method &amp; Provenance</h2>
<pre class=\"method-block\">{metadata['method']}</pre>
<ul>
<li>Heading/rudder effective-angle convention: α = δ − ψ.</li>
<li>Local-to-ship transform rotates local current-frame X/Y loads into ship-fixed CoG axes.</li>
<li>Current review equations: q=0.5·ρ·V²; Cxc, Cyc, Cxyc interpolated from OCIMF MEG4 (2018) [1] Annex A figures A9/A10/A11, loaded-tanker WD/T&gt;6 family; Xc=q·A_f·Cxc; Yc=q·A_l·Cyc; Nc=q·A_l·LBP·Cxyc.</li>
<li>Rudder Model A (default) per Whicker &amp; Fehlner [3]: F = β·A_R·V²·Cr·sin(α).</li>
<li>Rudder Model B (comparison) per Faltinsen [2] §6.5: F = 0.5·ρ·A_R·V²·Cn(α) with Cn(α) ≈ 2π·sin(α) and stall cap at ±28°.</li>
<li>Scope note: hull current loads are screening estimates using OCIMF tanker-class coefficients; not a validated whole-vessel current-load or oblique-current hull/rudder interaction model.</li>
</ul>
<pre class=\"method-block\">{summary_json}</pre>
</section>

<section id=\"appendix-a-rudder-model-comparison\" class=\"print-section\">
<h2>Appendix A — Rudder model comparison (Whicker-Fehlner vs thin-plate)</h2>
<p>Side-by-side comparison of the two screening rudder models at the default case (V={default_speed:.2f} kn, ψ=+{default_heading:.0f}°, δ=+{default_rudder:.0f}°, α=+{default_alpha:.0f}°). The chart below sweeps rudder angle 0..+{default_rudder:.0f}° at fixed ψ=+{default_heading:.0f}° and the selected current speed, comparing Model A and Model B transverse force <code>Y_rudder</code>.</p>
<table id=\"rudder-model-comparison-table\" class=\"data-table\" aria-label=\"Side-by-side rudder model A vs model B at default case\">
<thead><tr><th>Quantity</th><th>Model A — Whicker-Fehlner [3]</th><th>Model B — thin-plate (Faltinsen [2] §6.5)</th></tr></thead>
<tbody>
<tr><td>Normal force F (N)</td><td id=\"rudder-comp-a-normal\">{comp_default_a_normal}</td><td id=\"rudder-comp-b-normal\">{comp_default_b_normal}</td></tr>
<tr><td>Longitudinal X_rudder (kN)</td><td id=\"rudder-comp-a-x\">{comp_default_a_x}</td><td id=\"rudder-comp-b-x\">{comp_default_b_x}</td></tr>
<tr><td>Transverse Y_rudder (kN, port)</td><td id=\"rudder-comp-a-y\">{comp_default_a_y}</td><td id=\"rudder-comp-b-y\">{comp_default_b_y}</td></tr>
<tr><td>Yaw moment N_rudder (kN·m, +bow-to-port)</td><td id=\"rudder-comp-a-n\">{comp_default_a_n}</td><td id=\"rudder-comp-b-n\">{comp_default_b_n}</td></tr>
</tbody>
</table>
<p class=\"note-panel\"><strong>Both models are screening-level</strong>; neither is a validated rudder hydrodynamic model. Differences at large angles (&gt;20°) reflect each model's simplifying assumptions and the stall cap in Model B.</p>
<div id=\"rudder-model-comparison-chart\" class=\"chart\" aria-label=\"Rudder model A vs model B Y force vs rudder angle\"></div>
</section>

<section id=\"references-section\" class=\"print-section\">
<h2>References</h2>
<ol style=\"padding-left:24px; line-height:1.6\">
<li>OCIMF (2018). <em>Mooring Equipment Guidelines</em>, 4th edition (MEG4). Witherby Publishing Group, Livingston, UK. ISBN 978-1-85609-771-0. Annex A — Wind and Current Coefficients for VLCC; figures A9 (Cxc longitudinal), A10 (Cyc transverse), A11 (Cxyc yaw moment); loaded-tanker family, water-depth-to-draft regime applied per §A.1 conventions.</li>
<li>Faltinsen, O. M. (1990). <em>Sea Loads on Ships and Offshore Structures</em>. Cambridge University Press, Cambridge Ocean Technology Series. §6.5 — Rudder forces (thin-plate small-angle approximation Cn(α) ≈ 2π·sin(α)).</li>
<li>Whicker, L. F. and Fehlner, L. F. (1958). <em>Free-stream characteristics of a family of low-aspect-ratio, all-movable control surfaces for application to ship design</em>. David Taylor Model Basin Report 933. (Basis for the β·A_R·V²·Cr·sin(α) rudder normal-force formulation used as Model A.)</li>
</ol>
<h3>Project Documents</h3>
<ol style=\"padding-left:24px; line-height:1.6\">
<li><strong>[P1]</strong> B1528 SIROCCO — Vessel Geometry and Rudder Particulars (<a href=\"https://github.com/vamseeachanta/workspace-hub/issues/2569\">workspace-hub issue #2569</a> source-pack). Project-internal vessel particulars, rudder dimensions, and Whicker-Fehlner β/Cr inputs. (Proprietary; available on request.)</li>
<li><strong>[P2]</strong> digitalmodel (open-source). Report generator: <a href=\"{links['report_generator']}\">b1528_sirocco_current_heading_rudder_report.py</a>. Function <code>run_b1528_current_heading_rudder_report</code> interpolates Annex A figures from the OCIMF MEG4 workbook held off-repo at the licensed publisher path; calculation-time fail-closed if the workbook or citation route cannot be resolved.</li>
<li><strong>[P3]</strong> Approved revision plan: <a href=\"{links['plan']}\">workspace-hub issue #2760 plan</a> (dated 2026-05-20).</li>
</ol>
</section>
</main></div>
<script>
const ROWS = {rows_json};
const STANDARD_CHART_HEIGHT = 460;
const CHART_CONFIG = {{responsive: true, displaylogo: false}};
function selectedSpeed() {{ return parseFloat(document.getElementById('current-speed-select').value); }}
function selectedRudder() {{ return parseFloat(document.getElementById('rudder-angle-select').value); }}
function same(a, b) {{ return Math.abs(a - b) < 1e-9; }}
function byHeading(a, b) {{ return a.heading_offset_deg - b.heading_offset_deg; }}
function rowsForSpeed(speed) {{ return ROWS.filter(row => same(row.current_speed_kn, speed)); }}
function absMax(rows, key) {{ return rows.reduce((best, row) => Math.abs(row[key]) > Math.abs(best[key]) ? row : best, rows[0]); }}
function maxBy(rows, key) {{ return rows.reduce((best, row) => row[key] > best[key] ? row : best, rows[0]); }}
function caseLabel(row) {{ return `heading ${{row.heading_offset_deg}} deg · rudder ${{row.rudder_angle_deg}} deg`; }}
function updateText(id, text) {{ document.getElementById(id).textContent = text; }}
function updateSelectedSpeedEnvelope() {{
  const speed = selectedSpeed();
  const speedRows = rowsForSpeed(speed);
  const maxY = absMax(speedRows, 'force_y_ship_port_N');
  const maxN = absMax(speedRows, 'moment_n_yaw_bow_port_kN_m');
  const maxH = absMax(speedRows, 'force_x_ship_N');
  const extra = speedRows.some(row => row.is_chart_default_extra_speed) ? 'chart-default extra plane' : 'requested engineering sweep plane';
  updateText('selected-speed-case', extra);
  updateText('selected-speed-value', `${{speed}} kn · ${{speedRows.length}} rows`);
  updateText('selected-speed-max-y-case', caseLabel(maxY));
  updateText('selected-speed-max-y-value', `${{fmtKn(maxY.force_y_ship_port_N)}} kN`);
  updateText('selected-speed-max-n-case', caseLabel(maxN));
  updateText('selected-speed-max-n-value', `${{fmtMoment(maxN.moment_n_yaw_bow_port_kN_m)}} kN-m`);
  updateText('selected-speed-max-h-case', caseLabel(maxH));
  updateText('selected-speed-max-h-value', `${{fmtKn(maxH.force_x_ship_N)}} kN`);
}}
function updateForceChart() {{
  const speed = selectedSpeed();
  const rudder = selectedRudder();
  const rows = ROWS.filter(row => same(row.current_speed_kn, speed) && same(row.rudder_angle_deg, rudder)).sort(byHeading);
  const x = rows.map(row => row.heading_offset_deg);
  const hover = rows.map(row => `α=${{row.effective_rudder_inflow_angle_deg}}°<br>X_local=${{fmtKn(row.force_x_local_downstream_N)}} kN<br>Y_local=${{fmtKn(row.force_y_local_port_of_current_N)}} kN<br>N=${{fmtMoment(row.moment_n_yaw_bow_port_kN_m)}} kN·m`);
  const traces = [
    {{x, y: rows.map(row => row.force_x_ship_N/1000), name:'Longitudinal X (kN)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Longitudinal X=%{{y:.0f}} kN<br>%{{customdata}}<extra></extra>'}},
    {{x, y: rows.map(row => row.force_y_ship_port_N/1000), name:'Transverse Y (kN, port)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Transverse Y=%{{y:.0f}} kN<br>%{{customdata}}<extra></extra>'}}
  ];
  Plotly.newPlot('force-components-chart', traces, {{title:`Rudder-induced longitudinal X and transverse Y force (kN) · V=${{speed}} kn · δ=${{rudder}}°`, xaxis:{{title:'Heading offset ψ (deg)', zeroline:true, gridcolor:'#e5eaf1'}}, yaxis:{{title:'Rudder force component (kN)', zeroline:true, gridcolor:'#e5eaf1'}}, height:STANDARD_CHART_HEIGHT, margin:{{l:76,r:30,t:66,b:64}}, legend:{{orientation:'h', y:-0.22}}, hovermode:'x unified', template:'plotly_white'}}, CHART_CONFIG);
}}
function updateYawChart() {{
  const speed = selectedSpeed();
  const speedRows = rowsForSpeed(speed);
  const headings = [...new Set(speedRows.map(row => row.heading_offset_deg))].sort((a,b)=>a-b);
  const rudders = [...new Set(speedRows.map(row => row.rudder_angle_deg))].sort((a,b)=>a-b);
  const z = rudders.map(rudder => headings.map(heading => {{
    const row = speedRows.find(item => same(item.heading_offset_deg, heading) && same(item.rudder_angle_deg, rudder));
    return row ? row.moment_n_yaw_bow_port_kN_m : null;
  }}));
  const rudder = selectedRudder();
  const rows = ROWS.filter(row => same(row.current_speed_kn, speed) && same(row.rudder_angle_deg, rudder)).sort(byHeading);
  const trace = {{x: rows.map(row => row.heading_offset_deg), y: rows.map(row => row.moment_n_yaw_bow_port_kN_m), name:'Rudder yaw moment N (kN·m, +bow-to-port)', mode:'lines+markers', hovertemplate:'ψ=%{{x}}°<br>Yaw moment N=%{{y:.0f}} kN·m<extra></extra>'}};
  Plotly.newPlot('yaw-moment-chart', [trace], {{title:`Rudder-induced yaw moment N about CoG (kN·m) · V=${{speed}} kn · δ=${{rudder}}°`, xaxis:{{title:'Heading offset ψ (deg)', gridcolor:'#e5eaf1'}}, yaxis:{{title:'Yaw moment N about CoG (kN·m, +bow-to-port)', gridcolor:'#e5eaf1'}}, height:STANDARD_CHART_HEIGHT, margin:{{l:76,r:40,t:66,b:64}}, template:'plotly_white'}}, CHART_CONFIG);
}}
function selectedTraceRows() {{
  const speed = selectedSpeed();
  const rudder = selectedRudder();
  return ROWS.filter(row => same(row.current_speed_kn, speed) && same(row.rudder_angle_deg, rudder)).sort(byHeading);
}}
function updateComponentForceChart() {{
  const speed = selectedSpeed();
  const rudder = selectedRudder();
  const rows = selectedTraceRows();
  const x = rows.map(row => row.heading_offset_deg);
  const hover = rows.map(row => `Xc=${{fmtKn(row.ocimf_current_force_x_ship_N)}} kN<br>Yc=${{fmtKn(row.ocimf_current_force_y_ship_port_N)}} kN<br>Xr=${{fmtKn(row.force_x_ship_N)}} kN<br>Yr=${{fmtKn(row.force_y_ship_port_N)}} kN<br>Xt=${{fmtKn(row.total_force_x_ship_N)}} kN<br>Yt=${{fmtKn(row.total_force_y_ship_port_N)}} kN`);
  const traces = [
    {{x, y: rows.map(row => row.ocimf_current_force_y_ship_port_N/1000), name:'Current (OCIMF)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Current Y=%{{y:.0f}} kN<br>%{{customdata}}<extra></extra>'}},
    {{x, y: rows.map(row => row.force_y_ship_port_N/1000), name:'Rudder (Whicker-Fehlner)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Rudder Y=%{{y:.0f}} kN<br>%{{customdata}}<extra></extra>'}},
    {{x, y: rows.map(row => row.total_force_y_ship_port_N/1000), name:'Total (current + rudder)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Total Y=%{{y:.0f}} kN<br>%{{customdata}}<extra></extra>'}}
  ];
  Plotly.newPlot('ocimf-rudder-component-force-chart', traces, {{title:`Transverse force Y (kN, port) — current vs rudder vs total · V=${{speed}} kn · δ=${{rudder}}°`, xaxis:{{title:'Heading offset ψ (deg)', zeroline:true, gridcolor:'#e5eaf1'}}, yaxis:{{title:'Transverse force Y (kN, port)', rangemode:'tozero', gridcolor:'#e5eaf1'}}, height:STANDARD_CHART_HEIGHT, margin:{{l:76,r:30,t:66,b:64}}, legend:{{orientation:'h', y:-0.22}}, hovermode:'x unified', template:'plotly_white'}}, CHART_CONFIG);
}}
function updateComponentYawChart() {{
  const speed = selectedSpeed();
  const rudder = selectedRudder();
  const rows = selectedTraceRows();
  const x = rows.map(row => row.heading_offset_deg);
  const hover = rows.map(row => `Nc=${{fmtMoment(row.ocimf_current_moment_n_yaw_bow_port_kN_m)}} kN·m<br>Nr=${{fmtMoment(row.moment_n_yaw_bow_port_kN_m)}} kN·m<br>Nt=${{fmtMoment(row.total_moment_n_yaw_bow_port_kN_m)}} kN·m`);
  const traces = [
    {{x, y: rows.map(row => row.ocimf_current_moment_n_yaw_bow_port_kN_m), name:'Current (OCIMF direct)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Current N=%{{y:.0f}} kN·m<br>%{{customdata}}<extra></extra>'}},
    {{x, y: rows.map(row => row.moment_n_yaw_bow_port_kN_m), name:'Rudder', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Rudder N=%{{y:.0f}} kN·m<br>%{{customdata}}<extra></extra>'}},
    {{x, y: rows.map(row => row.total_moment_n_yaw_bow_port_kN_m), name:'Total (current + rudder)', mode:'lines+markers', customdata:hover, hovertemplate:'ψ=%{{x}}°<br>Total N=%{{y:.0f}} kN·m<br>%{{customdata}}<extra></extra>'}}
  ];
  Plotly.newPlot('ocimf-rudder-component-yaw-chart', traces, {{title:`Yaw moment N about CoG (kN·m, +bow-to-port) — current vs rudder vs total · V=${{speed}} kn · δ=${{rudder}}°`, xaxis:{{title:'Heading offset ψ (deg)', zeroline:true, gridcolor:'#e5eaf1'}}, yaxis:{{title:'Yaw moment N about CoG (kN·m, +bow-to-port)', zeroline:true, gridcolor:'#e5eaf1'}}, height:STANDARD_CHART_HEIGHT, margin:{{l:82,r:30,t:66,b:64}}, legend:{{orientation:'h', y:-0.22}}, hovermode:'x unified', template:'plotly_white'}}, CHART_CONFIG);
}}
function fmtKn(value) {{
  // Pass D: round kN display to 0 decimals when |val|>=1, 2 decimals otherwise
  const kn = value / 1000;
  return Math.abs(kn) >= 1 ? kn.toFixed(0) : kn.toFixed(2);
}}
function fmtMoment(value) {{
  // Pass D: round kN·m display to 0 decimals when |val|>=1, 2 decimals otherwise
  return Math.abs(value) >= 1 ? value.toFixed(0) : value.toFixed(2);
}}
function fmtAngle(value) {{ return value.toFixed(1); }}
function setSvgRotation(id, angleDeg, cx, cy) {{ document.getElementById(id).setAttribute('transform', `rotate(${{angleDeg}} ${{cx}} ${{cy}})`); }}
function updateHeadingRudderSchematic(row) {{
  const psi = row.heading_offset_deg;
  const delta = row.rudder_angle_deg;
  const alpha = row.effective_rudder_inflow_angle_deg;
  const totalForceAngle = Math.atan2(row.total_force_y_ship_port_N, row.total_force_x_ship_N) * 180 / Math.PI;
  setSvgRotation('schematic-current-heading-line', -psi, 310, 185);
  setSvgRotation('schematic-rudder-line', -delta, 310, 268);
  setSvgRotation('schematic-total-force-line', -totalForceAngle + 50, 310, 185);
  document.getElementById('schematic-rudder-blade').setAttribute('transform', `rotate(${{-delta}} 310 299)`);
  updateText('schematic-psi-label', `ψ = ${{psi}}°`);
  updateText('schematic-delta-label', `δ = ${{delta}}°`);
  updateText('schematic-alpha-label', `α = ${{alpha}}°`);
  updateText('schematic-speed-readout', `${{row.current_speed_kn}} kn`);
  updateText('schematic-heading-readout', `${{psi}} deg`);
  updateText('schematic-rudder-readout', `${{delta}} deg`);
  updateText('schematic-alpha-readout', `${{alpha}} deg`);
}}
function updateSelectedCaseBreakdown() {{
  const rows = selectedTraceRows();
  const selected = absMax(rows, 'total_moment_n_yaw_bow_port_kN_m');
  updateText('breakdown-current-x', fmtKn(selected.ocimf_current_force_x_ship_N));
  updateText('breakdown-current-y', fmtKn(selected.ocimf_current_force_y_ship_port_N));
  updateText('breakdown-current-n', fmtMoment(selected.ocimf_current_moment_n_yaw_bow_port_kN_m));
  updateText('breakdown-rudder-x', fmtKn(selected.force_x_ship_N));
  updateText('breakdown-rudder-y', fmtKn(selected.force_y_ship_port_N));
  updateText('breakdown-rudder-n', fmtMoment(selected.moment_n_yaw_bow_port_kN_m));
  updateText('breakdown-total-x', fmtKn(selected.total_force_x_ship_N));
  updateText('breakdown-total-y', fmtKn(selected.total_force_y_ship_port_N));
  updateText('breakdown-total-n', fmtMoment(selected.total_moment_n_yaw_bow_port_kN_m));
  updateHeadingRudderSchematic(selected);
  updateText('breakdown-case-label', `Selected breakdown case: V=${{selected.current_speed_kn}} kn · ${{caseLabel(selected)}} · total yaw reaction for mooring = ${{fmtMoment(-selected.total_moment_n_yaw_bow_port_kN_m)}} kN·m.`);
}}
function updateRudderForceVsAngleChart() {{
  // Pass H-2: rudder X_rudder and Y_rudder vs rudder angle δ from -28° to +28°
  // at fixed default heading ψ=+5° and selected current speed.
  const speed = selectedSpeed();
  const defaultHeading = 5.0;
  const defaultRudder = 28.0;
  const rows = ROWS
    .filter(row => same(row.current_speed_kn, speed) && same(row.heading_offset_deg, defaultHeading))
    .sort((a, b) => a.rudder_angle_deg - b.rudder_angle_deg);
  const x = rows.map(row => row.rudder_angle_deg);
  const xr = rows.map(row => row.force_x_ship_N / 1000);
  const yr = rows.map(row => row.force_y_ship_port_N / 1000);
  const selectedRow = rows.find(row => same(row.rudder_angle_deg, defaultRudder));
  const traces = [
    {{x, y: xr, name: 'X_rudder longitudinal (kN)', mode: 'lines+markers', line: {{color: '#155e95', width: 2.5}}, hovertemplate: 'δ=%{{x}}°<br>X_rudder=%{{y:.0f}} kN<extra></extra>'}},
    {{x, y: yr, name: 'Y_rudder transverse (kN, port)', mode: 'lines+markers', line: {{color: '#c2410c', width: 2.5}}, hovertemplate: 'δ=%{{x}}°<br>Y_rudder=%{{y:.0f}} kN<extra></extra>'}}
  ];
  if (selectedRow) {{
    traces.push({{
      x: [defaultRudder, defaultRudder],
      y: [selectedRow.force_x_ship_N / 1000, selectedRow.force_y_ship_port_N / 1000],
      name: `Selected δ = +${{defaultRudder}}°`,
      mode: 'markers',
      marker: {{color: '#7c3aed', size: 14, symbol: 'star'}},
      hovertemplate: 'Selected δ=%{{x}}°<br>%{{y:.0f}} kN<extra></extra>'
    }});
  }}
  Plotly.newPlot('rudder-force-vs-angle-chart', traces, {{
    title: `Rudder force vs rudder angle (Whicker-Fehlner [3]) · ψ=+${{defaultHeading}}°, V=${{speed}} kn`,
    xaxis: {{title: 'Rudder angle δ (deg)', dtick: 4, gridcolor: '#e5eaf1', zeroline: true}},
    yaxis: {{title: 'Rudder force component (kN)', gridcolor: '#e5eaf1', zeroline: true}},
    height: STANDARD_CHART_HEIGHT, margin: {{l: 82, r: 30, t: 66, b: 64}},
    legend: {{orientation: 'h', y: -0.22}}, template: 'plotly_white', hovermode: 'x unified'
  }}, CHART_CONFIG);
}}
function updateCurrentTransverseForceChart() {{
  // H2: Transverse current force Yc vs heading sweep at selected speed and default rudder.
  const speed = selectedSpeed();
  const defaultRudder = 28.0;
  const rows = ROWS
    .filter(row => same(row.current_speed_kn, speed) && same(row.rudder_angle_deg, defaultRudder))
    .sort(byHeading);
  const x = rows.map(row => row.heading_offset_deg);
  const trace = {{
    x, y: rows.map(row => row.ocimf_current_force_y_ship_port_N / 1000),
    name: 'Yc transverse current force (kN, port)',
    mode: 'lines+markers',
    line: {{color: '#155e95', width: 2.5}},
    hovertemplate: 'ψ=%{{x}}°<br>Yc=%{{y:.0f}} kN<extra></extra>'
  }};
  Plotly.newPlot('current-transverse-force-chart', [trace], {{
    title: `Transverse current force Yc vs heading (OCIMF MEG4 [1] Annex A Fig. A10) · V=${{speed}} kn`,
    xaxis: {{title: 'Heading offset ψ (deg)', dtick: 1, gridcolor: '#e5eaf1', zeroline: true}},
    yaxis: {{title: 'Transverse current force Yc (kN, port)', gridcolor: '#e5eaf1'}},
    height: STANDARD_CHART_HEIGHT, margin: {{l: 82, r: 30, t: 66, b: 64}},
    template: 'plotly_white', hovermode: 'x unified'
  }}, CHART_CONFIG);
}}
function updateSensitivityCurrentLoadChart() {{
  // H5: Current load sensitivity to current speed (0..5 kn) at default heading + rudder.
  const defaultHeading = 5.0;
  const defaultRudder = 28.0;
  const selectedV = selectedSpeed();
  // Average across rudder sweep at default heading to get speed-only dependence (all rudder
  // angles produce the same OCIMF current load — current is rudder-independent).
  const rowsByV = new Map();
  ROWS.forEach(row => {{
    if (same(row.heading_offset_deg, defaultHeading)) {{
      if (!rowsByV.has(row.current_speed_kn)) rowsByV.set(row.current_speed_kn, row);
    }}
  }});
  const speeds = [...rowsByV.keys()].sort((a, b) => a - b);
  const xc = speeds.map(v => rowsByV.get(v).ocimf_current_force_x_ship_N / 1000);
  const yc = speeds.map(v => rowsByV.get(v).ocimf_current_force_y_ship_port_N / 1000);
  const selectedRow = rowsByV.get(selectedV);
  const traces = [
    {{x: speeds, y: xc, name: 'Xc longitudinal (kN)', mode: 'lines+markers', line: {{color: '#155e95', width: 2.5}}, hovertemplate: 'V=%{{x}} kn<br>Xc=%{{y:.0f}} kN<extra></extra>'}},
    {{x: speeds, y: yc, name: 'Yc transverse (kN, port)', mode: 'lines+markers', line: {{color: '#c2410c', width: 2.5}}, hovertemplate: 'V=%{{x}} kn<br>Yc=%{{y:.0f}} kN<extra></extra>'}},
  ];
  if (selectedRow) {{
    traces.push({{
      x: [selectedV, selectedV],
      y: [selectedRow.ocimf_current_force_x_ship_N / 1000, selectedRow.ocimf_current_force_y_ship_port_N / 1000],
      name: `Selected V = ${{selectedV}} kn`,
      mode: 'markers',
      marker: {{color: '#7c3aed', size: 14, symbol: 'star'}},
      hovertemplate: 'Selected V=%{{x}} kn<br>%{{y:.0f}} kN<extra></extra>'
    }});
  }}
  Plotly.newPlot('sensitivity-current-load-chart', traces, {{
    title: `Current load sensitivity to current speed (OCIMF MEG4 [1]) · ψ=+${{defaultHeading}}°, δ=+${{defaultRudder}}° fixed`,
    xaxis: {{title: 'Current speed V (kn)', gridcolor: '#e5eaf1', zeroline: true}},
    yaxis: {{title: 'Current force component (kN)', gridcolor: '#e5eaf1', zeroline: true}},
    height: STANDARD_CHART_HEIGHT, margin: {{l: 82, r: 30, t: 66, b: 64}},
    legend: {{orientation: 'h', y: -0.22}}, template: 'plotly_white', hovermode: 'x unified'
  }}, CHART_CONFIG);
}}
function updateSensitivityRudderLoadChart() {{
  // H5: Rudder load sensitivity to current speed (0..5 kn) at default heading + rudder.
  const defaultHeading = 5.0;
  const defaultRudder = 28.0;
  const selectedV = selectedSpeed();
  const rows = ROWS
    .filter(row => same(row.heading_offset_deg, defaultHeading) && same(row.rudder_angle_deg, defaultRudder))
    .sort((a, b) => a.current_speed_kn - b.current_speed_kn);
  const speeds = rows.map(row => row.current_speed_kn);
  const xr = rows.map(row => row.force_x_ship_N / 1000);
  const yr = rows.map(row => row.force_y_ship_port_N / 1000);
  const selectedRow = rows.find(row => same(row.current_speed_kn, selectedV));
  const traces = [
    {{x: speeds, y: xr, name: 'X_rudder longitudinal (kN)', mode: 'lines+markers', line: {{color: '#155e95', width: 2.5}}, hovertemplate: 'V=%{{x}} kn<br>X_rudder=%{{y:.0f}} kN<extra></extra>'}},
    {{x: speeds, y: yr, name: 'Y_rudder transverse (kN, port)', mode: 'lines+markers', line: {{color: '#c2410c', width: 2.5}}, hovertemplate: 'V=%{{x}} kn<br>Y_rudder=%{{y:.0f}} kN<extra></extra>'}},
  ];
  if (selectedRow) {{
    traces.push({{
      x: [selectedV, selectedV],
      y: [selectedRow.force_x_ship_N / 1000, selectedRow.force_y_ship_port_N / 1000],
      name: `Selected V = ${{selectedV}} kn`,
      mode: 'markers',
      marker: {{color: '#7c3aed', size: 14, symbol: 'star'}},
      hovertemplate: 'Selected V=%{{x}} kn<br>%{{y:.0f}} kN<extra></extra>'
    }});
  }}
  Plotly.newPlot('sensitivity-rudder-load-chart', traces, {{
    title: `Rudder load sensitivity to current speed (Whicker-Fehlner [3]) · ψ=+${{defaultHeading}}°, δ=+${{defaultRudder}}° fixed`,
    xaxis: {{title: 'Current speed V (kn)', gridcolor: '#e5eaf1', zeroline: true}},
    yaxis: {{title: 'Rudder force component (kN)', gridcolor: '#e5eaf1', zeroline: true}},
    height: STANDARD_CHART_HEIGHT, margin: {{l: 82, r: 30, t: 66, b: 64}},
    legend: {{orientation: 'h', y: -0.22}}, template: 'plotly_white', hovermode: 'x unified'
  }}, CHART_CONFIG);
}}
function updateYawSideBySideChart() {{
  // Pass E: Method A (OCIMF direct Cxyc) vs Method B (Y_current × lever_arm)
  // across heading sweep, at default rudder angle for stability.
  const speed = selectedSpeed();
  const defaultRudder = 28.0;
  const sweepRows = ROWS
    .filter(row => same(row.current_speed_kn, speed) && same(row.rudder_angle_deg, defaultRudder))
    .sort(byHeading);
  const x = sweepRows.map(row => row.heading_offset_deg);
  const traces = [
    {{x, y: sweepRows.map(row => row.ocimf_current_moment_n_yaw_bow_port_kN_m), name: 'Method A: OCIMF direct Cxyc', mode: 'lines+markers', line: {{color: '#155e95'}}, hovertemplate: 'ψ=%{{x}}°<br>Method A N=%{{y:.0f}} kN·m<extra></extra>'}},
    {{x, y: sweepRows.map(row => row.current_y_times_lever_arm_moment_n_yaw_bow_port_kN_m), name: 'Method B: Y_current × ~0.3·LBP lever arm', mode: 'lines+markers', line: {{color: '#7c3aed', dash: 'dash'}}, hovertemplate: 'ψ=%{{x}}°<br>Method B N=%{{y:.0f}} kN·m<extra></extra>'}}
  ];
  Plotly.newPlot('yaw-side-by-side-chart', traces, {{
    title: `Current yaw moment about CoG — Method A vs Method B (not equality-based) · V=${{speed}} kn · δ=${{defaultRudder}}°`,
    xaxis: {{title: 'Heading offset ψ (deg)', zeroline: true, gridcolor: '#e5eaf1'}},
    yaxis: {{title: 'Current yaw moment N about CoG (kN·m, +bow-to-port)', zeroline: true, gridcolor: '#e5eaf1'}},
    height: STANDARD_CHART_HEIGHT,
    margin: {{l: 82, r: 30, t: 66, b: 64}},
    legend: {{orientation: 'h', y: -0.22}},
    hovermode: 'x unified',
    template: 'plotly_white'
  }}, CHART_CONFIG);
}}
function updateRudderModelComparison() {{
  // Default-case static table values (rendered into the table cells)
  const speed = selectedSpeed();
  const defaultHeading = 5.0;
  const defaultRudder = 28.0;
  const defaultCaseRow = ROWS.find(row =>
    same(row.current_speed_kn, speed)
    && same(row.heading_offset_deg, defaultHeading)
    && same(row.rudder_angle_deg, defaultRudder)
  );
  if (defaultCaseRow) {{
    // Model A (Whicker-Fehlner) values
    updateText('rudder-comp-a-normal', defaultCaseRow.normal_force_N.toFixed(1));
    updateText('rudder-comp-a-x', fmtKn(defaultCaseRow.force_x_ship_N));
    updateText('rudder-comp-a-y', fmtKn(defaultCaseRow.force_y_ship_port_N));
    updateText('rudder-comp-a-n', fmtMoment(defaultCaseRow.moment_n_yaw_bow_port_kN_m));
    // Model B (thin-plate) values
    updateText('rudder-comp-b-normal', defaultCaseRow.simple_plate_normal_force_N.toFixed(1));
    updateText('rudder-comp-b-x', fmtKn(defaultCaseRow.simple_plate_force_x_ship_N));
    updateText('rudder-comp-b-y', fmtKn(defaultCaseRow.simple_plate_force_y_ship_port_N));
    updateText('rudder-comp-b-n', fmtMoment(defaultCaseRow.simple_plate_moment_n_yaw_bow_port_kN_m));
  }}
  // Chart: rudder sweep at fixed heading=+5° and selected speed
  const sweepRows = ROWS
    .filter(row => same(row.current_speed_kn, speed) && same(row.heading_offset_deg, defaultHeading))
    .sort((a, b) => a.rudder_angle_deg - b.rudder_angle_deg);
  const x = sweepRows.map(row => row.rudder_angle_deg);
  const traces = [
    {{x, y: sweepRows.map(row => row.force_y_ship_port_N / 1000), name: 'Model A: Whicker-Fehlner Y (kN, port)', mode: 'lines+markers', line: {{color: '#155e95'}}, hovertemplate: 'δ=%{{x}}°<br>Model A Y=%{{y:.2f}} kN<extra></extra>'}},
    {{x, y: sweepRows.map(row => row.simple_plate_force_y_ship_port_N / 1000), name: 'Model B: thin-plate (Faltinsen §6.5) Y (kN, port)', mode: 'lines+markers', line: {{color: '#c2410c', dash: 'dash'}}, hovertemplate: 'δ=%{{x}}°<br>Model B Y=%{{y:.2f}} kN<extra></extra>'}}
  ];
  Plotly.newPlot('rudder-model-comparison-chart', traces, {{
    title: `Rudder model comparison — Y force vs rudder angle · ψ=+${{defaultHeading}}° · V=${{speed}} kn`,
    xaxis: {{title: 'Rudder angle δ (deg)', gridcolor: '#e5eaf1'}},
    yaxis: {{title: 'Transverse Y_rudder (kN, port)', gridcolor: '#e5eaf1', zeroline: true}},
    height: STANDARD_CHART_HEIGHT,
    margin: {{l: 82, r: 30, t: 66, b: 64}},
    legend: {{orientation: 'h', y: -0.22}},
    hovermode: 'x unified',
    template: 'plotly_white'
  }}, CHART_CONFIG);
}}
function updateCharts() {{ updateSelectedSpeedEnvelope(); updateRudderForceVsAngleChart(); updateForceChart(); updateYawChart(); updateCurrentTransverseForceChart(); updateComponentForceChart(); updateComponentYawChart(); updateYawSideBySideChart(); updateSelectedCaseBreakdown(); updateSensitivityCurrentLoadChart(); updateSensitivityRudderLoadChart(); updateRudderModelComparison(); }}
document.getElementById('current-speed-select').addEventListener('change', updateCharts);
document.getElementById('rudder-angle-select').addEventListener('change', () => {{ updateForceChart(); updateComponentForceChart(); updateComponentYawChart(); updateSelectedCaseBreakdown(); }});
updateCharts();
</script>
</body>
</html>
"""


def _speed_option(speed: float, default_speed: float) -> str:
    selected = " selected" if speed == default_speed else ""
    suffix = " (chart default extra)" if speed == default_speed else ""
    return f'<option value="{speed:g}"{selected}>{speed:g} kn{suffix}</option>'


def _design_rows_plain(design_data: dict[str, Any]) -> list[tuple[str, str]]:
    return [
        (str(key), str(value))
        for key, value in design_data.items()
        if not isinstance(value, dict)
    ]


def _design_rows_markdown(design_data: dict[str, Any]) -> list[str]:
    rows = ["| Field | Value |", "|---|---:|"]
    for key, value in design_data.items():
        if isinstance(value, dict):
            continue
        rows.append(f"| `{key}` | `{value}` |")
    return rows


def _case_label(row: dict[str, Any]) -> str:
    return (
        f"current {row['current_speed_kn']:g} kn, heading {row['heading_offset_deg']:g} deg, "
        f"rudder {row['rudder_angle_deg']:g} deg"
    )


def _positive(name: str, value: Any) -> float:
    number = float(value)
    if number <= 0:
        raise ValueError(f"{name} must be positive")
    return number


def _nonnegative(name: str, value: Any) -> float:
    number = float(value)
    if number < 0:
        raise ValueError(f"{name} must be nonnegative")
    return number
